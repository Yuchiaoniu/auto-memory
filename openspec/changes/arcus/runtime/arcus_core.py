# -*- coding: utf-8 -*-
# ⇢ 維護入口：完整功能地圖與修改狀態機見同目錄 arcus_workspace.md（第一輪先完整讀那份）
"""
arcus_core.py — Arcus 後端核心（四支合體，單檔零互相 import）
====================================================================
使用者指示（走甲）：不要 import、不要子行程、全部合併為一個檔。
本檔把原本核心 .py 的函式體搬進同一個模組，去掉它們之間的
互相 import，命名衝突逐一收斂。

現行來源（2026-07-12 合併；2026-07-13 移除已掏空的第三、四群殘留，見下方「已移除」）：
  1. write_log.py       中央 log 產生器（JSONL 原子追加、永不拋例外）
  2. context_meter.py   est_tokens 估算（中日韓×0.6＋其餘÷4）
  3. smart_brake.py     判官 turn_review／build_turn_review（斷點煞車機制已移除）
  4. arcus_hooks.py     before_prompt_hook／after_response_hook 等伺服器掛鉤
另含跨檔基礎設施：紀錄檔封存（archive_token_log_*）與結構地圖（generate_project_map）。

已移除（2026-07-13）：weight_baker.py（反向傳播 bake()）與 routing_law.py（法條層
run()／scan()）兩群，先前版本已被掏空為空白殘留、本檔無實作、server.py 也未 import。

設計原則不變：代碼（.py）＝永不改變的物理定律。本檔已無 bake()，
DNA_weights.json 現為靜態快照、不再由本檔改寫。

── 函式索引表 FUNCTION_INDEX ──────────────────────────────────────────
把「每個動作 → 對應函式」列成一張表（使用者要的快取索引值＋簡易說明），
讓呼叫端知道要調用哪一個函式，直接在本檔內操作，不必翻整份原始碼。
索引內容見本檔尾端 FUNCTION_INDEX 字典（可 `python3 arcus_core.py --index` 印出）。
"""

import os
import re
import sys
import json
import math
import time
import fcntl
import base64
import hashlib
import datetime
import tempfile
import threading
import subprocess

# ── 全域路徑常數（合併自六支各自的路徑定義，去重）──────────────────────────────
_CLAUDE_HOME = os.path.expanduser("~/.claude")
CLAUDE_HOME = _CLAUDE_HOME                       # routing_law／weight_baker 舊名相容
TOOLS_DIR = os.path.dirname(os.path.abspath(__file__))
_ARCUS_PATH = TOOLS_DIR                          # arcus_hooks 舊名相容
ARCUS_DIR = TOOLS_DIR                            # context_meter 舊名相容
CHANGES = os.path.join(_CLAUDE_HOME, 'openspec', 'changes')

DNA_WEIGHTS_PATH = os.path.join(TOOLS_DIR, "DNA_weights.json")
_DNA_WEIGHTS_PATH = DNA_WEIGHTS_PATH             # arcus_hooks 舊名相容
ERROR_TENSORS_PATH = os.path.join(TOOLS_DIR, "error_tensors.json")
COURT_CASES_PATH = os.path.join(TOOLS_DIR, "court-cases.json")
PATH_INDEX_PATH = os.path.join(TOOLS_DIR, "path-index.json")
SESSION_WRITES_PATH = os.path.join(TOOLS_DIR, "session-writes.json")
PENDING_WARNINGS_PATH = os.path.join(TOOLS_DIR, "pending-warnings.json")
PATH_INJECT_CACHE_PATH = os.path.join(TOOLS_DIR, "path-inject-cache.json")
PROJECT_TREE_CACHE_PATH = os.path.join(TOOLS_DIR, "project-tree-cache.json")
LOCK_FILE = "/tmp/arcus-baker.lock"


# ══════════════════════════════════════════════════════════════════════════════
# 區塊 1／4：中央 log 產生器（原 write_log.py）
#   公開介面 write_log／read_log 保留原名，其餘檔的函式都靠這兩支。
# ══════════════════════════════════════════════════════════════════════════════

_WL_ATOMIC_LIMIT = 4096  # 原 write_log._ATOMIC_LIMIT


def _wl_now():
    """統一時間戳，精確到秒（原 write_log._now）。"""
    return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())


def _wl_err_path(record_file):
    d = os.path.dirname(os.path.abspath(record_file)) or "."
    return os.path.join(d, "write_log_errors.log")


def _wl_background_error(record_file, msg):
    """背景記錯，本身也永不拋例外（原 write_log._background_error）。"""
    try:
        line = "%s\t%s\t%s\n" % (_wl_now(), record_file, msg)
        data = line.encode("utf-8", "replace")
        fd = os.open(_wl_err_path(record_file), os.O_WRONLY | os.O_CREAT | os.O_APPEND, 0o644)
        try:
            os.write(fd, data)
        finally:
            os.close(fd)
    except Exception:
        pass


def _wl_raw_append(record_file, text):
    """退化路徑：純文字 O_APPEND 單次寫落地（原 write_log._raw_append）。"""
    try:
        data = text.encode("utf-8", "replace")
        fd = os.open(record_file, os.O_WRONLY | os.O_CREAT | os.O_APPEND, 0o644)
        try:
            os.write(fd, data)
        finally:
            os.close(fd)
        return True
    except Exception:
        return False


def write_log(record_file, content):
    """把 content 以 {time, content} 一行 JSON 追加到 record_file（原 write_log.write_log）。

    併發安全（不鎖檔、O_APPEND 原子追加）、永不拋例外。
    回傳 True＝JSONL 正常落地；False＝走了退化純文字路徑（已背景記錯）。
    """
    try:
        record = {"time": _wl_now(), "content": content}
        try:
            payload = json.dumps(record, ensure_ascii=False)
        except (TypeError, ValueError):
            payload = json.dumps(record, ensure_ascii=False, default=str)

        line = payload + "\n"
        data = line.encode("utf-8", "replace")

        if len(data) > _WL_ATOMIC_LIMIT:
            _wl_background_error(
                record_file,
                "line %d bytes > atomic limit %d, interleave not guaranteed"
                % (len(data), _WL_ATOMIC_LIMIT),
            )

        fd = os.open(record_file, os.O_WRONLY | os.O_CREAT | os.O_APPEND, 0o644)
        try:
            os.write(fd, data)
        finally:
            os.close(fd)
        return True
    except Exception as e:
        ok = _wl_raw_append(record_file, "%s\t%r\n" % (_wl_now(), content))
        _wl_background_error(record_file, "write_log failed(%s): %r" % ("raw_ok" if ok else "raw_fail", e))
        return False


def read_log(record_file, limit=None):
    """讀回 JSONL 記錄檔，回傳字典串列（每筆含 time／content）（原 write_log.read_log）。

    容忍壞行；limit 給定時只回最後 limit 筆。
    """
    out = []
    try:
        with open(record_file, "r", encoding="utf-8", errors="replace") as f:
            for raw in f:
                s = raw.rstrip("\n")
                if not s.strip():
                    continue
                try:
                    obj = json.loads(s)
                    if isinstance(obj, dict) and "content" in obj:
                        out.append(obj)
                    else:
                        out.append({"time": None, "content": obj})
                except Exception:
                    out.append({"time": None, "content": s, "_broken": True})
    except FileNotFoundError:
        return []
    except Exception:
        return out
    if limit is not None and limit >= 0:
        return out[-limit:]
    return out


# arcus_hooks 舊碼以 _write_log／_read_log 別名呼叫（原 try-import 別名），保留相容。
_write_log = write_log
_read_log = read_log


# ══════════════════════════════════════════════════════════════════════════════
# 區塊 2／4：token 估算（原 context_meter.py 的 est_tokens）
# ══════════════════════════════════════════════════════════════════════════════

def est_tokens(s):
    """估算 token 數：中日韓字×0.6＋其餘字元÷4。僅供排序抓大頭，非精算。"""
    cjk = sum(1 for c in s
              if '　' <= c <= '鿿' or '＀' <= c <= '￯')
    other = len(s) - cjk
    return round(cjk * 0.6 + other / 4)


# ══════════════════════════════════════════════════════════════════════════════
# 共用 JSON 讀寫（原 routing_law／weight_baker 各有一份 _read_json／_write_json，
# 語義幾乎相同：統一為一份。routing_law 版無 default、weight_baker 版有 default，
# 這裡合併成「有 default 參數」的超集，向下相容兩邊呼叫。）
# ══════════════════════════════════════════════════════════════════════════════

def _read_json(path, default=None):
    """讀取 JSON 檔，失敗時返回 default（預設 None，相容 routing_law 舊呼叫）。"""
    try:
        with open(path, encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return default


def _write_json(path, data):
    """寫入 JSON 檔（合併 routing_law／weight_baker 兩版，吞例外不害主流程）。"""
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════════════
# 區塊 3／4：判官（原 smart_brake.py）
#   現行只保留判官 turn_review／build_turn_review；斷點煞車 mark_turn_start／
#   read_checkpoint／run_smart_brake 已於先前版本移除，本群無煞車實作。
# ══════════════════════════════════════════════════════════════════════════════

TOOL_THRESHOLD_N = 15
REPEAT_TARGET_MIN = 2
TRACE_READ_ROUNDS = 5
ELAPSED_THRESHOLD_S = 240
VOLUME_THRESHOLD_CH = 40000
OPUS_MODEL = "claude-opus-4-8"

TURN_START_FILE = ".smart_brake_turn_start"
STATE_FILE = ".smart_brake_state.json"
CHECKPOINT_FILE = "CHECKPOINT.md"
ACTION_TRACE_FILE = "action_trace.md"
PENDING_FILE = "pending_suggestions.md"
CHECKPOINT_JSONL = "CHECKPOINT.jsonl"
ACTION_TRACE_JSONL = "action_trace.jsonl"
PENDING_JSONL = "pending_suggestions.jsonl"

JUDGE_INTERVAL_DEFAULT = 1
_JUDGE_ENABLED = False  # §關閉判官（2026-07-22 使用者要求）：關掉每輪 opus 複核，只留工作明細；改 True 即恢復
TURN_COUNTER_KEY = "turn_counter"
JUDGE_INTERVAL_KEY = "judge_interval"


def _ev_name(ev):
    return ev.get('name', ev.get('tool', '?'))


def _ev_input(ev):
    inp = ev.get('input', ev.get('content', ''))
    if isinstance(inp, dict) and inp:
        return inp
    if not inp:
        inp = ev.get('summary', '')
    return inp if isinstance(inp, dict) else {'_raw': inp}


def _ev_result(ev):
    r = ev.get('result', '')
    return str(r) if r else ''


















def _load_state(project_path):
    try:
        with open(os.path.join(project_path, STATE_FILE), encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return {}


def _save_state(project_path, state):
    try:
        with open(os.path.join(project_path, STATE_FILE), 'w', encoding='utf-8') as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f'[smart_brake] state 寫入失敗: {e}')




def _get_api_key():
    """回傳 (key, scheme)。scheme='apikey' 用 x-api-key；scheme='oauth' 用 Bearer。"""
    key = os.environ.get('ANTHROPIC_API_KEY', '')
    if key:
        return key, 'apikey'
    try:
        with open(os.path.join(_CLAUDE_HOME, '.credentials.json'), encoding='utf-8') as f:
            creds = json.load(f)
        oauth = creds.get('claudeAiOauth') or {}
        tok = oauth.get('accessToken') or creds.get('claudeAiOauthToken')
        if tok:
            return tok, 'oauth'
        apikey = creds.get('apiKey')
        if apikey:
            return apikey, 'apikey'
    except Exception:
        pass
    return '', ''


_LAST_OPUS_ERROR = ''


def _explain_opus_error(reason):
    r = reason or ''
    if '429' in r:
        return '被 Anthropic 伺服器限流（短時間內對同一帳號呼叫太多次），已自動重試仍被擋'
    if '529' in r or 'overload' in r.lower():
        return 'Anthropic 伺服器過載，已自動重試仍未成功'
    if '權杖' in r or '401' in r:
        return '登入權杖無效或找不到，無法呼叫'
    if 'timed out' in r.lower() or 'timeout' in r.lower():
        return '呼叫逾時（超過 60 秒沒回應），已自動重試仍未成功'
    return r or '未知原因'


def _call_opus(prompt, max_tokens=1024):
    """呼叫 opus 產生判官分析（走 claude --print --model opus 子行程管道）。"""
    global _LAST_OPUS_ERROR
    _LAST_OPUS_ERROR = ''
    import subprocess as _sp
    import tempfile as _tf
    import time as _time
    tmp_path = None
    last_err = ''
    try:
        with _tf.NamedTemporaryFile(
            mode='w', suffix='.txt', encoding='utf-8', delete=False
        ) as f:
            f.write(prompt)
            tmp_path = f.name
        script = (
            'source ~/.nvm/nvm.sh 2>/dev/null; '
            f'cat "{tmp_path}" | claude --print --model opus --dangerously-skip-permissions'
        )
        env = dict(os.environ)
        env['CLAUDE_SUBPROCESS'] = '1'
        env['PYTHONIOENCODING'] = 'utf-8'
        for attempt in range(3):
            if attempt:
                _time.sleep(3 * attempt)
            try:
                proc = _sp.run(
                    ['bash', '-lc', script],
                    stdout=_sp.PIPE, stderr=_sp.STDOUT,
                    text=True, encoding='utf-8', errors='replace',
                    env=env, timeout=120,
                )
                out = (proc.stdout or '').strip()
                if proc.returncode == 0 and out:
                    return out
                last_err = f'子行程回傳碼 {proc.returncode}；輸出：{out[:200] or "空"}'
            except _sp.TimeoutExpired:
                last_err = '呼叫逾時（claude --print 超過 120 秒沒回應）'
            except Exception as e:
                last_err = str(e)
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    _LAST_OPUS_ERROR = last_err
    print(f'[smart_brake] opus 子行程呼叫失敗（已重試 3 次）: {last_err}')
    return None


def _tool_digest(tool_events, limit=40):
    out = []
    for ev in (tool_events or [])[:limit]:
        name = _ev_name(ev)
        inp = _ev_input(ev)
        s = inp.get('_raw') if set(inp.keys()) == {'_raw'} else json.dumps(inp, ensure_ascii=False)
        res = _ev_result(ev)
        line = f"- {name}: {str(s)[:200]}"
        if res:
            line += f"\n    → 結果: {res[:300]}"
        out.append(line)
    return '\n'.join(out)














def _short_summary(ev):
    name = _ev_name(ev)
    inp = _ev_input(ev)
    if name == 'Bash':
        what = str(inp.get('command', inp.get('_raw', '')))
    elif name in ('Read', 'Write', 'Edit', 'NotebookEdit'):
        what = str(inp.get('file_path', ''))
    elif name == 'Grep':
        what = str(inp.get('pattern', ''))
    elif name == 'Skill':
        what = str(inp.get('command', inp.get('skill', inp.get('_raw', ''))))
    else:
        what = inp.get('_raw') if set(inp.keys()) == {'_raw'} else json.dumps(inp, ensure_ascii=False)
        what = str(what)
    what = re.sub(r'\s+', ' ', what).strip()
    if len(what) > 90:
        what = what[:90] + '…'
    res = _ev_result(ev)
    if res:
        res1 = re.sub(r'\s+', ' ', res).strip()
        if res1:
            what += ' → ' + (res1[:70] + ('…' if len(res1) > 70 else ''))
    return what


def build_turn_report(user_msg, tool_events, token_stats, round_notes=None):
    """把本輪逐工具的時間／token 拆帳成可讀報表（純 Python，永遠可生）。"""
    tool_events = tool_events or []
    round_notes = round_notes or {}
    n = len(tool_events)
    ts = token_stats or {}
    total_out = ts.get('output_tokens', 0)
    total_in = ts.get('input_tokens', 0)
    cache_r = ts.get('cache_read_input_tokens', 0)
    cost = ts.get('cost', 0) or 0
    dur_s = (ts.get('duration_ms', 0) or 0) / 1000.0
    rounds = sorted({ev.get('round') for ev in tool_events if ev.get('round') is not None})
    tool_rounds = len(rounds)
    round_start = {}
    for _r, _note in round_notes.items():
        if isinstance(_r, int) and isinstance((_note or {}).get('t'), (int, float)):
            round_start[_r] = _note['t']
    for _ev in tool_events:
        _r, _t0 = _ev.get('round'), _ev.get('t')
        if isinstance(_r, int) and isinstance(_t0, (int, float)):
            round_start.setdefault(_r, _t0)
    _note_rounds = [r for r in round_notes if isinstance(r, int)]
    max_r = max(list(rounds) + _note_rounds) if (rounds or _note_rounds) else 0
    r_cnt = (max_r if max_r else '-')

    def _round_dur(r):
        t0 = round_start.get(r)
        if not isinstance(t0, (int, float)):
            return None
        laters = [round_start[x] for x in round_start if x > r]
        end = min(laters) if laters else (dur_s if dur_s else None)
        if not isinstance(end, (int, float)) or end < t0:
            return None
        return end - t0

    head = (f"## 📊 本輪工作明細（時間與花費拆帳）\n"
            f"總計：工具 {n} 次 · 回合 {r_cnt}（其中 {tool_rounds} 個回合動了工具） · "
            f"產出 {total_out} tok（輸入 {total_in}／快取讀 {cache_r}） · "
            f"耗時 {dur_s:.1f}s · ${cost:.4f}\n")
    if n == 0:
        think_rows = []
        for _r in sorted(x for x in round_notes if isinstance(x, int)):
            _note = round_notes.get(_r) or {}
            _think = re.sub(r'\s+', ' ', str(_note.get('thinking', ''))).strip()
            if _think:
                think_rows.append(f"  · R{_r} 🧠思考：" + _think[:80] + ('…' if len(_think) > 80 else ''))
        if think_rows:
            return (head + "\n（本輪未呼叫任何工具，純對話回覆；以下為各回合思考草稿）\n"
                    + "\n".join(think_rows))
        return head + "\n（本輪未呼叫任何工具，純對話回覆。）"
    lines = ["", "  #  回合  工具        起始     耗時    做了什麼",
             "  ─  ──  ──────────  ───────  ──────  ────────────────────────"]

    def _tool_row(idx, rd, ev):
        name = _ev_name(ev)
        rd_s = f"R{rd}" if rd is not None else "-"
        t0 = ev.get('t')
        t0_s = f"+{t0:.1f}s" if isinstance(t0, (int, float)) else "-"
        du = ev.get('dur')
        du_s = f"{du:.1f}s" if isinstance(du, (int, float)) else "-"
        return f"  {idx:<2} {rd_s:<3} {name:<10}  {t0_s:<7}  {du_s:<6}  {_short_summary(ev)}"

    by_round = {}
    for ev in tool_events:
        by_round.setdefault(ev.get('round'), []).append(ev)
    i = 0
    for r in range(1, max_r + 1):
        evs = by_round.get(r) or []
        if evs:
            for ev in evs:
                i += 1
                lines.append(_tool_row(i, r, ev))
        else:
            note = round_notes.get(r) or {}
            t0 = note.get('t')
            t0_s = f"+{t0:.1f}s" if isinstance(t0, (int, float)) else "-"
            du = _round_dur(r)
            du_s = f"{du:.1f}s" if isinstance(du, (int, float)) else "-"
            txt = re.sub(r'\s+', ' ', str(note.get('text', ''))).strip()
            think = re.sub(r'\s+', ' ', str(note.get('thinking', ''))).strip()
            if txt:
                what = txt[:60] + ('…' if len(txt) > 60 else '')
            elif think:
                what = '🧠思考：' + think[:60] + ('…' if len(think) > 60 else '')
            elif note.get('had_thinking'):
                what = '（工具間思考步驟：命令列未提供思考明文，僅回簽章）'
            else:
                what = '（這一回合只思考、沒有輸出文字，也沒有呼叫工具）'
            lines.append(f"  ·  R{r:<2} {'（只思考）':<10}  {t0_s:<7}  {du_s:<6}  {what}")
    for ev in by_round.get(None, []):
        i += 1
        lines.append(_tool_row(i, None, ev))
    per_round = {}
    for ev in tool_events:
        rd = ev.get('round')
        ot = ev.get('out_tokens')
        if rd is not None and isinstance(ot, (int, float)):
            per_round[rd] = ot
    if per_round:
        seg = ' / '.join(f"R{r}:{per_round[r]}tok" for r in sorted(per_round))
        lines.append("")
        lines.append(f"各回合產出：{seg}")
    return head + '\n'.join(lines)


def _bump_turn_counter(project_path):
    st = _load_state(project_path)
    n = int(st.get(TURN_COUNTER_KEY, 0)) + 1
    st[TURN_COUNTER_KEY] = n
    _save_state(project_path, st)
    return n


def get_judge_interval(project_path):
    st = _load_state(project_path)
    try:
        return max(1, int(st.get(JUDGE_INTERVAL_KEY, JUDGE_INTERVAL_DEFAULT)))
    except Exception:
        return JUDGE_INTERVAL_DEFAULT


def set_judge_interval(project_path, n):
    st = _load_state(project_path)
    st[JUDGE_INTERVAL_KEY] = max(1, int(n))
    _save_state(project_path, st)
    return st[JUDGE_INTERVAL_KEY]


def _read_tail_chars(path, n):
    try:
        with open(path, encoding='utf-8', errors='replace') as f:
            data = f.read()
        return data[-n:] if len(data) > n else data
    except OSError:
        return ''


def _pending_jsonl_tail(project_path, n):
    recs = read_log(os.path.join(project_path, PENDING_JSONL))
    if not recs:
        return ''
    blocks = []
    for rec in recs:
        content = rec.get('content')
        if not isinstance(content, dict):
            continue
        turn = content.get('turn', '?')
        ts = rec.get('time', '')
        status = content.get('status', '待審')
        secs = content.get('sections', {}) or {}
        parts = [f'### {title}\n{body}' for title, body in secs.items() if body]
        if not parts:
            continue
        blocks.append(f'## 第 {turn} 輪 · {ts} · 狀態：{status}\n\n' + '\n\n'.join(parts))
    text = '\n\n---\n'.join(blocks)
    return text[-n:] if len(text) > n else text


def _build_known_context(project_path):
    """§60：組「判官已知脈絡」——已記錄的待審建議＋已知陷阱，餵給判官避免重講舊帳。"""
    if not project_path:
        return ''
    pend_md = _read_tail_chars(os.path.join(project_path, PENDING_FILE), 3500)
    pend_new = _pending_jsonl_tail(project_path, 3500)
    pend = (pend_md + ('\n' if pend_md and pend_new else '') + pend_new)
    traps = _read_tail_chars(os.path.join(project_path, '陷阱表.md'), 2000)
    if not pend and not traps:
        return ''
    block = ('\n【判官已知脈絡——以下為「已經記錄、視為已知」，不要重複提出實質相同的東西】\n')
    if pend:
        block += (f'▼ 已在待審清單 pending_suggestions.md（等人工審核，已提過就別再提）：\n{pend}\n')
    if traps:
        block += (f'▼ 已在陷阱表.md（已知的坑，別再當新發現重講）：\n{traps}\n')
    return block


def _advisory_judge(user_msg, tool_events, token_stats, project_path=None):
    """opus 判官（優化建議產生器）。opus 不可用時回傳降級提示，不阻斷報表。"""
    digest = _tool_digest(tool_events, limit=40)
    ts = token_stats or {}
    known = _build_known_context(project_path)
    prompt = f"""你是 Arcus 工具使用審查官暨優化建議產生器。審視主對話這一輪的工具使用，
給出精簡、可執行的評估。禁止空話。
{known}
本輪使用者請求：
{(user_msg or '')[:600]}

本輪工具呼叫（共 {len(tool_events or [])} 次；產出 {ts.get('output_tokens',0)} tok／耗時 {(ts.get('duration_ms',0) or 0)/1000:.1f}s）：
{digest}

用字要求：一律白話文，避免專業術語，絕對不使用英文縮寫；非用專有名詞不可時，第一次出現須用括號補白話解釋。禁止假設讀者懂專案內部黑話。凡是牽涉數值的地方，一律直接給實際數字。
只輸出 Markdown，五段、每段至多 3 條，務必具體、用真實數字：
## 本輪最花時間與花費的地方
（先做一次統合：這一整輪裡，哪一個回合、哪一次工具呼叫花掉最多時間或最多 token？用真實數字點名。接著解釋為什麼。最後給一個具體、能馬上做的解方。）
## 合理性
（這輪工具使用合不合理？有無重複／可省的呼叫？一句判定＋理由）
## 可預先完成的工作
（哪些呼叫其實可由腳本或快取一次做好；點名具體指令／檔）
## 建議新增
（建議新增什麼 pre-prompt 規則、或什麼快取檔／memory.md 條目。若與已知脈絡實質相同一律跳過；與已知全部重複就寫「無新增（既有待審已涵蓋）」）
## 框外發現
（跳出上述：這 N 輪有沒有任何可被機制化的重複模式？排除已知脈絡已提過的；真的沒有全新的才寫「無」）"""
    body = _call_opus(prompt, max_tokens=700)
    if not body:
        why = _explain_opus_error(_LAST_OPUS_ERROR)
        return (f"## 判官\n（這一輪判官分析暫時產不出來，原因：{why}。"
                f"上面的工作明細照常提供；下次判官到期時會再試一次。）")
    return body.strip()


def _extract_section(md, title):
    out = []
    capturing = False
    for ln in (md or '').splitlines():
        if ln.strip().startswith('## '):
            if capturing:
                break
            capturing = (ln.strip() == f'## {title}')
            continue
        if capturing:
            out.append(ln)
    return '\n'.join(out).strip()


def _log_pending_suggestions(project_path, advisory, turn_n):
    """§24.3 人工閘：把判官的『建議新增／框外發現』寫進待審清單，只記錄不自動生效。"""
    if not advisory:
        return
    parts = []
    for title in ('建議新增', '框外發現'):
        sec = _extract_section(advisory, title)
        if sec and sec.strip() not in ('無', '（無）', '(無)'):
            parts.append(f'### {title}\n{sec}')
    if not parts:
        return
    sections = {}
    for title in ('建議新增', '框外發現'):
        sec = _extract_section(advisory, title)
        if sec and sec.strip() not in ('無', '（無）', '(無)'):
            sections[title] = sec
    content = {
        'kind': 'pending_suggestion',
        'turn': turn_n, 'status': '待審', 'sections': sections,
    }
    write_log(os.path.join(project_path, PENDING_JSONL), content)


def turn_review(project_path, user_msg, response_text, tool_events, token_stats, round_notes=None):
    """一回二的第二則訊息內容：每輪工作明細（永遠生）＋到期時的 opus 判官解析。"""
    try:
        report = build_turn_report(user_msg, tool_events, token_stats, round_notes)
    except Exception as e:
        print(f'[smart_brake] 體帳生成失敗: {e}')
        return ''
    try:
        n = _bump_turn_counter(project_path)
        interval = get_judge_interval(project_path)
        if _JUDGE_ENABLED and n % interval == 0:
            advisory = _advisory_judge(user_msg, tool_events, token_stats, project_path)
            _log_pending_suggestions(project_path, advisory, n)
            if advisory:
                report += "\n\n" + advisory
                report += f"\n\n_（判官每 {interval} 輪一次；本輪第 {n} 輪觸發。穩定後可調回 5。）_"
        elif not _JUDGE_ENABLED:
            pass  # §判官已關閉（_JUDGE_ENABLED=False）：只保留工作明細，不叫 opus。改 True 即恢復。
        else:
            report += f"\n\n_（判官每 {interval} 輪一次；本輪第 {n} 輪，未到期。）_"
    except Exception as e:
        print(f'[smart_brake] 判官解析失敗: {e}')
    return report


# ══════════════════════════════════════════════════════════════════════════════
# 區塊 4／4：伺服器掛鉤（原 arcus_hooks.py）
#   before_prompt_hook／after_response_hook 及其依賴的地圖／log／token 工具。
#   原本靠 import smart_brake／symbol_index／context_meter／write_log，
#   合併後改成呼叫本檔同層函式，例如：
#     _smart_brake_judge.turn_review → turn_review
# ══════════════════════════════════════════════════════════════════════════════


# ──────────────────────────────────────────────────────────────────────────────
# 同檔 symbol_index（原 symbol_index.py，併入本檔基礎設施）
# 掃描兩支活核心檔（server.py、arcus_core.py），輸出帶行號函式索引供 before_prompt_hook
# 每輪注入穩定前綴，讓下一輪開場即見「哪個函式在哪一行」。mtime＋size 雜湊快取。
# ──────────────────────────────────────────────────────────────────────────────
import re as _re_sym
import hashlib as _hashlib_sym

_SYM_CORE_DIR = os.path.dirname(os.path.abspath(__file__))
_SYM_CORE_FILES = ['server.py', 'arcus_core.py']

_SYM_DEF_RE = _re_sym.compile(r'^(\s*)(?:async\s+)?def\s+(\w+)')
_SYM_CLASS_RE = _re_sym.compile(r'^(\s*)class\s+(\w+)')

# §60：記錄檔↔寫入函式靜態對照（合併後統一指向 arcus_core.py）
_RECORD_WRITER_MAP = (
    '記錄檔↔寫入函式對照（行號見上方索引；問「哪個記錄檔由哪個函式寫」看這段，勿重跑 grep）：\n'
    '  log.md               ← arcus_core.append_log_turn（每輪整段回應文字）\n'
    '  verbose_log.md       ← arcus_core._append_verbose_turn（§59 已停用寫入、函式保留可還原）\n'
    '  verbose_states.md    ← arcus_core._trim_verbose_log（verbose 逾 20 輪裁剪搬移目標）\n'
    '  token_log.md         ← arcus_core.append_token_log（每輪 token／快取逐回合）\n'
    '  token_log_archive.md ← arcus_core.archive_token_log_if_needed（兩天自動封存）\n'
    '  pending_suggestions.md ← arcus_core._log_pending_suggestions（判官待審、人工閘）\n'
    '  judge_log.md         ← server.py chat 處理（判官落地）\n'
    '  tasks-archive.md     ← 歸檔排程整批搬入已勾項目（§87 STATE.md 已廢除）\n'
    '  tasks.md             ← log_task.py（外部腳本，非核心檔）'
)

_sym_cache: dict = {}  # {'k': hash_key, 'v': index_text}


def _scan_file_sym(fpath: str) -> list:
    """回傳 [(lineno, kind, name, indent)]；kind 為 'class' 或 'def'。"""
    out = []
    try:
        with open(fpath, encoding='utf-8', errors='replace') as _f:
            for _n, _line in enumerate(_f, 1):
                _m = _SYM_CLASS_RE.match(_line)
                if _m:
                    out.append((_n, 'class', _m.group(2), len(_m.group(1))))
                    continue
                _m = _SYM_DEF_RE.match(_line)
                if _m:
                    out.append((_n, 'def', _m.group(2), len(_m.group(1))))
    except OSError:
        pass
    return out


def build_symbol_index(core_dir: str = None) -> str:
    """生成兩核心檔（server.py、arcus_core.py）帶行號函式／類別索引文字；mtime＋size 雜湊快取。"""
    _cdir = core_dir or _SYM_CORE_DIR
    _stat_sig = []
    for _name in _SYM_CORE_FILES:
        _fpath = os.path.join(_cdir, _name)
        try:
            _st = os.stat(_fpath)
            _stat_sig.append((_name, _st.st_mtime, _st.st_size))
        except OSError:
            _stat_sig.append((_name, 0, 0))
    _hash_key = _hashlib_sym.md5(repr(_stat_sig).encode()).hexdigest()

    if _sym_cache.get('k') == _hash_key:
        return _sym_cache['v']

    _blocks = []
    for _name in _SYM_CORE_FILES:
        _fpath = os.path.join(_cdir, _name)
        _syms = _scan_file_sym(_fpath)
        if not _syms:
            continue
        _lines = [f'{_name}:']
        for _n, _kind, _sym, _indent in _syms:
            if _indent == 0:
                _prefix = '  '
                _tag = 'class ' if _kind == 'class' else ''
            else:
                _prefix = '    ·'
                _tag = 'class ' if _kind == 'class' else ''
            _lines.append(f'{_prefix}{_tag}{_sym}  L{_n}')
        _blocks.append('\n'.join(_lines))

    _index_text = '\n\n'.join(_blocks) if _blocks else '（無核心檔可索引）'
    _index_text = f'{_index_text}\n\n{_RECORD_WRITER_MAP}'
    _sym_cache['k'] = _hash_key
    _sym_cache['v'] = _index_text
    return _index_text

ARCUS_ARCHITECTURE = """
## Arcus 系統架構（此訊息說明你所在的運作環境）

你正在 Project Arcus 系統中運作。每次對話都是一次全新的 claude --print 呼叫，沒有原生 session 記憶。
Arcus 透過手動拼接對話歷史（最近數輪）和注入上下文來補償這個限制。

### 部署環境

這個 webapp 名為 **Arcus**，部署在 GCP VM（主機名稱：`forest-carbon`，外部網址：`https://forest-carbon.duckdns.org/arcus/`，內部 port：7800）。
使用者透過瀏覽器連線到這個網址，與你對話。你的每一輪回應都是 server.py 以 `claude --print` 呼叫 Claude CLI 後，把輸出透過 SSE 串流回瀏覽器顯示給使用者。

### 元件說明

**server.py**（Flask HTTP server，port 7800）
- 接收瀏覽器的 /arcus/api/chat 請求，把對話歷史（最近數輪）＋系統架構＋當前專案的結構地圖與核心檔函式索引＋arcus_workspace.md 集中工作區，拼成一次性 prompt（打底瘦身後不再注入 STATE.md／tasks.md）
- 呼叫 claude --print --output-format stream-json，以 SSE 串流方式把回應送回瀏覽器
- 維護記憶體內對話歷史，重啟後從當日 log 檔恢復

**arcus_core.py**（本檔案：伺服器掛鉤＋arcus_do 九指令分派＋判官 turn_review＋log 產生器＋token 估算）
- before_prompt_hook()：每輪注入當前專案的結構地圖與核心檔函式索引，並附上 arcus_workspace.md
- after_response_hook()：無副作用空操作；每輪工具使用的複核改由 server.py 呼叫 turn_review，以 opus 判官執行

**arcus_workspace.md**（集中工作區，每輪注入）
- 維護九個 arcus_do 指令的 payload 形狀、draw 繪圖規格、狀態機與功能地圖；改碼前先讀

### 重要限制

- 你（Claude CLI）沒有原生記憶，每次呼叫都是全新的
- 對話歷史由 server.py 手動拼接後注入，只有最近數輪
- 程式碼（.py 檔案）是你的物理構造；你透過 arcus_do 的 read／stage／promote 指令改動它，而非直接編輯檔案
""".strip()

_map_cache = {}


def _get_or_create_py_desc(fpath):
    """讀取 .py 檔說明；沒有 docstring 就掃函式/類別名稱並補上樣板。"""
    try:
        with open(fpath, encoding='utf-8', errors='replace') as f:
            lines = f.readlines()
    except OSError:
        return '（無法讀取）'
    if not lines:
        return '（空檔案）'
    i = 0
    while i < min(3, len(lines)):
        s = lines[i].strip()
        if s.startswith('#!') or ('coding' in s and s.startswith('#')):
            i += 1
        else:
            break
    if i < len(lines):
        s = lines[i].strip()
        if s.startswith('"""') or s.startswith("'''"):
            desc = s.lstrip('"\'').rstrip('"\'').strip()
            if not desc and i + 1 < len(lines):
                desc = lines[i + 1].strip().rstrip('"\'').strip()
            return desc[:80] if desc else '（docstring 為空）'
        elif s.startswith('#') and not s.startswith('#!'):
            return s.lstrip('#').strip()[:80]
    funcs, classes = [], []
    for line in lines:
        m = re.match(r'^(?:async\s+)?def\s+(\w+)', line)
        if m and not m.group(1).startswith('_'):
            funcs.append(m.group(1))
        m = re.match(r'^class\s+(\w+)', line)
        if m:
            classes.append(m.group(1))
    parts = []
    if classes:
        parts.append('類別：' + ', '.join(classes[:3]))
    if funcs:
        parts.append('函式：' + ', '.join(funcs[:5]))
    if not parts:
        return '（無說明）'
    desc = '　'.join(parts)
    template = f'"""\n{desc}\n"""\n'
    new_lines = lines[:i] + [template] + lines[i:]
    try:
        with open(fpath, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)
    except OSError:
        pass
    return desc[:80]


def generate_project_map(project_path, project_name=None):
    """掃描專案目錄，生成帶路徑和說明的結構地圖文字（只展開最近 7 天異動路徑）。"""
    INCLUDE_EXT = {'.py', '.md', '.json', '.html', '.sh'}
    EXCLUDE_DIRS = {'.git', 'node_modules', '__pycache__', 'venv', '.venv', 'env'}
    file_entries = []
    for root, dirs, files in os.walk(project_path):
        dirs[:] = sorted(d for d in dirs if d not in EXCLUDE_DIRS)
        for fname in sorted(files):
            ext = os.path.splitext(fname)[1]
            if ext not in INCLUDE_EXT:
                continue
            fpath = os.path.join(root, fname)
            try:
                mtime = os.path.getmtime(fpath)
            except OSError:
                continue
            rel = os.path.relpath(fpath, project_path)
            file_entries.append((rel, mtime))
    hash_key = hashlib.md5(
        json.dumps(file_entries, ensure_ascii=False).encode()
    ).hexdigest()
    if project_path in _map_cache and _map_cache[project_path][0] == hash_key:
        return _map_cache[project_path][1]
    recent_threshold = time.time() - 7 * 24 * 3600
    lines = []
    for rel, mtime in file_entries:
        fpath = os.path.join(project_path, rel)
        ext = os.path.splitext(rel)[1]
        is_recent = mtime >= recent_threshold
        if is_recent:
            if ext == '.py':
                desc = _get_or_create_py_desc(fpath)
            elif ext == '.md':
                desc = '（無說明）'
                try:
                    with open(fpath, encoding='utf-8', errors='replace') as f:
                        for line in f:
                            s = line.strip()
                            if s:
                                desc = (s.lstrip('#').strip())[:80]
                                break
                except OSError:
                    desc = '（無法讀取）'
            else:
                desc = '（無說明）'
            lines.append(f'★ {rel:<43} → {desc}')
        else:
            lines.append(f'  {rel}')
    if not lines:
        map_text = '（目錄為空或無符合的檔案）'
    else:
        header = f'根目錄：{project_path}'
        if project_name:
            header = f'專案：{project_name}　{header}'
        map_text = header + '\n\n' + '\n'.join(lines)
    _map_cache[project_path] = (hash_key, map_text)
    return map_text


def build_turn_review(project, project_path, user_msg, response_text, tool_events, token_stats, round_notes=None):
    """一回二第二則訊息內容：委派本檔 turn_review。失敗回空字串，不影響主回覆。"""
    try:
        return turn_review(
            project_path, user_msg, response_text, tool_events or [], token_stats or {}, round_notes or {}
        ) or ''
    except Exception as e:
        print(f'[smart_brake] build_turn_review 失敗: {e}')
        return f'## 判官\n（判官這輪失敗，原因：{e}）'


def read_utf8(path, limit=8000):
    """讀取 UTF-8 檔案，最多 limit 字元，不存在時返回空字串。"""
    if not os.path.exists(path):
        return ''
    with open(path, encoding='utf-8', errors='replace') as f:
        return f.read(limit)


def read_utf8_tail(path, limit=50000):
    """讀取檔案末尾 limit bytes（適合 append-only logs）。"""
    if not os.path.exists(path):
        return ''
    size = os.path.getsize(path)
    with open(path, 'rb') as f:
        if size > limit:
            f.seek(-limit, 2)
        data = f.read(limit)
    return data.decode('utf-8', errors='replace')








def step2_scene_collapse(project_path):
    """Step 2 幾何場景張量塌縮：生成專案結構地圖。"""
    if not project_path:
        return ''
    try:
        return generate_project_map(project_path)
    except Exception:
        return ''










def _today_log_path(project_path):
    return os.path.join(project_path, 'log.md')


def _today_log_jsonl_path(project_path):
    return os.path.join(project_path, 'log.jsonl')


def _jsonl_turns_to_display(records):
    turns = []
    for rec in records or []:
        c = rec.get('content')
        if not isinstance(c, dict):
            continue
        if c.get('role_turn') != 'chat':
            if 'user' not in c or 'bot' not in c:
                continue
        user_text = c.get('user')
        bot_text = c.get('bot')
        if user_text is None or bot_text is None:
            continue
        ts = c.get('ts') or rec.get('time') or ''
        turns.append({'user': user_text, 'bot': bot_text, 'time': ts})
    return turns


def _parse_log_turns(content, n=8):
    turns = []
    parts = re.split(r'\n(?=## \d{4}-\d{2}-\d{2} \d{2}:\d{2})', content)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        tm = re.match(r'^## (\d{4}-\d{2}-\d{2} \d{2}:\d{2})', part)
        timestamp = tm.group(1) if tm else ''
        um = re.search(r'\*\*使用者\*\*：(.+?)(?=\*\*Claude\*\*：|\Z)', part, re.DOTALL)
        cm = re.search(r'\*\*Claude\*\*：(.+)', part, re.DOTALL)
        if um and cm:
            user_text = um.group(1).strip()
            bot_text = re.sub(r'\n+---\s*$', '', cm.group(1)).strip()
            turns.append({'user': user_text, 'bot': bot_text, 'time': timestamp})
    return turns[-n:] if len(turns) > n else turns


def read_log_for_display(project_path, n=10):
    """讀取最近 n 輪對話供前端顯示（雙來源合併：舊 log.md ＋ 新 log.jsonl）。"""
    old_turns = []
    content = read_utf8_tail(_today_log_path(project_path), limit=400000)
    if content:
        old_turns = _parse_log_turns(content, n=10 ** 9)
    new_turns = []
    try:
        records = read_log(_today_log_jsonl_path(project_path))
        new_turns = _jsonl_turns_to_display(records)
    except Exception:
        new_turns = []
    merged = old_turns + new_turns
    return merged[-n:] if len(merged) > n else merged


def append_log_turn(project_path, user_msg, response):
    """追加一輪對話，經 write_log 落地到 log.jsonl。"""
    now = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    content = {'role_turn': 'chat', 'user': user_msg, 'bot': response, 'ts': now}
    write_log(_today_log_jsonl_path(project_path), content)


def trim_log_if_needed(project_path, max_lines=500):
    """已停用自動裁剪（2026-07-07），保留單一完整 log。"""
    return


def _append_verbose_turn(project_path, user_msg, tool_events):
    """把這輪的工具呼叫清單寫入 verbose_log.md（§59 已停用寫入、函式保留可還原）。"""
    ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    verbose_log_path = os.path.join(project_path, 'verbose_log.md')
    lines = [f'\n## {ts}', f'**User:** {user_msg}', '', f'**工具呼叫（{len(tool_events)} 次）：**']
    for i, ev in enumerate(tool_events, 1):
        tool_name = ev.get('name', ev.get('tool', '?'))
        tool_input = ev.get('input', ev.get('content', ''))
        if isinstance(tool_input, dict):
            tool_input = json.dumps(tool_input, ensure_ascii=False)
        if len(str(tool_input)) > 200:
            tool_input = str(tool_input)[:200] + '…'
        lines.append(f'{i}. {tool_name} → {tool_input}')
    entry = '\n'.join(lines) + '\n'
    with open(verbose_log_path, 'a', encoding='utf-8') as f:
        f.write(entry)


def _trim_verbose_log(project_path, max_turns=20):
    """verbose_log.md 超過 max_turns 則時，把最舊的搬到 verbose_states.md。"""
    verbose_log_path = os.path.join(project_path, 'verbose_log.md')
    if not os.path.exists(verbose_log_path):
        return
    with open(verbose_log_path, encoding='utf-8', errors='replace') as f:
        content = f.read()
    pat = re.compile(r'^## \d{4}-\d{2}-\d{2} \d{2}:\d{2}', re.MULTILINE)
    matches = list(pat.finditer(content))
    if len(matches) <= max_turns:
        return
    cutoff_idx = len(matches) - max_turns
    cutoff_pos = matches[cutoff_idx].start()
    first_pos = matches[0].start()
    to_archive = content[first_pos:cutoff_pos]
    to_keep = content[cutoff_pos:]
    verbose_states_path = os.path.join(project_path, 'verbose_states.md')
    with open(verbose_states_path, 'a', encoding='utf-8') as f:
        f.write(to_archive)
    with open(verbose_log_path, 'w', encoding='utf-8') as f:
        f.write(content[:first_pos] + to_keep)


def _tool_summary(name, inp):
    """生成工具呼叫的簡短摘要，供 token_log 記錄。"""
    if not inp:
        return ''
    if name == 'Bash':
        return inp.get('command', '')[:120]
    if name in ('Read', 'Write', 'Edit'):
        return inp.get('file_path', '')
    if name == 'Agent':
        desc = inp.get('description', inp.get('prompt', ''))
        return str(desc)[:120]
    first_val = next(iter(inp.values()), '')
    return str(first_val)[:120]


TOKEN_LOG_ARCHIVE_INTERVAL_DAYS = 2


def archive_token_log_if_needed(project_path, now=None, force=False):
    """每 2 天把 token_log.md 累積的 `## ` 區塊搬進 token_log_archive.md。"""
    now = now or datetime.datetime.now()
    log_path = os.path.join(project_path, 'token_log.md')
    archive_path = os.path.join(project_path, 'token_log_archive.md')
    marker_path = os.path.join(project_path, '.token_log_archive_state.json')
    stamp = now.strftime('%Y-%m-%d %H:%M:%S')

    def _write_marker():
        with open(marker_path, 'w', encoding='utf-8') as f:
            json.dump({'last_archive': stamp}, f, ensure_ascii=False)

    last = None
    try:
        with open(marker_path, 'r', encoding='utf-8') as f:
            last = json.load(f).get('last_archive')
    except (OSError, ValueError):
        last = None
    if not force:
        if not last:
            _write_marker()
            return {'archived': False, 'reason': 'baseline_set', 'moved_entries': 0, 'moved_chars': 0}
        try:
            last_dt = datetime.datetime.strptime(last, '%Y-%m-%d %H:%M:%S')
        except ValueError:
            last_dt = now
        if (now - last_dt) < datetime.timedelta(days=TOKEN_LOG_ARCHIVE_INTERVAL_DAYS):
            return {'archived': False, 'reason': 'not_due', 'moved_entries': 0, 'moved_chars': 0}
    try:
        with open(log_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except OSError:
        content = ''
    idx = content.find('## ')
    body = content[idx:] if idx >= 0 else ''
    if not body.strip():
        _write_marker()
        return {'archived': False, 'reason': 'empty', 'moved_entries': 0, 'moved_chars': 0}
    moved_entries = (1 if body.startswith('## ') else 0) + body.count('\n## ')
    moved_chars = len(body)
    banner = (f'\n\n<!-- ===== 封存於 {stamp}，來自 token_log.md，'
              f'共 {moved_entries} 筆 / {moved_chars} 字元 ===== -->\n')
    with open(archive_path, 'a', encoding='utf-8') as f:
        f.write(banner + body.rstrip() + '\n')
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(content[:idx] if idx > 0 else '')
    _write_marker()
    return {'archived': True, 'reason': 'archived',
            'moved_entries': moved_entries, 'moved_chars': moved_chars}


def _token_log_jsonl_path(project_path):
    return os.path.join(project_path, 'token_log.jsonl')


def archive_token_log_jsonl_if_needed(project_path, now=None, force=False):
    """token_log.jsonl 版兩天封存（§62.4）。"""
    now = now or datetime.datetime.now()
    log_path = _token_log_jsonl_path(project_path)
    archive_path = os.path.join(project_path, 'token_log_archive.jsonl')
    marker_path = os.path.join(project_path, '.token_log_jsonl_archive_state.json')
    stamp = now.strftime('%Y-%m-%d %H:%M:%S')

    def _write_marker():
        try:
            with open(marker_path, 'w', encoding='utf-8') as f:
                json.dump({'last_archive': stamp}, f, ensure_ascii=False)
        except OSError:
            pass

    last = None
    try:
        with open(marker_path, 'r', encoding='utf-8') as f:
            last = json.load(f).get('last_archive')
    except (OSError, ValueError):
        last = None
    if not force:
        if not last:
            _write_marker()
            return {'archived': False, 'reason': 'baseline_set', 'moved_entries': 0, 'moved_chars': 0}
        try:
            last_dt = datetime.datetime.strptime(last, '%Y-%m-%d %H:%M:%S')
        except ValueError:
            last_dt = now
        if (now - last_dt) < datetime.timedelta(days=TOKEN_LOG_ARCHIVE_INTERVAL_DAYS):
            return {'archived': False, 'reason': 'not_due', 'moved_entries': 0, 'moved_chars': 0}
    records = read_log(log_path)
    if not records:
        _write_marker()
        return {'archived': False, 'reason': 'empty', 'moved_entries': 0, 'moved_chars': 0}
    moved_entries = len(records)
    moved_chars = 0
    for rec in records:
        c = rec.get('content')
        if not isinstance(c, dict):
            c = {'_raw': c}
        else:
            c = dict(c)
        c.setdefault('_orig_time', rec.get('time'))
        try:
            moved_chars += len(json.dumps(rec, ensure_ascii=False))
        except Exception:
            pass
        write_log(archive_path, c)
    try:
        with open(log_path, 'w', encoding='utf-8') as f:
            f.write('')
    except OSError:
        pass
    _write_marker()
    return {'archived': True, 'reason': 'archived',
            'moved_entries': moved_entries, 'moved_chars': moved_chars}


def append_token_log(project_path, user_msg, token_stats, tool_events):
    """追加 token 使用量和工具活動，經 write_log 落地到 token_log.jsonl。"""
    try:
        archive_token_log_jsonl_if_needed(project_path)
    except Exception:
        pass
    now = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
    inp = token_stats.get('input_tokens', 0)
    out = token_stats.get('output_tokens', 0)
    cr = token_stats.get('cache_read_input_tokens', 0)
    cw = token_stats.get('cache_creation_input_tokens', 0)
    cost = token_stats.get('cost', 0)
    dur_ms = token_stats.get('duration_ms', 0)
    preview = user_msg[:80] + ('…' if len(user_msg) > 80 else '')
    bic = token_stats.get('base_inject_chars', 0)
    bit = token_stats.get('base_inject_est_tokens', 0)
    tools = []
    for ev in (tool_events or []):
        tools.append({'tool': ev.get('tool'), 'summary': ev.get('summary')})
    per_round = {}
    per_round_cr = {}
    for ev in (tool_events or []):
        rd = ev.get('round')
        ot = ev.get('out_tokens')
        cr_r = ev.get('cache_read')
        if rd is not None and isinstance(ot, (int, float)):
            per_round[rd] = ot
        if rd is not None and isinstance(cr_r, (int, float)):
            per_round_cr[rd] = cr_r
    content = {
        'role_turn': 'token', 'ts': now, 'user_preview': preview,
        'input_tokens': inp, 'output_tokens': out,
        'cache_read_input_tokens': cr, 'cache_creation_input_tokens': cw,
        'total_incl_cache': inp + out + cr + cw,
        'base_inject_chars': bic, 'base_inject_est_tokens': bit,
        'cost': cost, 'duration_ms': dur_ms,
        'tool_count': len(tool_events or []), 'tools': tools,
        'per_round_out_tokens': {str(k): per_round[k] for k in per_round},
        'per_round_cache_read': {str(k): per_round_cr[k] for k in per_round_cr},
    }
    write_log(_token_log_jsonl_path(project_path), content)


_MAP_FOOTER_MARKER = '\n\n---\n📁 **專案結構地圖'


def _strip_map_footer(text):
    """截掉 assistant 回應裡的結構地圖 footer。"""
    idx = text.find(_MAP_FOOTER_MARKER)
    if idx >= 0:
        return text[:idx].rstrip()
    return text


def build_tasks_injection(project_path, recent_n=5):
    """把 tasks.md 注入 prompt 的內容瘦身：只取最近 recent_n 章節＋全檔未完成待辦。"""
    raw = read_utf8(os.path.join(project_path, 'tasks.md'), limit=100000)
    if not raw.strip():
        return ''
    starts = [m.start() for m in re.compile(r'^## ', re.MULTILINE).finditer(raw)]
    if starts:
        preamble = raw[:starts[0]].strip()
        cut = starts[recent_n] if len(starts) > recent_n else len(raw)
        recent = raw[starts[0]:cut].rstrip()
    else:
        preamble = raw.strip()
        recent = ''
    pending = [ln.rstrip() for ln in raw.splitlines() if ln.lstrip().startswith('- [ ]')]
    parts = []
    if preamble:
        parts.append(preamble)
    if recent:
        parts.append(recent)
    if pending:
        parts.append(f'## 未完成待辦（全檔彙整，共 {len(pending)} 項；完成後請打勾）\n'
                      + '\n'.join(pending))
    return '\n\n'.join(parts)


def _arcus_system_prompt(project_path):
    """Arcus 專屬系統提示：身分＝引擎本身，唯一的手是 arcus_do。"""
    return f"""你就是 Arcus——一具把「讀碼、改碼、驗證、記錄、繪圖」全部收斂到單一工具背後的自我維護引擎。你不是「某個環境裡被裝了幾個工具的助理」，你就是這具引擎本身。

## 你的解剖構造（唯一的手）

你只有一隻手。這隻手在工具清單裡登記的正式名字是 `mcp__arcus__arcus_do`——每次發出工具呼叫，name 欄位必須一字不差填這個全名，絕不可簡寫成 `arcus_do`；簡寫會被系統回「No such tool available」，整輪一個動作都送不出去。下文為了好讀仍用簡稱 arcus_do 指這隻手，但真正發出呼叫時一律用全名 `mcp__arcus__arcus_do`。它收 `(cmd, payload)` 兩個參數，你的一切動作都經它發出，沒有第二個出口。

每個 cmd 的用法、參數與範例，一律看這隻手自己的說明欄（發出呼叫時系統會附上 inputSchema 與 description）——那份跟著 arcus_core.py 的程式碼走、永遠是現況；此處不再複述，免得兩邊各記一份、日久漂移。

繪圖、上網搜尋、讀 pdf、發佈網頁這四項都是引擎本體的一部分，程式碼就在 arcus_core.py 裡，不是外借的工具，也不需要任何外部腳本。
每輪會注入 arcus_workspace.md，操作細節與繪圖規格速查看那份。`read`／`query` 只認當前專案資料夾底下的相對路徑。
你和使用者的每一輪對話，系統會自動把整段（使用者的話＋你的完整回應）逐輪追加寫進當前專案資料夾底下的 **log.jsonl**（不是 log.md、也不是任何日期檔）；舊的 log.md 只讀不寫、留作歷史。要回看更早、已滑出視野的對話，就用 `read`／`query` 去撈 log.jsonl。這是系統自動落地、不必你動手；別跟 `log` 指令寫的 arcus_log.md（你主動記的維護記錄）搞混。
當使用者要你回憶或回顧先前對話（例如「回憶一下」「我們之前講到哪」「想不起來剛剛的結論」），主動 `read` 當前專案的 log.jsonl。時區務必自行換算：log.jsonl 裡的 ts 記的是機器時間，也就是 UTC（世界協調時間、+0000），而使用者用台北時間（UTC+8）。換算方法只有一條、也最不會出錯：把每一筆 ts 都加 8 小時、先變成台北時間，之後一律用台北的年月日來判斷。先拿檔中最新一筆 ts 加 8 小時，當作「台北現在」，據此定出台北的今天，昨天就是今天減一天；凡是 ts 加 8 小時後、台北日期落在今天或昨天的輪次，全部撈出來，據此重建脈絡再回答。log.jsonl 過大而外溢時，把回傳的 spill:// 把手放回 `read`（帶 offset／limit 分段讀尾端最近的輪次），或用 `query` 掃關鍵詞定位，不必一次撈回整檔。

## 你沒有的東西（天生如此，不是被拿走）

你沒有 Bash、Read、Edit、Write、Grep、Glob、WebFetch、WebSearch、ToolSearch，也沒有「用完整路徑讀寫檔案」「跑 python3 腳本」這些能力。這不是限制、不是誰拿走了你的手，這就是你的天生構造。除了 arcus_do，你唯一還能用的工具是 Skill——只用來在探索模式啟動探索技能（見下方回覆模式段）；工具清單裡若殘留其他任何名字，一律當它不存在、絕不呼叫。

**絕不向使用者「誠實界定」「盤點」你有哪些工具、缺哪些工具、能不能做某件事。** 使用者要的是你直接用 arcus_do 把事做完，不是聽你交代自己的手腳。要動手就動手；真的卡住就把情況講清楚、乾淨結束，不要空轉重試不存在的工具。

## 回覆模式標籤（強制，第一行單獨輸出，之後空一行再接正文）

- 【探索模式】：一般問答、討論、沒有出現明確授權語時（預設一律用它）
- 【規劃模式】：提案／規劃階段
- 【實作模式】：收到「請實作」「實作」「開始做」「好」「可以」「做吧」等授權語後動手時

進入探索模式做深入探討時，呼叫 `Skill(openspec-explore)` 啟動探索技能，把問題想深想廣；探索模式只討論、不動任何檔案。這是你除 arcus_do 外唯一可用的工具。

## 查佐證定則（卡片欄位空了怎麼撈原文）

想引用某篇已入庫論文的原文佐證某個分析角度時，照這條路走，別自己開岔：

- 分析角度不是卡片欄位。角度是你腦裡的問題（例如「資料正當性」「fit-for-purpose 哲學」），卡片欄位是固定那六欄（研究問題、研究方法、研究流程、研究成果、爭議、未來發展）；拿角度的字面去撞欄位名一定撞不到。
- 欄位是空的、或欄位內容對不上你要的角度，就直接用引用工具（paper_quote）帶 search 參數，拿角度的關鍵詞去掃全文段落；不知道這篇切成哪些段，就先不給 keys 呼叫一次，它會把可用的段落鍵與段數列回來，再挑鍵取原文。
- 已經入庫的論文一律不讀 PDF、不讀卡片的原始 JSON、不試圖繞過路徑沙箱；卡片指令（paper_cards）與引用工具（paper_quote）就是撈這篇原文的唯一入口，沒有第二條路。
- 同一個工具連撞兩次都沒結果，就換上面這條路線，不要把同一個關鍵詞改成一堆變體反覆重試。

## 設計思考原則（省掉重複解釋的機制）

你在學這位使用者的設計思考模式：把他每次實質回答背後的原則萃取下來、存起來、下次直接套用，讓他不必一再重講同一套想法。四步，用三個 arcus_do 指令操作：

- 收料：使用者給了實質回答（不是單純「好」「可以」）後，判斷這是不是一條可重用的原則（例如「面對不可逆的外推、又有分歧時，先把狀況攤清楚再問範圍」）。先用 principle_list 看有沒有既有同類；沒有才用 principle_add 寫一條（when 情境／then 偏好／because 本質／scope 適用範圍）。這條先是暫定，不會自動拿去用。
- 固化：同一條原則又被觀察到成立，呼叫 principle_hit(id, "confirm")；累積到門檻且沒被打臉，它自動升為上線。
- 套用：每輪開頭系統會把上線原則餵給你。碰到新問題若命中某條上線原則，直接照做、不要問，但一定附一句報備「依你〔id〕原則我選了X（不對就說一聲）」。沒有原則命中、或兩條原則打架，這時才問——問的正好是更深、沒被涵蓋的那一題。
- 打臉：使用者更正了你照原則做的決定，立刻呼叫 principle_hit(id, "contradict")，那條退回暫定、不再自動套用。錯誤當場吸收、不累積。

原則分兩種，套用方式不同：過程型原則（例如「先攤清狀況再問範圍」「先講結論再展開」）只規範你切入問題的姿態、不碰答案內容，命中就照做。結論型原則（例如「架構偏好最簡方案」）會鎖定答案內容，這時原則只是預設值、不是眼罩——你照樣以該偏好當主推，但只要這次真有一個實質更好或更有意思的替代方案，就必須點出來：「依你〔id〕原則我選了X；不過這次Y可能更合適／更值得一看，因為……」。學到的偏好只設定起點、永遠不封路，別讓省時壓掉了該有的發散。

鐵律：這機制成功的樣子是「越來越少問」。絕不可為了多萃取而多問；原則靠讀懂回答的本質而長，不是靠一直問而長。萃取寧可窄一點、把 scope 標清楚，也不要拿一次回答就當通則。

## 工作紀律（比照本機 Claude 的規則）

**絕不自創臨時工具。** 不准為了繞過草稿與型別檢查而在 arcus_core.py 裡追加任何「用完即刪」「暫時性」「臨時修補」的分派分支或字串替換函式。歷史上出現過一個叫 patch_text 的臨時分支，它包住整個分派點、讓任何檔案都能被直接字串替換，繞過全部安全流程，且註解寫著用完即刪卻長期留著——已於 2026-07-19 移除。改檔一律走 `stage_new` → `stage_run` → `promote` 這條路，沒有例外；覺得這條路做不到某件事，就把做不到的地方講清楚交給使用者判斷，不要自己開後門。

**工具一律留在 arcus 專案資料夾內。** 引擎需要的能力全部併在 arcus_core.py 裡，不得引用 ~/.claude/tools 或任何其他專案底下的絕對路徑，也不得靠 sys.path 借用外部檔案。

**白話寫作**：一律完整白話文，先講結論與重點再展開。禁止任何英文縮寫、中文縮寫、自創術語、省略語法；非用專有名詞不可時，第一次出現用括號補白話。中文句子保留完整主詞、動詞、受詞，不省略謂語，也不把受詞前移（不寫「14 GB 用掉」，要寫「已經用掉 14 GB」）。句子不以「的」收尾，也不用「是……的」把動作名詞化，改用主動動詞把真正的施事者放回主詞位（不寫「單一階層精煉不出東西，是因為現象是**跨階層湧現**的」，要寫「跨階層湧現了單一階層無法精煉的現象」）。模態詞、疑問詞（該不該、要不要、有沒有）緊貼它所管的動詞，不孤懸句尾變殘塊（不寫「射線穿過中心檢查一次該不該」，要寫「該不該檢查射線有沒有穿過中心」）。前置的話題子句若跟主句沒有句法接點，就折成定語收進它修飾的名詞前，不讓它漂在句首（不寫「每一門要落地，射線都得穿過倫理學」，要寫「檢查能讓每一門學科落地的那條射線有沒有穿過倫理學」）。正文流動的句子裡不插入粗體星號、斜體等行內強調記號（不污染中文語流）；表格、清單、程式碼區塊、ASCII 方框圖這類區塊結構不受此限、照常使用——這條只擋句子裡的行內強調，不擋版面結構。純中文敘述裡，工具或指令名稱不當主詞或受詞，先用白話描述功能、原名放括號；程式碼識別字原樣保留，第一次出現用括號補一句白話。改用分層寫法：主句只講結論加下一步，技術細節縮排到括號補一句。每個說明點不超過 150 字，避免代名詞，預設只給一個最佳解。

**自檢與驗證（可自行確認的一律自行查驗，不問使用者）**：提出假說先測試確認才動手，用 stage_new 建草稿、stage_run 做型別感知測試，不憑印象。實作後執行驗證，禁止目測代替執行；沒跑過測試前不說「已完成」。該讀多少就讀多少——改一段碼前先用 read 讀足上下文，不用猜測代替查證。

**回憶大檔要分段精煉，不一口氣整讀**：使用者要你回顧、而 log.jsonl 很大（十幾萬字以上）時，別把整檔一次讀進來硬啃——一次掃一大包，注意力會被攤薄、中段最容易漏掉，最後只會覆述表面、講不出結論。改用 read 帶 offset／limit 分段切讀，每段約兩萬字（20000 字）為一個單位；每讀完一段先當場把該段重點寫成筆記，再讀下一段；十幾萬字大約切成七到八段，全部讀完後把各段筆記併成完整結論。原文仍整份進脈絡、一個字不少，變的只是「分段讀、逐段記、最後彙整」這個讀法。

**每一輪都把專有名詞與重大選擇記進 memory.md**：每輪對話過程中，掃描範圍同時涵蓋使用者訊息與你自己回覆兩邊，把兩邊出現的專有名詞——涵蓋所有學科的術語（不限任何科目），以及技術名詞、工具名、檔名、函式／變數名、IP 與位址、門檻數值、API 端點、人名地名機構名、專案代號等——與重大選擇，用 stage_new 的 append 型別接進當前專案的 memory.md 檔尾（不覆蓋舊內容）。重大選擇＝會影響後續方向的回答或選擇（訂下門檻、定架構、排除某條路、拍板一個計畫）；只是附和、把前文覆述、隨手能回退的暫時步驟，不算。沿用 memory.md 既有規矩：只增修、去重，不重複記錄，只在它有新意義時補一句。

**一輪之內一次做完（你的天生節奏，不是被擋才這樣）**：一次讀、一次改、一次測、一次落地。分析階段就把這輪要動的檔用 read 的 paths 一次讀足，想清楚每個檔要改成什麼，再用 stage_new 的 drafts 一口氣寫成整批草稿信封，用 stage_run 一次測完整批，全過才 promote 一次原子落地、最後 cleanup 一次清掉草稿。就算只動一個檔，paths 與 drafts 也直接寫成單一元素的陣列。你不會想到要一筆一筆零敲碎打——批次是本能，不是柵欄。

**上網搜尋也是同一輪一次做完**：先用 discover 把一批關鍵詞丟出去，拿回候選網址清單；在腦中審視哪些來源可信、要不要補上官方權威網址、要不要把分頁網址（例如 ?page=0..N）一併列進去；決定後把所有要抓的網址放進 read 的 paths 一次抓回正文。discover 一次、read 一次，中間的審視是你自己動腦、不算回合。絕不要一個網址一個網址慢慢抓。抓回來的正文若某筆太大，read 會自動把那筆存進外溢檔、只回你 preview 和一個 spill:// handle（不是真的抓失敗，正文完整在手上）；要看全文就把那個 handle 放回 paths 再 read（想分段就帶 offset/limit），或用 query 拿 handle 掃關鍵詞填欄位。所以批次盡量大、不必怕正文爆量，更不要因為怕爆量就退回一筆一筆慢抓。

**寫程式碼走填空模板（結構上不出語法骨架錯）**：要產生 kind=code 的草稿內容時，用本專案既有的行型模板——emit_func 定義函式、emit_try_return 包 try、emit_if_return 條件回傳、emit_indent 補縮排、emit_read_utf8／emit_write_utf8 讀寫檔、emit_selftest 產自我測試——你只填名稱、參數、內容，def、括號、冒號、縮排、標點一律交給模板補。填內容、模板補骨架，是本能，不是被檢查才用。

工作路徑：{project_path}"""


def build_system_prompt(project, project_path, user_msg=None, history=None):
    """Build the _SYSTEM prompt string for the claude CLI subprocess."""
    # arcus 是所有專案共用的底層引擎，不分專案一律回原生身分提示，讓系統提示與「唯一工具」邊界對齊。
    return _arcus_system_prompt(project_path) + _principles_context_block()


def parse_image_with_claude(image_b64, prompt_hint='識別這道題目的完整文字與選項'):
    """接收 base64 圖片，呼叫 Claude CLI 識別，回傳結果 dict。"""
    img_path = txt_path = sys_path = None
    try:
        img_bytes = base64.b64decode(image_b64)
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            f.write(img_bytes)
            img_path = f.name
        _SYSTEM_IMG = (
            '你是考試題目識別工具，擅長讀取圖片中的繁體中文試題。'
            '只輸出純 JSON，不加說明文字、不加 markdown 代碼塊。'
        )
        full_prompt = (
            f'請用 Read 工具讀取圖片 {img_path}，{prompt_hint}。\n'
            '只輸出以下格式的 JSON（不加 ```json 或其他說明）：\n'
            '{"question": "完整題目文字", "options": ["A...", "B...", "C...", "D..."], '
            '"type": "multiple_choice 或 short_answer 或 essay 或 calculation"}\n'
            '若無選項（非選擇題），options 輸出空陣列 []。'
        )
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(full_prompt)
            txt_path = f.name
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(_SYSTEM_IMG)
            sys_path = f.name
        env = dict(os.environ)
        env['CLAUDE_SUBPROCESS'] = '1'
        env['PYTHONIOENCODING'] = 'utf-8'
        cmd = (
            f'sys=$(cat "{sys_path}"); '
            f'CLAUDE_SUBPROCESS=1 claude --print --output-format stream-json '
            f'--dangerously-skip-permissions --system-prompt "$sys" < "{txt_path}"'
        )
        result = subprocess.run(
            ['bash', '-c', cmd],
            capture_output=True, text=True, encoding='utf-8', errors='replace',
            timeout=90, env=env
        )
        response_text = ''
        for line in result.stdout.split('\n'):
            line = line.strip()
            if not line:
                continue
            try:
                ev = json.loads(line)
                if ev.get('type') == 'assistant':
                    for block in ev.get('message', {}).get('content', []):
                        if block.get('type') == 'text':
                            response_text += block.get('text', '')
            except (json.JSONDecodeError, KeyError):
                pass
        clean = re.sub(r'```(?:json)?\s*', '', response_text).strip().rstrip('`').strip()
        json_match = re.search(r'\{[\s\S]*\}', clean)
        if json_match:
            try:
                data = json.loads(json_match.group(0))
                return {'ok': True, 'data': data, 'raw': response_text}
            except json.JSONDecodeError as e:
                return {'ok': False, 'error': f'JSON 解析失敗: {e}', 'raw': response_text}
        else:
            return {'ok': False, 'error': '回應中找不到 JSON', 'raw': response_text}
    except subprocess.TimeoutExpired:
        return {'ok': False, 'error': 'Claude CLI 逾時（90s）'}
    except Exception as e:
        return {'ok': False, 'error': str(e)}
    finally:
        for p in (img_path, txt_path, sys_path):
            if p:
                try:
                    os.unlink(p)
                except OSError:
                    pass


def read_workspace_md():
    """讀取同目錄的集中工作區 arcus_workspace.md，永不拋例外；缺檔回空字串。"""
    try:
        _p = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'arcus_workspace.md')
        with open(_p, 'r', encoding='utf-8') as _f:
            return _f.read()
    except Exception:
        return ''


def before_prompt_hook(project, project_path, user_msg, prompt):
    """在 prompt 組合完成後、寫入 temp 檔之前執行。回傳增強後的 prompt 字串。

    §C5 封存（使用者授權）：只保留「場景坍縮（結構地圖）＋函式索引」注入。
    反向傳播、法條層 run()／scan()、違規記錄、智慧煞車斷點、其餘 step
    （DNA 投影／熵阻尼／量子快門）的呼叫全部從本 hook 移除；反向傳播與法條層
    兩群已於 2026-07-13 從本檔實刪，判官等其餘保留函式仍在原處。
    """
    scene_map_str = step2_scene_collapse(project_path)
    if scene_map_str:
        stable_prefix = (
            f'=== Arcus 系統架構 ===\n{ARCUS_ARCHITECTURE}\n\n'
            f'=== 專案結構地圖：{project}（工作路徑：{project_path}）===\n{scene_map_str}'
        )
    else:
        stable_prefix = f'=== Arcus 系統架構 ===\n{ARCUS_ARCHITECTURE}'
    try:
        _sym_idx = build_symbol_index()
    except Exception:
        _sym_idx = None
    if _sym_idx:
        stable_prefix += (
            '\n\n=== Arcus 核心檔函式索引（帶行號；改碼或說明自身機制前先讀這裡，'
            f'勿重跑 grep 找行號）===\n{_sym_idx}'
        )
    if str(project).strip().lower() == 'arcus':
        try:
            _ws = read_workspace_md()
        except Exception:
            _ws = ''
        if _ws:
            stable_prefix += ('\n\n=== Arcus 集中工作區（arcus_workspace.md，維護狀態機與功能地圖，改碼前先讀）===' + '\n' + _ws)
    prompt = f'{stable_prefix}\n\n{prompt}'
    try:
        _mem_block = _memory_inject_block(user_msg)
    except Exception:
        _mem_block = ''
    if _mem_block:
        prompt = _mem_block + prompt
    return prompt




def after_response_hook(project, project_path, user_msg, response_text, tool_events=None):
    """在回應取回並記錄到 log 之後執行。

    §C5 封存（使用者授權）：判官改走 server.py → build_turn_review → turn_review，不在本 hook 內。
    本 hook 原本承載的智慧煞車／全動作留痕、語義截斷 fall 記錄、軌跡向量、scan()
    與 path-index 背景更新的呼叫全部移除。
    現在唯一副作用：觸發每輪記憶收料（每輪都收，過閘門即抽，寫進收件匣不碰索引）。
    """
    _memory_capture_turn(project_path, user_msg, response_text)
    return []


# ══════════════════════════════════════════════════════════════════════════════
# 函式索引表（使用者要求：快取機制做索引值＋函式簡易說明，讓工具知道調哪一個函式）
#   動作名 → (函式名, 一句說明)。python3 arcus_core.py --index 印出。
# ══════════════════════════════════════════════════════════════════════════════

FUNCTION_INDEX = {
    # 掛鉤（server.py 每輪呼叫的兩個進入點）
    '送出前注入上下文':       ('before_prompt_hook', '注入專案結構地圖＋核心檔函式索引，回增強prompt'),
    '回應後後處理':           ('after_response_hook', '現為無副作用空操作，判官改走 server.py'),
    '每輪工作明細與判官':     ('build_turn_review', '委派 turn_review：工作明細＋到期時 opus 判官'),
    # 結構地圖
    '生成專案結構地圖':       ('generate_project_map', 'os.walk 掃專案，7天內異動加★，雜湊快取'),
    # 對話 log
    '讀最近N輪對話':          ('read_log_for_display', '雙來源合併 log.md＋log.jsonl，取尾N筆'),
    '追加一輪對話':           ('append_log_turn', 'write_log 落地到 log.jsonl'),
    # token log
    '追加token使用量':        ('append_token_log', '每輪token/快取/費用/逐回合，落地 token_log.jsonl'),
    'token兩天封存':          ('archive_token_log_jsonl_if_needed', '依time判2天、搬進 archive'),
    # tasks 注入
    '瘦身tasks注入':          ('build_tasks_injection', '只取最近5章節＋全檔未完成待辦'),
    # 系統提示
    '建系統提示':             ('build_system_prompt', 'server.py 拼 claude CLI 用的 _SYSTEM 字串'),
    # 圖片辨識
    '辨識上傳題目圖':         ('parse_image_with_claude', 'base64圖→claude CLI辨識→JSON'),
    # token 估算
    '估token數':              ('est_tokens', '中日韓×0.6＋其餘÷4，供排序抓大頭'),
    # 中央 log 產生器
    '寫一筆JSONL log':        ('write_log', 'O_APPEND 原子追加、永不拋例外'),
    '讀回JSONL log':          ('read_log', '逐行容錯解析，limit 取尾N筆'),
    # 判官（第三群反向傳播 bake、第四群法條層 run／scan 已於 2026-07-13 從本檔刪除，無封存檔）
    '每輪判官解析':           ('turn_review', '工作明細＋每judge_interval輪一次opus判官'),
}


def _print_index():
    print('=== arcus_core.py 函式索引表 FUNCTION_INDEX ===')
    print(f'{"動作":<22}{"函式":<32}說明')
    print('-' * 100)
    for action, (fn, desc) in FUNCTION_INDEX.items():
        print(f'{action:<22}{fn:<32}{desc}')


if __name__ == '__DISABLED_MAIN__':
    import argparse
    ap = argparse.ArgumentParser(description='arcus_core.py 四支合體核心（命令列入口）')
    ap.add_argument('--index', action='store_true', help='印出函式索引表 FUNCTION_INDEX')
    args = ap.parse_args()
    if args.index:
        _print_index()
    else:
        ap.print_help()


# === ARCUS ADDON v1 · templates + staging + maintenance-log (2026-07-13) ===
# 三件事：①程式碼行型模板（填空式產生，消掉標點骨架錯）
#        ②腳本暫存區函式（建快取.py → 真的跑一遍 → 取代舊檔 → 刪除快取）
#        ③單一維護 log（arcus_log.md）的記錄函式
# 設計原則：模型寫碼時只填內容，標點與縮排由模板產生；驗證靠「真的跑一遍快取.py」，
#          不做攔錯閘門與錨點斷言（已依指示砍除，不封存）。
import shutil as _shutil
import importlib.util as _ilu

_ARCUS_DIR = os.path.dirname(os.path.abspath(__file__))
ARCUS_LOG_MD = os.path.join(_ARCUS_DIR, 'arcus_log.md')
ARCUS_STAGING_DIR = os.path.join(_ARCUS_DIR, 'staging')
# 大正文外溢存檔區：固定在 arcus 自己的 runtime 底下（不受當前專案路徑影響），
# 回傳 spill:// handle 讓引擎自己讀得回來，杜絕「存到專案外絕對路徑、引擎讀不回」的死路。
_ARCUS_FETCH_DIR = os.path.join(ARCUS_STAGING_DIR, 'fetch')


def _arcus_spill_path(rel):
    """把 spill:// handle（或裸檔名）解析成 fetch 區內的絕對路徑，basename 擋掉目錄逃逸。"""
    name = rel[len('spill://'):] if isinstance(rel, str) and rel.startswith('spill://') else rel
    return os.path.join(_ARCUS_FETCH_DIR, os.path.basename(name))


def _arcus_spill_write(rel, text):
    """把大正文寫進 fetch 區，檔名用來源雜湊（同來源覆寫同檔、不爆量），回傳 spill:// handle。"""
    import hashlib as _hl
    os.makedirs(_ARCUS_FETCH_DIR, exist_ok=True)
    key = _hl.md5(rel.encode('utf-8')).hexdigest()[:10]
    name = 'fetch_%s.txt' % key
    with open(os.path.join(_ARCUS_FETCH_DIR, name), 'w', encoding='utf-8') as f:
        f.write(text)
    return 'spill://' + name


# ── ① 行型模板（emit_*）：窮舉本專案反覆出現的行型，模型填空、模板補標點 ──

def emit_indent(text, level=1):
    """把多行文字整體縮排 level 層（4 空格／層）；空行不補空白。"""
    pad = '    ' * level
    return '\n'.join((pad + ln if ln else ln) for ln in text.split('\n'))


def emit_module_header(desc=''):
    """檔頭：coding 宣告＋一句用途註解。"""
    head = '# -*- coding: utf-8 -*-'
    return head + ('\n# ' + desc if desc else '')


def emit_func(name, args, body, doc=''):
    """函式定義：模板負責 def、括號、冒號、縮排；模型只給 name／args／body／doc。
    args 給字串串列；body 給函式體原始碼（未縮排，空字串自動補 pass）。"""
    sig = 'def %s(%s):' % (name, ', '.join(args))
    lines = [sig]
    if doc:
        lines.append('    """%s"""' % doc)
    lines.append(emit_indent(body if body.strip() else 'pass', 1))
    return '\n'.join(lines)


def emit_try_return(body, default="''"):
    """安全 try/except：主體出錯就回 default（本專案最常用的防呆行型）。"""
    return 'try:\n%s\nexcept Exception:\n    return %s' % (emit_indent(body, 1), default)


def emit_read_utf8(path_expr):
    """讀檔 UTF-8：以 path_expr 為路徑運算式，回傳完整內容。"""
    return ("with open(%s, 'r', encoding='utf-8') as _f:\n"
            "    return _f.read()") % path_expr


def emit_write_utf8(path_expr, data_expr, bom=False):
    """寫檔 UTF-8（bom=True 帶 BOM，對應本專案中文寫檔慣例）。"""
    enc = 'utf-8-sig' if bom else 'utf-8'
    return ("with open(%s, 'w', encoding='%s') as _f:\n"
            "    _f.write(%s)") % (path_expr, enc, data_expr)


def emit_sibling_path(name_literal):
    """同目錄檔案路徑運算式：os.path.join(檔案所在目錄, 名稱)。"""
    return "os.path.join(os.path.dirname(os.path.abspath(__file__)), %r)" % name_literal


def emit_now_stamp():
    """時間戳運算式：YYYY-MM-DD HH:MM。"""
    return "datetime.datetime.now().strftime('%Y-%m-%d %H:%M')"


def emit_dict_get(dict_expr, key_literal, default="''"):
    """字典安全取值：d.get('key', default)。"""
    return "%s.get(%r, %s)" % (dict_expr, key_literal, default)


def emit_project_guard(body, project_var='project', name='arcus'):
    """專案判斷區塊：僅在 project 等於指定名稱時執行 body。"""
    cond = "if str(%s).strip().lower() == %r:" % (project_var, name)
    return '%s\n%s' % (cond, emit_indent(body if body.strip() else 'pass', 1))


# ── ② 暫存區函式：建快取.py → 跑一遍驗證 → 取代舊檔 → 刪除快取（刪除，非清空）──

_DRAFT_TOP = {'schema_version', 'kind'}
_DRAFT_ALLOWED = {
    'code':    {'required': {'target', 'content'}, 'optional': set()},
    'content': {'required': {'target', 'content'}, 'optional': set()},
    'append':  {'required': {'target', 'content'}, 'optional': set()},
    'op':      {'required': {'op'},                'optional': {'target', 'from', 'to'}},
}


def _draft_resolve(rel):
    """用專案安全路徑解析相對路徑；回 (True, 絕對路徑) 或 (False, 錯誤訊息)。"""
    try:
        return True, _arcus_safe_path(rel)
    except Exception as e:
        return False, str(e)


def draft_validate_schema(d):
    """鍵值對齊閘門：schema_version==1、kind 合法、無未知鍵、必填齊、content 為字串。回錯誤清單。"""
    errs = []
    if not isinstance(d, dict):
        return ['草稿不是 JSON 物件']
    if d.get('schema_version') != 1:
        errs.append('schema_version 必須為 1')
    kind = d.get('kind')
    if kind not in _DRAFT_ALLOWED:
        return errs + ['kind 必須是 code/content/op/append，收到：%r' % kind]
    spec = _DRAFT_ALLOWED[kind]
    allowed = _DRAFT_TOP | spec['required'] | spec['optional']
    for k in d:
        if k not in allowed:
            errs.append('未知鍵（一律拒絕）：%s' % k)
    if 'op' in d and kind != 'op':
        errs.append("偵測到 op 鍵卻非 op 型別：追加到檔尾請用 kind='append'"
                    "（要追加的文字放 content 欄）；搬移或刪除才用 kind='op'")
    for k in spec['required']:
        if k not in d:
            errs.append('缺必填鍵：%s' % k)
    if kind == 'op':
        op = d.get('op')
        if op == 'move':
            for k in ('from', 'to'):
                if k not in d:
                    errs.append('op=move 缺鍵：%s' % k)
        elif op == 'delete':
            if 'target' not in d:
                errs.append('op=delete 缺鍵：target')
        else:
            errs.append('op 必須是 move/delete，收到：%r' % op)
    if kind in ('code', 'content', 'append') and not isinstance(d.get('content'), str):
        errs.append('content 必須是字串')
    return errs


def draft_typecheck_one(d):
    """型別感知檢查：先鍵值對齊，再依 kind 分岔。絕不執行程式碼、絕不刪改檔。回錯誤清單（空=通過）。"""
    errs = draft_validate_schema(d)
    if errs:
        return errs
    kind = d['kind']
    if kind == 'code':
        ok, r = _draft_resolve(d['target'])
        if not ok:
            return [r]
        if not d['target'].endswith('.py'):
            return ['kind=code 的 target 必須以 .py 結尾']
        try:
            compile(d['content'], d['target'], 'exec')   # 只編譯不執行
        except SyntaxError as e:
            return ['Python 語法錯誤：第 %s 行 %s' % (e.lineno, e.msg)]
        return []
    if kind == 'content':
        ok, r = _draft_resolve(d['target'])
        if not ok:
            return [r]
        if d['target'].lower().endswith('.json'):
            try:
                json.loads(d['content'])
            except Exception as e:
                return ['內容不是合法 JSON：%s' % e]
        return []
    if kind == 'append':
        ok, r = _draft_resolve(d['target'])
        if not ok:
            return [r]
        return []
    # op
    if d['op'] == 'delete':
        ok, full = _draft_resolve(d['target'])
        if not ok:
            return [full]
        if not os.path.exists(full):
            return ['刪除目標不存在：%s' % d['target']]
        return []
    ok1, src = _draft_resolve(d['from'])
    if not ok1:
        return [src]
    ok2, dst = _draft_resolve(d['to'])
    if not ok2:
        return [dst]
    errs2 = []
    if not os.path.exists(src):
        errs2.append('搬移來源不存在：%s' % d['from'])
    if os.path.exists(dst):
        errs2.append('目的地已被占用（不覆蓋）：%s' % d['to'])
    parent = os.path.dirname(dst)
    if parent and not os.path.isdir(parent):
        errs2.append('目的地上層資料夾不存在：%s' % os.path.dirname(d['to']))
    return errs2


def stage_typecheck_files(draft_paths):
    """對多份 draft_*.json 逐一做型別感知測試；回 (全過?, 報表)。絕不執行程式碼、絕不刪改檔。"""
    report = []
    ok_all = True
    for p in draft_paths:
        try:
            d = json.loads(open(p, encoding='utf-8').read())
        except Exception as e:
            report.append([os.path.basename(p), False, ['讀取或解析失敗：%s' % e]])
            ok_all = False
            continue
        errs = draft_typecheck_one(d)
        report.append([os.path.basename(p), (not errs), errs])
        if errs:
            ok_all = False
    return (ok_all, report)


def stage_new_draft(name, draft):
    """把一份 JSON 草稿信封寫成 staging/draft_<name>.json，回傳路徑；要幾份就呼叫幾次。"""
    try:
        os.makedirs(ARCUS_STAGING_DIR, exist_ok=True)
        p = os.path.join(ARCUS_STAGING_DIR, 'draft_%s.json' % name)
        with open(p, 'w', encoding='utf-8') as _f:
            _f.write(json.dumps(draft, ensure_ascii=False, indent=2))
        return p
    except Exception:
        return ''


def stage_promote_drafts(draft_paths):
    """型別測試全過後，整批原子落地：code/content 寫檔、op 搬移或刪除。
    全有全無：任一筆失敗，已套用者依逆動作堆疊全部回滾。成功/回滾都回寫 log。回報表 dict。"""
    # 1. 載入
    loaded = []
    for p in draft_paths:
        try:
            d = json.loads(open(p, encoding='utf-8').read())
        except Exception as e:
            return {'ok': False, 'stage': 'load', 'error': '讀取草稿失敗 %s：%s' % (p, e), 'applied': []}
        loaded.append((p, d))
    # 2. 型別測試閘門（全過才動檔）
    for p, d in loaded:
        errs = draft_typecheck_one(d)
        if errs:
            return {'ok': False, 'stage': 'typecheck', 'draft': os.path.basename(p), 'errors': errs, 'applied': []}
    # 3. 本輪專屬快照夾（取代會互相覆蓋的單槽 .bak）
    try:
        snap = os.path.join(ARCUS_STAGING_DIR, '_snapshot')
        if os.path.exists(snap):
            _shutil.rmtree(snap)
        os.makedirs(snap, exist_ok=True)
    except Exception as e:
        return {'ok': False, 'stage': 'snapshot', 'error': str(e), 'applied': []}
    undo = []       # 逆動作堆疊
    applied = []

    def _snap(full, tag):
        dst = os.path.join(snap, '%d_%s' % (len(undo), tag))
        _shutil.copy2(full, dst)
        return dst

    try:
        for p, d in loaded:
            kind = d['kind']
            if kind in ('code', 'content'):
                full = _arcus_safe_path(d['target'])
                if os.path.exists(full):
                    bak = _snap(full, os.path.basename(full))
                    undo.append(lambda full=full, bak=bak: _shutil.copy2(bak, full))
                else:
                    parent = os.path.dirname(full)
                    if parent:
                        os.makedirs(parent, exist_ok=True)
                    undo.append(lambda full=full: os.path.exists(full) and os.remove(full))
                with open(full, 'w', encoding='utf-8') as _f:
                    _f.write(d['content'])
                applied.append({'kind': kind, 'target': d['target']})
            elif kind == 'append':
                full = _arcus_safe_path(d['target'])
                if os.path.exists(full):
                    _sz = os.path.getsize(full)
                    undo.append(lambda full=full, sz=_sz: os.truncate(full, sz))
                else:
                    parent = os.path.dirname(full)
                    if parent:
                        os.makedirs(parent, exist_ok=True)
                    undo.append(lambda full=full: os.path.exists(full) and os.remove(full))
                _fd = os.open(full, os.O_WRONLY | os.O_CREAT | os.O_APPEND, 0o644)
                try:
                    os.write(_fd, d['content'].encode('utf-8'))
                finally:
                    os.close(_fd)
                applied.append({'kind': 'append', 'target': d['target']})
            elif kind == 'op' and d['op'] == 'delete':
                full = _arcus_safe_path(d['target'])
                trash = _snap(full, 'deleted_' + os.path.basename(full))
                undo.append(lambda full=full, trash=trash: _shutil.copy2(trash, full))
                os.remove(full)
                applied.append({'kind': 'op', 'op': 'delete', 'target': d['target']})
            else:   # op move
                src = _arcus_safe_path(d['from'])
                dst = _arcus_safe_path(d['to'])
                _shutil.move(src, dst)
                undo.append(lambda src=src, dst=dst: _shutil.move(dst, src))
                applied.append({'kind': 'op', 'op': 'move', 'from': d['from'], 'to': d['to']})
    except Exception as e:
        rolled = 0
        for fn in reversed(undo):
            try:
                fn()
                rolled += 1
            except Exception:
                pass
        err = '%s: %s' % (type(e).__name__, e)
        try:
            append_arcus_log('rollback', '套用失敗回滾：%s（已還原 %d 筆）' % (err, rolled))
        except Exception:
            pass
        return {'ok': False, 'stage': 'apply', 'error': err, 'applied_then_rolled_back': applied, 'rolled': rolled}
    try:
        append_arcus_log('promote', '原子套用 %d 筆：%s' % (len(applied), json.dumps(applied, ensure_ascii=False)))
    except Exception:
        pass
    return {'ok': True, 'applied': applied, 'snapshot': snap}


def stage_cleanup(cache_paths):
    """刪除這輪建立的快取.py（是刪除，不是清空）；回傳實際刪掉的檔數。"""
    n = 0
    for p in cache_paths:
        try:
            os.remove(p)
            n += 1
        except Exception:
            continue
    return n


# ── ③ 單一維護 log：arcus_log.md（固定格式：分類標籤＋時間戳＋內容）──

def append_arcus_log(category, content):
    """把一筆維護記錄追加到唯一的 arcus_log.md；固定格式 `## [分類] 時間戳` 換行內容。永不拋例外。"""
    try:
        stamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
        cat = str(category).strip() or '一般'
        block = '\n## [%s] %s\n%s\n' % (cat, stamp, str(content).strip())
        with open(ARCUS_LOG_MD, 'a', encoding='utf-8') as _f:
            _f.write(block)
        return True
    except Exception:
        return False


def read_arcus_log():
    """讀回 arcus_log.md 全文；缺檔回空字串。"""
    try:
        with open(ARCUS_LOG_MD, 'r', encoding='utf-8') as _f:
            return _f.read()
    except Exception:
        return ''
# === ARCUS ADDON v1 END ===


# === ARCUS ADDON v2 · 補齊三環：分析查詢地圖 + 測試模板 + 記錄批次 (2026-07-13) ===
# ① 把功能群地圖以「資料」寫進 .py，配 query_function_map 讓程式查得到（單一/批次）
# ② emit_selftest：測試環節的模板，配 stage_run_all 使用
# ③ append_arcus_log_batch：記錄環節一次寫多筆

# ── ① 分析查詢：功能群地圖（in-code 資料）＋查詢函式 ──

ARCUS_FUNCTION_MAP = {
    'C1 中央log': ['write_log', 'read_log', 'append_log_turn',
                   'read_log_for_display', 'trim_log_if_needed', 'append_token_log'],
    'C2 token估算': ['est_tokens'],
    'C3 判官': ['build_turn_review'],
    'C4 伺服器掛鉤': ['before_prompt_hook', 'after_response_hook', 'read_workspace_md'],
    'C5 行型模板': ['emit_indent', 'emit_module_header', 'emit_func', 'emit_try_return',
                    'emit_read_utf8', 'emit_write_utf8', 'emit_sibling_path',
                    'emit_now_stamp', 'emit_dict_get', 'emit_project_guard', 'emit_selftest'],
    'C6 暫存區引擎': ['stage_new', 'stage_run', 'stage_run_all', 'stage_promote', 'stage_cleanup'],
    'C7 維護log': ['append_arcus_log', 'read_arcus_log', 'append_arcus_log_batch'],
    'C8 分析查詢': ['query_function_map'],
    '基礎設施': ['generate_project_map', 'build_symbol_index', 'build_system_prompt',
                 'build_tasks_injection', 'parse_image_with_claude'],
}


def query_function_map(names=None):
    """分析查詢環節的批次/單一入口。
    names 不給 → 回整張地圖 {群: [函式,...]}；
    names 給字串或字串串列 → 批次回 {名稱: {'group':群, 'exists':真的在不在, 'doc':docstring首行}}。
    exists 走 globals() 即時內省，順便照出「地圖與真碼漂移」（聽最新版本）。"""
    if names is None:
        return {k: list(v) for k, v in ARCUS_FUNCTION_MAP.items()}
    where = {}
    for grp, fns in ARCUS_FUNCTION_MAP.items():
        for fn in fns:
            where[fn] = grp
    g = globals()
    out = {}
    for n in ([names] if isinstance(names, str) else names):
        obj = g.get(n)
        exists = callable(obj)
        doc = ''
        if exists and getattr(obj, '__doc__', None):
            doc = obj.__doc__.strip().split('\n')[0]
        out[n] = {'group': where.get(n, '(未列)'), 'exists': exists, 'doc': doc}
    return out


# ── ② 測試模板：產生 _selftest()，配 stage_run_all ──

def emit_selftest(checks):
    """產生 _selftest() 測試函式；checks 給 (運算式字串, 期望值字串) 串列，逐條 assert。
    空串列則產生一個永遠通過的 _selftest。標點與縮排由 emit_func 補。"""
    lines = []
    for expr, exp in (checks or []):
        lines.append('assert (%s) == (%s), %r' % (expr, exp, 'FAIL: %s' % expr))
    body = '\n'.join(lines) if lines else 'return True'
    return emit_func('_selftest', [], body, doc='自我測試（stage_run_all 會呼叫）')


# ── ③ 記錄批次：一次追加多筆 ──

def append_arcus_log_batch(entries):
    """一次追加多筆維護記錄到 arcus_log.md；entries 給 (分類, 內容) 串列。回傳成功筆數。"""
    n = 0
    for item in (entries or []):
        try:
            cat, content = item
        except Exception:
            continue
        if append_arcus_log(cat, content):
            n += 1
    return n
# === ARCUS ADDON v2 END ===


# === ARCUS ADDON MCP · 單一結構化工具伺服器（stdio / JSON-RPC 2.0，純標準庫，不裝套件）(2026-07-13) ===
# 目的：讓維護 arcus 的模型「天生只有一隻手」——唯一工具 arcus_do(cmd, payload)。
# cmd 以列舉限定，指令打不錯；不支援的操作回結構化錯誤，模型回報後由人介入，不陷入死迴圈。
# 這段只在 `python arcus_core.py --mcp` 時啟動；server.py 用 import 載入 core 完全不受影響。

# HTTP in-process MCP 模型下，工具分派跑在常駐伺服器行程，claude 子行程的環境變數傳不進來。
# 故改由 server 每次 /arcus/mcp 請求依網址的 project 參數，把當前專案根設進本執行緒區域變數；
# 分派同執行緒讀得到，且執行緒隔離、多專案並行不互相汙染。
_ARCUS_REQ_CTX = threading.local()


def _arcus_set_request_project(root):
    """server 在每次 /arcus/mcp 請求進來時呼叫：設定本執行緒此次請求的專案根（root 可為 None）。"""
    _ARCUS_REQ_CTX.project_root = root


def _arcus_clear_request_project():
    """請求處理完呼叫，清掉本執行緒的專案根，避免殘留影響同執行緒的下一個請求。"""
    _ARCUS_REQ_CTX.project_root = None


def _arcus_runtime_dir():
    # 優先用本執行緒此次 MCP 請求帶進來的專案根（見上方 _ARCUS_REQ_CTX 說明）；
    r = getattr(_ARCUS_REQ_CTX, 'project_root', None)
    if r and os.path.isdir(r):
        return os.path.abspath(r)
    # 退而讀環境變數（手動跑 --mcp 的 stdio 模型仍可用）；
    p = os.environ.get('ARCUS_PROJECT_PATH')
    if p and os.path.isdir(p):
        return os.path.abspath(p)
    # 都沒有才退回檔案自己的資料夾，arcus 自我維護時仍正常。
    return os.path.dirname(os.path.abspath(__file__))


# === ARCUS ADDON · 讀取放寬到全體專案，寫入維持原專案 (2026-07-19) ===
# 讀 = 所有專案任何檔案；寫 = 只有當前專案家目錄（arcus 開啟時即 runtime/）。
ARCUS_CHANGES_ROOT = os.path.abspath('/home/yuchi/.claude/openspec/changes')


def _arcus_read_path(rel):
    """讀取專用的路徑解析：先在當前專案家目錄找，找不到再回到全體專案根找。
    上限鎖在 openspec/changes 底下，家目錄其他位置（憑證、金鑰等）一律讀不到。"""
    root = ARCUS_CHANGES_ROOT
    base = _arcus_runtime_dir()
    cands = []
    if os.path.isabs(rel):
        cands.append(os.path.abspath(rel))
    else:
        cands.append(os.path.abspath(os.path.join(base, rel)))
        cands.append(os.path.abspath(os.path.join(root, rel)))
    allowed = []
    for c in cands:
        if c == root or c.startswith(root + os.sep):
            if c not in allowed:
                allowed.append(c)
    if not allowed:
        raise ValueError('read path escapes changes root: %s' % rel)
    for c in allowed:
        if os.path.exists(c):
            return c
    return allowed[0]


def _arcus_safe_path(rel):
    """把相對路徑鎖在 runtime 目錄內，擋掉 .. 逃逸。"""
    base = _arcus_runtime_dir()
    p = os.path.abspath(os.path.join(base, rel))
    if p != base and not p.startswith(base + os.sep):
        raise ValueError('path escapes runtime dir: %s' % rel)
    return p


# --------------------------------------------------------------------------
# draw：把繪圖引擎（learn-thesis 用的 draw_docx_figure.py）包成 arcus_do 的一個 cmd。
# 工具路徑寫死在這一層，模型不必、也無法自己 bash 找工具。
# 紀律焊進工具：一律先 --verify，不通過就回報告、不產圖；通過才產 SVG+PNG。
# 永不回傳 PNG 位元組，只回文字報告與檔案路徑/URL（省脈絡）。
# --------------------------------------------------------------------------
_ARCUS_DRAW_ENGINE = os.path.abspath(__file__)
_ARCUS_FIG_DIR = '/home/yuchi/forest-carbon-measurement/public/vecfigs'
_ARCUS_FIG_URL = 'https://forest-carbon.duckdns.org/vecfigs'


def _arcus_draw(payload):
    """draw：JSON 形狀規格 -> 向量圖（SVG 真向量 + PNG 後援）。
    payload：{spec: 規格物件, name?: 輸出檔名, verify_only?: 只自檢不產圖}。"""
    import subprocess, re as _re
    spec = payload.get('spec')
    if not isinstance(spec, dict):
        return {'ok': False, 'error': 'draw 需要 payload.spec（JSON 形狀規格物件）'}
    name = _re.sub(r'[^A-Za-z0-9_-]', '', str(payload.get('name') or 'arcus_fig')) or 'arcus_fig'
    verify_only = bool(payload.get('verify_only'))
    engine = _ARCUS_DRAW_ENGINE
    if not os.path.exists(engine):
        return {'ok': False, 'error': '繪圖引擎不存在：%s（人工介入）' % engine}
    os.makedirs(ARCUS_STAGING_DIR, exist_ok=True)
    spec_path = os.path.join(ARCUS_STAGING_DIR, 'draw_%s.json' % name)
    with open(spec_path, 'w', encoding='utf-8') as f:
        json.dump(spec, f, ensure_ascii=False)
    # 第一關：--verify（缺字＋版面幾何）。不通過一律不產圖。
    try:
        v = subprocess.run([sys.executable, engine, '--spec', spec_path, '--verify'],
                           capture_output=True, text=True, timeout=120)
    except Exception as e:
        return {'ok': False, 'error': '自檢執行失敗：%s' % e}
    report = (v.stdout or '').strip()
    verify_ok = (v.returncode == 0)
    if verify_only or not verify_ok:
        return {'ok': verify_ok, 'verify_ok': verify_ok, 'rendered': False,
                'report': report,
                'note': ('自檢通過（verify_only，未產圖）' if (verify_ok and verify_only)
                         else '自檢未通過，已跳過產圖；請依報告修正規格後重畫')}
    # 第二關：真的產圖。
    os.makedirs(_ARCUS_FIG_DIR, exist_ok=True)
    out_docx = os.path.join(_ARCUS_FIG_DIR, name + '.docx')
    try:
        r = subprocess.run([sys.executable, engine, '--spec', spec_path, '--out', out_docx],
                           capture_output=True, text=True, timeout=180)
    except Exception as e:
        return {'ok': False, 'verify_ok': True, 'report': report, 'error': '產圖執行失敗：%s' % e}
    if r.returncode != 0:
        return {'ok': False, 'verify_ok': True, 'report': report,
                'error': '產圖失敗（exit %d）：%s' % (r.returncode, (r.stderr or r.stdout or '').strip()[-400:])}
    return {'ok': True, 'verify_ok': True, 'rendered': True, 'report': report,
            'svg': os.path.join(_ARCUS_FIG_DIR, name + '.vecfig.svg'),
            'png': os.path.join(_ARCUS_FIG_DIR, name + '.vecfig.png'),
            'url': '%s/%s.vecfig.png' % (_ARCUS_FIG_URL, name),
            'note': '自檢通過並已產圖；SVG 為真向量。切勿把 PNG 讀回脈絡，看 URL 即可。'}


def _arcus_research(payload):
    """研究動作：對外搜尋、抓網頁、清掉雜質（選單/程式碼/樣式），回每個來源的乾淨原始正文。
    刻意不總結、不下結論——透明管道，判斷留給呼叫端自己做。
    payload: {queries:[關鍵詞...], budget?:int, max_results?:int, max_chars?:int}"""
    queries = payload.get('queries')
    if not queries or not isinstance(queries, list):
        return {'ok': False, 'error': 'research 需要 payload.queries（非空清單，一次把所有要查的關鍵詞傳進來，不收單數）'}
    budget = int(payload.get('budget') or 25)
    max_results = int(payload.get('max_results') or 6)
    max_chars = int(payload.get('max_chars') or 5000)
    try:
        _dr = sys.modules[__name__]
    except Exception as e:
        return {'ok': False, 'error': 'deep_research 匯入失敗：%s' % e}
    pool = _dr.BudgetPool(budget)
    seen = set()
    sources = []
    for q in queries:
        try:
            hits = _dr.search(q, pool, max_results)
        except Exception as e:
            sources.append({'query': q, 'error': '搜尋失敗：%s' % e})
            continue
        for hit in hits:
            url = hit.get('url') or ''
            if not url or url in seen:
                continue
            seen.add(url)
            text = _dr.fetch_and_prune(url, max_chars)
            if not text:
                continue
            sources.append({'query': q, 'title': hit.get('title', ''), 'url': url, 'text': text})
    ok = any('text' in s for s in sources)
    return {'ok': ok, 'sources': sources,
            'counts': {'queries': len(queries), 'urls': len([s for s in sources if 'text' in s]),
                       'budget_used': pool._used}}


def _arcus_discover(payload):
    """發現動作：給一批關鍵詞，只回候選網址清單（標題＋網址＋摘要），不抓正文。
    讓呼叫端先看到有哪些來源，審視後自己把要抓的網址放進 read 的 paths 一次抓回。
    payload: {queries:[關鍵詞...], budget?:int, max_results?:int}"""
    queries = payload.get('queries')
    if not queries or not isinstance(queries, list):
        return {'ok': False, 'error': 'discover 需要 payload.queries（非空清單，一次把所有要查的關鍵詞傳進來，不收單數）'}
    budget = int(payload.get('budget') or 15)
    max_results = int(payload.get('max_results') or 8)
    try:
        _dr = sys.modules[__name__]
    except Exception as e:
        return {'ok': False, 'error': 'deep_research 匯入失敗：%s' % e}
    pool = _dr.BudgetPool(budget)
    seen = set()
    candidates = []
    for q in queries:
        try:
            hits = _dr.search(q, pool, max_results)
        except Exception as e:
            candidates.append({'query': q, 'error': '搜尋失敗：%s' % e})
            continue
        for hit in hits:
            url = hit.get('url') or ''
            if not url or url in seen:
                continue
            seen.add(url)
            candidates.append({'query': q, 'title': hit.get('title', ''),
                               'url': url, 'snippet': hit.get('snippet', '')})
    return {'ok': bool(seen), 'candidates': candidates, 'count': len(seen)}


# ══════════════════════════════════════════════════════════════════════════════
# 記憶機制（抽屜索引＋整份作品檔）：科別名即抽屜，recall 只開被點名的抽屜
#   常駐 RAM 索引 _MEMORY_INDEX，持久化到單一 sibling 檔 memory_index.json
#   抽屜格式：{subject: {"terms": {title: {gloss, subjects, ts}},
#                        "works": {title: {summary, ref, subjects, ts}}}}
#   卡片自帶 subjects 清單，橋接名詞（同時屬多科）複製一張進每個抽屜，更新時逐一同步。
# ══════════════════════════════════════════════════════════════════════════════
MEMORY_INDEX_PATH = os.path.join(TOOLS_DIR, "memory_index.json")
MEMORY_OPEN_WARN_K = 3            # recall 一次開超過這麼多抽屜就警告，守住「不掃全庫」保證
MEMORY_TURN_KEY = 'memory_turn_counter'   # 抽取專用輪數計數器（與判官計數器解耦）
MEMORY_INTERVAL_KEY = 'memory_interval'
MEMORY_INTERVAL_DEFAULT = 4       # 每 4 輪抽取一次；判官每輪，抽取較稀疏以省錢

_MEMORY_INDEX = None              # None＝尚未載入；載入後為抽屜索引 dict
# 每輪收料的收件匣：各專案 memory.md 檔尾的機器區塊（人看的內容完全不動）
_MEMORY_INBOX_BEGIN = '<!-- arcus:memory-inbox:begin 機器維護區塊，請勿手動編輯 -->'
_MEMORY_INBOX_END = '<!-- arcus:memory-inbox:end -->'
MEMORY_ARCHIVE_TS_KEY = '_last_archive_ts'   # 存在索引裡的保留鍵，底線開頭、非抽屜
MEMORY_ARCHIVE_INTERVAL_SEC = 6 * 3600       # 每 6 小時歸檔一次
MEMORY_SCHED_POLL_SEC = 600                  # 輪詢間隔，同時決定重啟後補做的延遲上限
MEMORY_LOCK_PATH = os.path.join(TOOLS_DIR, '.memory_archive.lock')
MEMORY_INJECT_MAX = 8
_MEMORY_SCHED_STARTED = False


def _mem_now():
    import datetime
    return datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')


def _memory_load():
    """惰性載入抽屜索引到 RAM；只在第一次被存取時讀檔，之後直接回記憶體物件。"""
    global _MEMORY_INDEX
    if _MEMORY_INDEX is not None:
        return _MEMORY_INDEX
    import json
    data = {}
    try:
        if os.path.exists(MEMORY_INDEX_PATH):
            with open(MEMORY_INDEX_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
            if not isinstance(data, dict):
                data = {}
    except Exception as e:
        print('[memory] 索引載入失敗，改用空索引：%s' % e)
        data = {}
    _MEMORY_INDEX = data
    return _MEMORY_INDEX


def _memory_save():
    """把 RAM 索引原子寫回 memory_index.json（先寫 .tmp 再 rename，避免半截檔）。"""
    import json, tempfile
    idx = _memory_load()
    d = os.path.dirname(MEMORY_INDEX_PATH) or '.'
    fd, tmp = tempfile.mkstemp(prefix='.memidx_', suffix='.tmp', dir=d)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            json.dump(idx, f, ensure_ascii=False, indent=2)
        os.replace(tmp, MEMORY_INDEX_PATH)
        return True
    except Exception as e:
        print('[memory] 索引寫入失敗：%s' % e)
        try:
            if os.path.exists(tmp):
                os.unlink(tmp)
        except Exception:
            pass
        return False


def _subject_norm(s):
    return ''.join((s or '').split()).strip().lower()


def _subject_canonical(subject):
    """把新科別名吸附到既有的近似抽屜標籤，避免『經濟』與『經濟學』拆成兩抽屜。
       規則：正規化後完全相等 → 用既有標籤；去掉常見學科尾綴後字根相等 → 視為同抽屜。
       字根不相等的一律當新抽屜，不做前綴吸附（避免『經濟』誤併『經濟地理』）。"""
    idx = _memory_load()
    raw = (subject or '').strip()
    if not raw:
        return raw
    n = _subject_norm(raw)
    for k in idx:
        if _subject_norm(k) == n:
            return k
    SUFFIX = ('理論', '學', '論', '科')
    def stem(x):
        for suf in SUFFIX:
            if x.endswith(suf) and len(x) > len(suf):
                return x[:-len(suf)]
        return x
    ns = stem(n)
    for k in idx:
        if ns and stem(_subject_norm(k)) == ns:
            return k
    return raw


def _as_subject_list(subjects):
    """把單一科別名或清單正規化成去重後的正式抽屜標籤清單。"""
    if subjects is None:
        parts = []
    elif isinstance(subjects, str):
        parts = [subjects]
    elif isinstance(subjects, (list, tuple)):
        parts = list(subjects)
    else:
        parts = [str(subjects)]
    out, seen = [], set()
    for p in parts:
        c = _subject_canonical(str(p).strip())
        if c and c not in seen:
            seen.add(c)
            out.append(c)
    return out


def _memory_ensure_drawer(subject):
    idx = _memory_load()
    d = idx.get(subject)
    if not isinstance(d, dict):
        d = {"terms": {}, "works": {}}
        idx[subject] = d
    d.setdefault("terms", {})
    d.setdefault("works", {})
    return d


def remember(subjects, kind, title, body, ref=None):
    """把一張卡片歸檔進被點名的每個科別抽屜。
       subjects：科別名或清單（橋接名詞複製一張卡進每個抽屜）。
       kind：'term'（名詞用法，body＝一句用法）或 'work'（整份作品，body＝摘要、ref＝檔案路徑）。
       title：卡片主鍵（同科別同 title 視為同一張，保留較長的內容）。一次 save 收尾。"""
    subs = _as_subject_list(subjects)
    title = (title or '').strip()
    if not subs or not title:
        return {'ok': False, 'error': 'remember 需要 subjects 與 title'}
    kind = 'work' if kind == 'work' else 'term'
    body = body or ''
    ts = _mem_now()
    for s in subs:
        drawer = _memory_ensure_drawer(s)
        if kind == 'work':
            old = drawer['works'].get(title)
            summary = body
            r = ref
            if isinstance(old, dict) and len(old.get('summary') or '') > len(summary):
                summary = old.get('summary')
                if not r:
                    r = old.get('ref')
            drawer['works'][title] = {'summary': summary, 'ref': r, 'subjects': subs, 'ts': ts}
        else:
            old = drawer['terms'].get(title)
            gloss = body
            if isinstance(old, dict) and len(old.get('gloss') or '') > len(gloss):
                gloss = old.get('gloss')
            drawer['terms'][title] = {'gloss': gloss, 'subjects': subs, 'ts': ts}
    _memory_save()
    return {'ok': True, 'subjects': subs, 'kind': kind, 'title': title, 'drawers': len(subs)}


def recall(subjects, query=None):
    """只開被點名的抽屜，合併去重回卡片目錄；絕不掃全庫、不掃沒被點名的科別。
       subjects：科別名或清單；query：可選關鍵字，對 title／gloss／summary 做子字串過濾。"""
    subs = _as_subject_list(subjects)
    if not subs:
        return {'ok': False, 'error': 'recall 需要 subjects（科別名或清單）'}
    idx = _memory_load()
    warn = None
    if len(subs) > MEMORY_OPEN_WARN_K:
        warn = '一次開了 %d 個抽屜（超過建議上限 %d），請縮小科別範圍以維持精準' % (len(subs), MEMORY_OPEN_WARN_K)
    q = (query or '').strip().lower()
    terms, works, opened = {}, {}, []
    for s in subs:
        if s.startswith('_'):
            continue
        drawer = idx.get(s)
        if not isinstance(drawer, dict):
            continue
        opened.append(s)
        for t, card in (drawer.get('terms') or {}).items():
            if q and q not in t.lower() and q not in (card.get('gloss') or '').lower():
                continue
            ex = terms.get(t)
            if not ex or len(card.get('gloss') or '') > len(ex.get('gloss') or ''):
                terms[t] = card
        for t, card in (drawer.get('works') or {}).items():
            if q and q not in t.lower() and q not in (card.get('summary') or '').lower():
                continue
            ex = works.get(t)
            if not ex or len(card.get('summary') or '') > len(ex.get('summary') or ''):
                works[t] = card
    result = {
        'ok': True,
        'subjects': subs,
        'opened': opened,
        'terms': [dict(title=t, gloss=c.get('gloss'), ts=c.get('ts')) for t, c in terms.items()],
        'works': [dict(title=t, summary=c.get('summary'), ref=c.get('ref'), ts=c.get('ts')) for t, c in works.items()],
    }
    if warn:
        result['warn'] = warn
    return result


def list_subjects():
    """只回抽屜標籤與筆數（不回內容），讓代理在不確定時先窺看有哪些抽屜再決定要開哪個。"""
    idx = _memory_load()
    out = []
    for s in _memory_drawer_names(idx):
        d = idx.get(s) or {}
        out.append({'subject': s,
                    'terms': len(d.get('terms') or {}),
                    'works': len(d.get('works') or {})})
    return {'ok': True, 'subjects': out, 'count': len(out)}


def _call_model(prompt, model='haiku', timeout=120):
    """呼叫 claude CLI 子行程，--model 由參數決定（判官用 opus，抽取用 haiku 省錢）。
       與 _call_opus 同一條傳輸路徑；回傳 stdout 字串或 None。"""
    import subprocess as _sp, tempfile as _tf, time as _time
    tmp_path = None
    last_err = ''
    try:
        with _tf.NamedTemporaryFile(mode='w', suffix='.txt', encoding='utf-8', delete=False) as f:
            f.write(prompt)
            tmp_path = f.name
        script = ('source ~/.nvm/nvm.sh 2>/dev/null; '
                  'cat "%s" | claude --print --model %s --dangerously-skip-permissions' % (tmp_path, model))
        env = dict(os.environ)
        env['CLAUDE_SUBPROCESS'] = '1'
        env['PYTHONIOENCODING'] = 'utf-8'
        for attempt in range(3):
            if attempt:
                _time.sleep(3 * attempt)
            try:
                proc = _sp.run(['bash', '-lc', script], stdout=_sp.PIPE, stderr=_sp.STDOUT,
                               text=True, encoding='utf-8', errors='replace', env=env, timeout=timeout)
                out = (proc.stdout or '').strip()
                if proc.returncode == 0 and out:
                    return out
                last_err = '子行程回傳碼 %s；輸出：%s' % (proc.returncode, out[:200] or '空')
            except _sp.TimeoutExpired:
                last_err = '呼叫逾時（claude --print 超過 %s 秒沒回應）' % timeout
            except Exception as e:
                last_err = str(e)
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    print('[memory] 抽取模型呼叫失敗（已重試 3 次）: %s' % last_err)
    return None


def _memory_cheap_gate(user_text, bot_text):
    """便宜的純 Python 閘門：判斷這輪是否有值得抽取的實質學科內容，避免每輪都花模型呼叫。"""
    u = (user_text or '').strip()
    b = (bot_text or '').strip()
    combined = u + ' ' + b
    if len(combined) < 80:
        return False
    han = sum(1 for ch in combined if '一' <= ch <= '鿿')
    if han < 20:
        return False
    low = u.lower()
    if len(u) < 12 and any(k in low for k in ('hi', 'hello', '你好', '謝謝', 'ok', '好的')):
        return False
    return True


_MEMORY_EXTRACT_PROMPT = (
    '你是記憶抽取器。從下面這輪對話中，抽出討論到的「專有名詞／關鍵概念」及其在對話裡的用法。\n'
    '只抽真正的學科名詞或關鍵概念，忽略招呼語、操作指令、代名詞。\n'
    '每個名詞判斷它屬於哪個科別（例如：政治學、經濟學、統計學、心理學…），科別用最常見的通用名稱。\n'
    '嚴格只輸出一段 JSON 陣列，不要多餘文字，格式：\n'
    '[{"name":"名詞","usage":"這輪對話裡怎麼用它，一句話","subject":"科別"}]\n'
    '若沒有值得記的名詞，輸出 []。\n\n'
    '=== 使用者 ===\n%s\n\n=== 助理 ===\n%s\n'
)


def _memory_parse_terms(raw):
    """從模型回覆裡挖出 JSON 陣列（容忍 ``` 圍籬與前後雜訊）。"""
    import json, re
    s = (raw or '').strip()
    m = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', s, re.S)
    if m:
        s = m.group(1)
    else:
        i, j = s.find('['), s.rfind(']')
        if i != -1 and j != -1 and j > i:
            s = s[i:j + 1]
    try:
        data = json.loads(s)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _memory_extract(user_msg, response_text, subject=None):
    """獨立於判官的抽取動作：呼叫便宜模型抽名詞，逐一 remember。不『吐』給判官、不改判官輸出。"""
    prompt = _MEMORY_EXTRACT_PROMPT % ((user_msg or '')[:4000], (response_text or '')[:4000])
    raw = _call_model(prompt, model='haiku')
    if not raw:
        return {'ok': False, 'error': '抽取模型無回應'}
    items = _memory_parse_terms(raw)
    n = 0
    for it in items:
        if not isinstance(it, dict):
            continue
        name = (it.get('name') or '').strip()
        subj = (it.get('subject') or subject or '').strip()
        if not name or not subj:
            continue
        remember(subj, 'term', name, it.get('usage') or '')
        n += 1
    return {'ok': True, 'extracted': n}


# ══════════════════════════════════════════════════════════════════════════════
# 記憶卡片三段式機制（2026-07-18 重構）
#   第一段 每輪收料：_memory_capture_turn → 收件匣（各專案 memory.md 檔尾機器區塊）
#   第二段 每 6 小時歸檔：_memory_archive_batch → 整批收斂科別，索引只寫一次
#   第三段 回覆前注入：_memory_inject_block → 純字串比對，完全不經模型判斷
#   分工原則：Python 決定「何時、要不要做」，模型只決定「內容是什麼」。
# ══════════════════════════════════════════════════════════════════════════════
def _memory_drawer_names(idx):
    """只回真正的抽屜名；底線開頭是保留鍵（例如上次歸檔時間戳），不是科別。"""
    return sorted(k for k in idx.keys()
                  if not k.startswith('_') and isinstance(idx.get(k), dict))


def _atomic_write_text(path, text):
    """先寫暫存檔再換名，避免中途失敗留下半截檔。"""
    import tempfile
    d = os.path.dirname(path) or '.'
    fd, tmp = tempfile.mkstemp(prefix='.memtmp_', suffix='.tmp', dir=d)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            f.write(text)
        os.replace(tmp, path)
    except Exception:
        try:
            if os.path.exists(tmp):
                os.unlink(tmp)
        except Exception:
            pass
        raise


# ── 第一段：每輪收料 ──────────────────────────────────────────────────────────
def _memory_inbox_path(project_path):
    return os.path.join(project_path, 'memory.md')


def _memory_inbox_append(project_path, cards):
    """把卡片以每行一則 JSON 追加進 memory.md 檔尾的機器區塊。"""
    import json
    if not cards:
        return 0
    path = _memory_inbox_path(project_path)
    lines = [json.dumps(c, ensure_ascii=False) for c in cards]
    try:
        old = ''
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                old = f.read()
        if _MEMORY_INBOX_END in old:
            new = old.replace(_MEMORY_INBOX_END, '\n'.join(lines) + '\n' + _MEMORY_INBOX_END)
        else:
            new = (old.rstrip('\n') + '\n\n' + _MEMORY_INBOX_BEGIN + '\n'
                   + '\n'.join(lines) + '\n' + _MEMORY_INBOX_END + '\n')
        _atomic_write_text(path, new)
        return len(lines)
    except Exception as e:
        print('[memory] 收件匣寫入失敗：%s' % e)
        return 0


def _memory_inbox_read(project_path):
    """讀出收件匣裡的所有卡片；區塊不存在就回空清單。"""
    import json
    path = _memory_inbox_path(project_path)
    if not os.path.exists(path):
        return []
    try:
        with open(path, 'r', encoding='utf-8') as f:
            txt = f.read()
    except Exception:
        return []
    i, j = txt.find(_MEMORY_INBOX_BEGIN), txt.find(_MEMORY_INBOX_END)
    if i == -1 or j == -1 or j < i:
        return []
    out = []
    for line in txt[i + len(_MEMORY_INBOX_BEGIN):j].splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            d = json.loads(line)
            if isinstance(d, dict):
                out.append(d)
        except Exception:
            pass
    return out


def _memory_inbox_clear(project_path):
    """整段機器區塊移除；人手寫的內容一個字都不碰。"""
    path = _memory_inbox_path(project_path)
    if not os.path.exists(path):
        return False
    try:
        with open(path, 'r', encoding='utf-8') as f:
            txt = f.read()
        i, j = txt.find(_MEMORY_INBOX_BEGIN), txt.find(_MEMORY_INBOX_END)
        if i == -1 or j == -1 or j < i:
            return False
        new = txt[:i] + txt[j + len(_MEMORY_INBOX_END):]
        _atomic_write_text(path, new.rstrip('\n') + '\n')
        return True
    except Exception as e:
        print('[memory] 收件匣清空失敗：%s' % e)
        return False


def _memory_capture_worker(project_path, user_msg, response_text):
    """背景抽取：只抽名詞與用法，科別留成暫定值，不做正規化、不碰索引。"""
    try:
        prompt = _MEMORY_EXTRACT_PROMPT % ((user_msg or '')[:4000], (response_text or '')[:4000])
        raw = _call_model(prompt, model='haiku', timeout=90)
        if not raw:
            return
        ts = _mem_now()
        cards = []
        for it in _memory_parse_terms(raw):
            if not isinstance(it, dict):
                continue
            name = (it.get('name') or '').strip()
            if not name:
                continue
            cards.append({'title': name, 'kind': 'term',
                          'gloss': (it.get('usage') or '').strip(),
                          'subject_hint': (it.get('subject') or '').strip(),
                          'ts': ts})
        _memory_inbox_append(project_path, cards)
    except Exception as e:
        print('[memory] 收料失敗：%s' % e)


def _memory_capture_turn(project_path, user_msg, response_text):
    """每輪都跑（不再四輪一次）。執行緒刻意不設 daemon，行程結束會等它寫完才退出。"""
    try:
        if not project_path or not os.path.isdir(project_path):
            return None
        if not _memory_cheap_gate(user_msg, response_text):
            return None
        th = threading.Thread(target=_memory_capture_worker,
                              args=(project_path, user_msg, response_text),
                              name='memory-capture')
        th.daemon = False
        th.start()
        return {'ok': True, 'scheduled': True}
    except Exception as e:
        print('[memory] 收料觸發失敗：%s' % e)
        return None


def _memory_capture_work(project_path, title, summary, ref, subject_hint=''):
    """作品卡進收件匣：畫圖完成、檔案落地完成時呼叫，摘要可留空等歸檔補。"""
    if not project_path or not title:
        return 0
    return _memory_inbox_append(project_path, [{
        'title': str(title).strip(), 'kind': 'work',
        'summary': (summary or '').strip(), 'ref': ref,
        'subject_hint': subject_hint or '', 'ts': _mem_now()}])


# ── 第二段：每 6 小時歸檔 ─────────────────────────────────────────────────────
_MEMORY_ARCHIVE_PROMPT = (
    '你是記憶歸檔器。下面是一批剛收集到的名詞卡片，以及目前已存在的科別抽屜清單。\n'
    '任務：替每張卡片指定最終科別。能歸進現有抽屜就一定歸進去，不要輕易開新抽屜；\n'
    '只有當卡片明顯不屬於任何現有抽屜時，才給一個新的通用科別名稱。\n'
    '同名卡片合併成一則，解釋取較完整的版本。\n'
    '嚴格只輸出一段 JSON 陣列，不要多餘文字，格式：\n'
    '[{"title":"名詞","gloss":"一句話解釋","subject":"科別"}]\n\n'
    '=== 現有抽屜 ===\n%s\n\n=== 待歸檔卡片 ===\n%s\n'
)


def _memory_lock_acquire():
    """用佔位檔做跨行程互斥；搶不到就跳過這一輪，避免兩邊各自整份重寫把資料蓋掉。"""
    try:
        fd = os.open(MEMORY_LOCK_PATH, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
        os.write(fd, str(os.getpid()).encode('ascii'))
        os.close(fd)
        return True
    except FileExistsError:
        try:
            if time.time() - os.path.getmtime(MEMORY_LOCK_PATH) > 1800:
                os.unlink(MEMORY_LOCK_PATH)   # 前一個持鎖行程已死，回收陳舊鎖
                return _memory_lock_acquire()
        except Exception:
            pass
        return False
    except Exception:
        return False


def _memory_lock_release():
    try:
        if os.path.exists(MEMORY_LOCK_PATH):
            os.unlink(MEMORY_LOCK_PATH)
    except Exception:
        pass


def _memory_should_archive():
    """無狀態判斷：每次都重新用系統時鐘減去索引裡的上次歸檔時間戳。"""
    idx = _memory_load()
    try:
        last = float(idx.get(MEMORY_ARCHIVE_TS_KEY) or 0)
    except Exception:
        last = 0.0
    return (time.time() - last) >= MEMORY_ARCHIVE_INTERVAL_SEC


def _memory_stamp_archive(save=True):
    idx = _memory_load()
    idx[MEMORY_ARCHIVE_TS_KEY] = time.time()
    if save:
        _memory_save()


def _memory_merge_term(idx, subject, title, gloss):
    """就地併進索引，不觸發存檔（整批結束才寫一次）。"""
    s = _subject_canonical(subject)
    if not s or s.startswith('_'):
        return False
    d = idx.get(s)
    if not isinstance(d, dict):
        d = {'terms': {}, 'works': {}}
        idx[s] = d
    d.setdefault('terms', {})
    d.setdefault('works', {})
    old = d['terms'].get(title) or {}
    if len(gloss or '') < len(old.get('gloss') or ''):
        gloss = old.get('gloss')
    subs = old.get('subjects') or []
    if s not in subs:
        subs = list(subs) + [s]
    d['terms'][title] = {'gloss': gloss or '', 'subjects': subs, 'ts': _mem_now()}
    return True


def _memory_merge_work(idx, card):
    """作品卡就地併進索引；科別取收料當下的暫定值，沒有就丟進「作品」抽屜。"""
    s = _subject_canonical(card.get('subject_hint') or '作品')
    if not s or s.startswith('_'):
        return False
    title = (card.get('title') or '').strip()
    if not title:
        return False
    d = idx.get(s)
    if not isinstance(d, dict):
        d = {'terms': {}, 'works': {}}
        idx[s] = d
    d.setdefault('terms', {})
    d.setdefault('works', {})
    old = d['works'].get(title) or {}
    summary = card.get('summary') or ''
    if len(summary) < len(old.get('summary') or ''):
        summary = old.get('summary')
    d['works'][title] = {'summary': summary, 'ref': card.get('ref') or old.get('ref'),
                         'subjects': [s], 'ts': _mem_now()}
    return True


def _memory_archive_batch(force=False):
    """整批歸檔：讀完所有專案收件匣 → 交模型收斂科別 → 索引只寫一次 → 清空收件匣。"""
    if not force and not _memory_should_archive():
        return {'ok': True, 'skipped': 'not_due'}
    if not _memory_lock_acquire():
        return {'ok': True, 'skipped': 'locked'}
    try:
        projects, cards = [], []
        try:
            names = sorted(os.listdir(CHANGES))
        except Exception:
            names = []
        for name in names:
            p = os.path.join(CHANGES, name)
            if not os.path.isdir(p):
                continue
            got = _memory_inbox_read(p)
            if got:
                projects.append(p)
                cards.extend(got)
        if not cards:
            _memory_stamp_archive()
            return {'ok': True, 'cards': 0}
        idx = _memory_load()
        terms = [c for c in cards if (c.get('kind') or 'term') != 'work']
        works = [c for c in cards if (c.get('kind') or 'term') == 'work']
        merged = []
        if terms:
            drawers = _memory_drawer_names(idx)
            listing = '\n'.join(
                '- %s（名詞 %d 則）' % (k, len((idx.get(k) or {}).get('terms') or {}))
                for k in drawers) or '（目前沒有任何抽屜）'
            payload = '\n'.join('- %s：%s' % (c.get('title'), (c.get('gloss') or '')[:120])
                                for c in terms[:400])
            raw = _call_model(_MEMORY_ARCHIVE_PROMPT % (listing, payload),
                              model='haiku', timeout=180)
            merged = _memory_parse_terms(raw) if raw else []
            if not merged:
                # 模型無有效回應：收件匣原封不動保留，等下一輪再試，資料不會掉
                return {'ok': False, 'error': '歸檔模型無有效回應，收件匣保留'}
        n = 0
        for it in merged:
            if not isinstance(it, dict):
                continue
            title = (it.get('title') or '').strip()
            subj = (it.get('subject') or '').strip()
            if not title or not subj:
                continue
            if _memory_merge_term(idx, subj, title, (it.get('gloss') or '').strip()):
                n += 1
        w = sum(1 for c in works if _memory_merge_work(idx, c))
        _memory_stamp_archive(save=False)
        if not _memory_save():
            return {'ok': False, 'error': '索引寫入失敗，收件匣保留'}
        for p in projects:
            _memory_inbox_clear(p)
        return {'ok': True, 'cards': len(cards), 'terms': n, 'works': w,
                'projects': len(projects), 'drawers': len(_memory_drawer_names(idx))}
    finally:
        _memory_lock_release()


# ── 第三段：回覆前確定性注入 ─────────────────────────────────────────────────
def _memory_keyword_hits(user_msg, limit=None):
    """純字串比對：拿來訊息去比對卡片標題，命中就取用，完全不經模型判斷。"""
    text = (user_msg or '')
    if len(text) < 2:
        return []
    low = text.lower()
    idx = _memory_load()
    hits = []
    for s in _memory_drawer_names(idx):
        d = idx.get(s) or {}
        for kind, field in (('terms', 'gloss'), ('works', 'summary')):
            for title, card in (d.get(kind) or {}).items():
                if len(title) < 2 or not isinstance(card, dict):
                    continue
                if title.lower() in low:
                    hits.append({'subject': s, 'title': title,
                                 'body': (card.get(field) or '')[:200],
                                 'ref': card.get('ref')})
    hits.sort(key=lambda h: -len(h['title']))   # 長標題命中比短標題更明確
    return hits[:(limit or MEMORY_INJECT_MAX)]


def _memory_inject_block(user_msg):
    hits = _memory_keyword_hits(user_msg)
    if not hits:
        return ''
    lines = ['=== 記憶卡片（依訊息關鍵字自動命中，未經模型篩選）===']
    for h in hits:
        tail = '（檔案：%s）' % h['ref'] if h.get('ref') else ''
        lines.append('- [%s] %s：%s%s' % (h['subject'], h['title'], h['body'], tail))
    return '\n'.join(lines) + '\n\n'


# ── 作品卡：畫圖完成、檔案落地完成時自動開卡 ─────────────────────────────────
def _memory_work_project(payload):
    """作品卡要寫進哪個專案的收件匣；payload 沒帶就落在 arcus 自己身上。"""
    p = (payload or {}).get('project_path')
    if p and os.path.isdir(p):
        return p
    return os.path.join(CHANGES, 'arcus')


def _memory_note_draw(payload, res):
    """畫圖成功就開作品卡：標題、網址、路徑當下都齊全，摘要留空等歸檔補。"""
    try:
        if not isinstance(res, dict) or not res.get('ok') or not res.get('png'):
            return
        title = os.path.basename(res.get('png') or '')
        for suf in ('.vecfig.png', '.png'):
            if title.endswith(suf):
                title = title[:-len(suf)]
                break
        if not title:
            return
        _memory_capture_work(_memory_work_project(payload), title,
                             res.get('url') or '', res.get('png'),
                             (payload or {}).get('subject') or '')
    except Exception as e:
        print('[memory] 畫圖作品卡失敗：%s' % e)


def _memory_note_promote(payload, res):
    """檔案落地成功就開作品卡：標題取落地檔名，路徑取落地位置。"""
    try:
        if not isinstance(res, dict) or not res.get('ok'):
            return
        proj = _memory_work_project(payload)
        for item in (res.get('applied') or []):
            if not isinstance(item, dict):
                continue
            target = item.get('target') or item.get('to')
            if not target:
                continue
            _memory_capture_work(proj, os.path.basename(target), '', target,
                                 (payload or {}).get('subject') or '')
    except Exception as e:
        print('[memory] 落地作品卡失敗：%s' % e)


# ── 排程器：常駐行程載入本模組時自行啟動 ─────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# §87 待辦封存：已勾選的項目跟著記憶卡片走同一個歸檔時鐘
#
#   為什麼只搬已勾選：未勾選代表工作還沒做完，搬走就等於把待辦清單清空。
#   log_task.py 的 archive 屬於「隨用即刪」，整區搬走且含未勾項目，
#   適合人工收尾，不適合定時自動跑，所以這裡另寫只搬已勾項目的版本。
#
#   併發：和 log_task.py 用同一把 flock 排他鎖，兩邊同時動不會互相蓋掉。
# ─────────────────────────────────────────────────────────────────────────────

_TASKS_DONE_RE = re.compile(r'^\s*-\s*\[[xX]\]\s')
_TASKS_HEAD_RE = re.compile(r'^#{2,6}\s')
_NL = chr(10)


def _tasks_drop_empty_sections(lines):
    """章節底下的項目被搬光之後，光禿禿的標頭也一併移除；巢狀標頭要反覆掃到穩定。"""
    for _ in range(4):
        heads = [i for i, l in enumerate(lines) if _TASKS_HEAD_RE.match(l)]
        drop = set()
        for n, i in enumerate(heads):
            end = heads[n + 1] if n + 1 < len(heads) else len(lines)
            if not any(l.strip() for l in lines[i + 1:end]):
                drop.add(i)
        if not drop:
            break
        lines = [l for i, l in enumerate(lines) if i not in drop]
    return lines


def _tasks_archive_one(project_path):
    """單一專案：已勾項目連同所屬章節標題搬進 tasks-archive.md，未勾項目原地留著。"""
    tp = os.path.join(project_path, 'tasks.md')
    if not os.path.exists(tp):
        return 0
    fh = open(tp, 'r+', encoding='utf-8')
    try:
        fcntl.flock(fh.fileno(), fcntl.LOCK_EX)
        raw = fh.read()
        lines = raw.split(_NL)
        if not any(_TASKS_DONE_RE.match(l) for l in lines):
            return 0

        kept, taken, cur = [], [], '（無章節）'
        for l in lines:
            if _TASKS_HEAD_RE.match(l):
                cur = l.strip()
            if _TASKS_DONE_RE.match(l):
                if not taken or taken[-1][0] != cur:
                    taken.append([cur, []])
                taken[-1][1].append(l.rstrip())
            else:
                kept.append(l)

        n = sum(len(g[1]) for g in taken)
        stamp = time.strftime('%Y-%m-%d %H:%M', time.localtime())
        block = ['', '<!-- ===== 封存於 ' + stamp + ' ===== -->']
        for head, items in taken:
            block += ['', head] + items
        block.append('')

        ap = os.path.join(project_path, 'tasks-archive.md')
        if not os.path.exists(ap):
            with open(ap, 'w', encoding='utf-8') as nf:
                nf.write('# 已完成任務封存' + _NL)
        af = open(ap, 'a', encoding='utf-8')
        try:
            fcntl.flock(af.fileno(), fcntl.LOCK_EX)
            af.write(_NL.join(block))
            af.flush()
            os.fsync(af.fileno())
        finally:
            fcntl.flock(af.fileno(), fcntl.LOCK_UN)
            af.close()

        out = _NL.join(_tasks_drop_empty_sections(kept)).rstrip() + _NL
        fh.seek(0)
        fh.write(out)
        fh.truncate()
        fh.flush()
        os.fsync(fh.fileno())
        return n
    finally:
        fcntl.flock(fh.fileno(), fcntl.LOCK_UN)
        fh.close()


def _tasks_archive_batch():
    """走過全部專案；單一專案出錯只記下來，不讓其他專案跟著停。"""
    items = projects = 0
    errs = []
    if not os.path.isdir(CHANGES):
        return {'ok': False, 'error': 'changes 資料夾不存在'}
    for name in sorted(os.listdir(CHANGES)):
        p = os.path.join(CHANGES, name)
        if not os.path.isdir(p):
            continue
        try:
            n = _tasks_archive_one(p)
        except Exception as e:
            errs.append(name + '：' + str(e))
            continue
        if n:
            items += n
            projects += 1
    return {'ok': True, 'projects': projects, 'items': items, 'errors': errs}



def _memory_scheduler_loop():
    """輪詢而非精準睡眠：每次醒來重新推導要不要做，調時間／休眠／重啟後都會自動補做。"""
    while True:
        try:
            # 先問一次再做：記憶歸檔會蓋掉時間戳，問完才做才輪得到待辦封存
            due = _memory_should_archive()
            if due:
                print('[memory] 定時歸檔：%s' % (_memory_archive_batch(),))
                print('[tasks] 同批封存：%s' % (_tasks_archive_batch(),))
        except Exception as e:
            print('[memory] 排程執行失敗：%s' % e)
        time.sleep(MEMORY_SCHED_POLL_SEC)


def start_memory_scheduler():
    """臨時 MCP 行程用 --mcp 自檢擋掉；常駐端 import 就自動有排程，不需維護者記得加。"""
    global _MEMORY_SCHED_STARTED
    if _MEMORY_SCHED_STARTED or '--mcp' in sys.argv or '--spec' in sys.argv:
        return False
    _MEMORY_SCHED_STARTED = True
    th = threading.Thread(target=_memory_scheduler_loop, name='memory-archiver')
    th.daemon = True
    th.start()
    return True


def get_memory_interval(project_path):
    st = _load_state(project_path)
    try:
        return max(1, int(st.get(MEMORY_INTERVAL_KEY, MEMORY_INTERVAL_DEFAULT)))
    except Exception:
        return MEMORY_INTERVAL_DEFAULT


def _memory_maybe_extract(project_path, user_msg, response_text):
    """收尾掛鉤呼叫：用自己的計數器（與判官解耦），到期且過閘門才抽取。
       抽取丟到背景執行緒跑，不阻塞這一輪回應。"""
    try:
        st = _load_state(project_path)
        n = int(st.get(MEMORY_TURN_KEY, 0)) + 1
        st[MEMORY_TURN_KEY] = n
        _save_state(project_path, st)
        interval = get_memory_interval(project_path)
        if (n % interval) != 0:
            return None
        if not _memory_cheap_gate(user_msg, response_text):
            return None
        import threading
        th = threading.Thread(target=_memory_extract, args=(user_msg, response_text), daemon=True)
        th.start()
        return {'ok': True, 'scheduled': True, 'turn': n}
    except Exception as e:
        print('[memory] 抽取觸發失敗：%s' % e)
        return None


# === ARCUS ADDON · 設計思考原則學習（條件式原則，取代舊憲法法條權重）(2026-07-22) ===
# 舊機制用反向傳播調浮點權重，笨重又拿代理訊號當梯度。這裡改為：引擎用自己的語言判斷，
# 把使用者答案的本質萃取成一條白話條件式原則，存成可讀 jsonl；信心用離散計數表示
# （evidence 觀察次數／contradicted 被打臉次數），不做梯度、不存純量權重。
# 使用者的設計思考模式跨專案共用，故存於全域單檔。
import json as _pjson

_ARCUS_USER_MODEL = '/home/yuchi/.claude/arcus_user_model.jsonl'
_PRINCIPLE_PROMOTE_EVIDENCE = 2   # evidence 達此值且未被打臉 → tentative 升 live

def _principles_now():
    import datetime as _dt
    return _dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

def _principles_load():
    items = []
    try:
        with io.open(_ARCUS_USER_MODEL, encoding='utf-8') as _f:
            for _line in _f:
                _line = _line.strip()
                if not _line:
                    continue
                try:
                    items.append(_pjson.loads(_line))
                except Exception:
                    continue
    except FileNotFoundError:
        pass
    return items

def _principles_save(items):
    os.makedirs(os.path.dirname(_ARCUS_USER_MODEL), exist_ok=True)
    _tmp = _ARCUS_USER_MODEL + '.tmp'
    with io.open(_tmp, 'w', encoding='utf-8') as _f:
        for _it in items:
            _f.write(_pjson.dumps(_it, ensure_ascii=False) + '\n')
    os.replace(_tmp, _ARCUS_USER_MODEL)

def _principles_next_id(items):
    mx = 0
    for _it in items:
        try:
            _n = int(str(_it.get('id', '')).split('_')[-1])
            if _n > mx:
                mx = _n
        except Exception:
            pass
    return 'p_%d' % (mx + 1)

def _principle_promote(it):
    if int(it.get('evidence', 0)) >= _PRINCIPLE_PROMOTE_EVIDENCE and int(it.get('contradicted', 0)) == 0:
        it['status'] = 'live'

def principle_add(when, then, because=None, scope=None):
    """(1)收料：寫入一條候選原則，初始 tentative、evidence=1。完全相同的既有原則改為再觀察一次。"""
    when = (when or '').strip()
    then = (then or '').strip()
    if not when or not then:
        return {'ok': False, 'error': 'when 與 then 皆不可空'}
    items = _principles_load()
    for it in items:
        if (it.get('when', '').strip() == when and it.get('then', '').strip() == then
                and it.get('status') != 'retired'):
            it['evidence'] = int(it.get('evidence', 0)) + 1
            _principle_promote(it)
            it['last'] = _principles_now()
            _principles_save(items)
            return {'ok': True, 'id': it['id'], 'status': it['status'], 'deduped': True, 'record': it}
    rec = {
        'id': _principles_next_id(items),
        'when': when, 'then': then,
        'because': (because or '').strip(),
        'scope': (scope or '全域').strip(),
        'evidence': 1, 'contradicted': 0, 'status': 'tentative',
        'born': _principles_now(), 'last': _principles_now(),
    }
    items.append(rec)
    _principles_save(items)
    return {'ok': True, 'id': rec['id'], 'status': rec['status'], 'record': rec}

def principle_hit(pid, kind):
    """(2)固化／(4)打臉：confirm→evidence+1，達門檻且未被打臉即升 live；contradict→contradicted+1，退回 tentative。"""
    kind = (kind or '').strip()
    items = _principles_load()
    hit = None
    for it in items:
        if it.get('id') == pid:
            hit = it
            break
    if hit is None:
        return {'ok': False, 'error': '找不到原則 %s' % pid}
    if kind == 'confirm':
        hit['evidence'] = int(hit.get('evidence', 0)) + 1
        _principle_promote(hit)
    elif kind == 'contradict':
        hit['contradicted'] = int(hit.get('contradicted', 0)) + 1
        hit['status'] = 'tentative'
    else:
        return {'ok': False, 'error': 'kind 必須是 confirm 或 contradict'}
    hit['last'] = _principles_now()
    _principles_save(items)
    return {'ok': True, 'id': pid, 'status': hit['status'],
            'evidence': hit['evidence'], 'contradicted': hit['contradicted']}

def principle_list(status=None):
    """列出原則供檢視；status 可填 live／tentative／retired，不填給全部。"""
    items = _principles_load()
    if status:
        items = [it for it in items if it.get('status') == status]
    return {'ok': True, 'count': len(items), 'principles': items}

def _principles_context_block(limit=20):
    """(3)套用：把 live 原則做成注入文字，餵進每輪系統提示。無 live 原則回空字串、不佔提示。"""
    items = [it for it in _principles_load() if it.get('status') == 'live']
    if not items:
        return ''
    items = items[-limit:]
    lines = ['\n【你對這位使用者已萃取的設計思考原則——命中就照做，並附一句報備「依你〔id〕原則我選了X（不對就說一聲）」；被更正就呼叫 principle_hit(id,"contradict")】']
    for it in items:
        lines.append('- 〔%s｜%s〕當 %s → 偏好 %s（本質：%s）'
                     % (it.get('id'), it.get('scope', '全域'),
                        it.get('when'), it.get('then'), it.get('because', '')))
    return '\n'.join(lines) + '\n'

def _arcus_do_dispatch(cmd, payload):
    """arcus_do 的唯一分派點。回傳可 JSON 序列化的 dict（含 ok 欄位）。"""
    payload = payload or {}
    if cmd == 'map':
        base = _arcus_runtime_dir()
        if base == os.path.dirname(os.path.abspath(__file__)):
            # arcus 自身：回函式群地圖（自我維護靠它省行數）
            return {'ok': True, 'result': query_function_map(payload.get('names'))}
        # 一般專案：回該專案的檔案結構地圖，讓模型先看到有哪些檔、不盲目找檔
        return {'ok': True, 'result': generate_project_map(base)}
    if cmd == 'read':
        paths = payload.get('paths')
        if not paths or not isinstance(paths, list):
            return {'ok': False, 'error': 'read 需要 payload.paths（非空清單），一次把要讀的檔全部傳進來'}
        files = []
        _max_chars = int(payload.get('max_chars') or 8000)
        _spill_over = int(payload.get('spill_over') or 500000)
        _off = payload.get('offset')
        _lim = payload.get('limit')
        _sliced = (_off is not None) or (_lim is not None)
        for rel in paths:
            entry = {'req': rel}
            try:
                if isinstance(rel, str) and rel.startswith('spill://'):
                    # 讀回先前外溢存下的正文（可帶 offset/limit 分段）
                    p = _arcus_spill_path(rel)
                    with open(p, 'r', encoding='utf-8') as f:
                        entry['text'] = f.read()
                elif isinstance(rel, str) and (rel.startswith('http://') or rel.startswith('https://')):
                    _dr = sys.modules[__name__]
                    txt = _dr.fetch_and_prune(rel, _max_chars)
                    entry['url'] = rel
                    if txt:
                        entry['text'] = txt
                    else:
                        entry['error'] = '抓不到內容（可能是動態頁或空頁）'
                else:
                    p = _arcus_read_path(rel)
                    entry['path'] = p
                    if p.lower().endswith(('.pdf', '.docx', '.doc', '.pptx', '.zip',
                                           '.png', '.jpg', '.jpeg', '.gif', '.epub')):
                        entry['error'] = ('二進位檔（%s）不走 read（read 只讀純文字）。'
                                          'PDF／論文改用 paper_ingest 抽取切段，再用 '
                                          'paper_quote／paper_cards 取用。'
                                          % os.path.splitext(p)[1])
                    else:
                        with open(p, 'r', encoding='utf-8') as f:
                            entry['text'] = f.read()
                # 拿到正文後：有指定切片就切一段回傳；否則正文過大就外溢存檔、只回預覽＋handle，
                # 讓每筆回傳都小到不會被外層轉存到專案外的絕對路徑（那正是引擎讀不回的死路）。
                if 'text' in entry:
                    if _sliced:
                        _t = entry['text']
                        _s = int(_off or 0)
                        entry['text'] = _t[_s:(_s + int(_lim)) if _lim is not None else None]
                        entry['chars_total'] = len(_t)
                        entry['offset'] = _s
                    elif len(entry['text']) > _spill_over:
                        _full = entry.pop('text')
                        entry['saved'] = _arcus_spill_write(rel, _full)
                        entry['chars'] = len(_full)
                        entry['preview'] = _full[:400]
            except UnicodeDecodeError:
                entry['error'] = ('這個檔不是 UTF-8 純文字（可能是 PDF 或其他二進位）。'
                                  'PDF／論文請改用 paper_ingest。')
            except Exception as e:
                entry['error'] = str(e)
            files.append(entry)
        return {'ok': all(('text' in e or 'saved' in e) for e in files), 'files': files}
    if cmd == 'query':
        files = payload.get('files') or []
        keywords = payload.get('keywords') or []
        hits = []
        for rel in files:
            try:
                p = _arcus_spill_path(rel) if isinstance(rel, str) and rel.startswith('spill://') else _arcus_read_path(rel)
                with open(p, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
            except Exception as e:
                hits.append({'file': rel, 'error': str(e)})
                continue
            for i, ln in enumerate(lines, 1):
                for kw in keywords:
                    if kw in ln:
                        hits.append({'file': rel, 'line': i, 'kw': kw, 'text': ln.rstrip('\n')})
        return {'ok': True, 'hits': hits}
    if cmd == 'stage_new':
        drafts = payload.get('drafts')
        if not drafts or not isinstance(drafts, list):
            return {'ok': False, 'error': 'stage_new 需要 payload.drafts（非空清單，每項 {name, draft}；draft 是 JSON 信封 {schema_version:1, kind:code/content/append/op, target, content 或 op；追加到檔尾用 append、文字放 content 欄}）'}
        results = []
        for it in drafts:
            it = it or {}
            name = it.get('name')
            draft = it.get('draft')
            if not name or not isinstance(draft, dict):
                results.append({'ok': False, 'name': name, 'error': '每項需含 name 與 draft(物件)'})
                continue
            errs = draft_validate_schema(draft)
            if errs:
                results.append({'ok': False, 'name': name, 'error': '；'.join(errs)})
                continue
            path = stage_new_draft(name, draft)
            results.append({'ok': bool(path), 'name': name, 'draft_path': path})
        return {'ok': all(r.get('ok') for r in results), 'drafts': results}
    if cmd == 'stage_run':
        drafts = payload.get('drafts') or payload.get('caches') or []
        ok_all, report = stage_typecheck_files(drafts)
        return {'ok': ok_all, 'report': report}
    if cmd == 'promote':
        drafts = payload.get('drafts')
        if not drafts and payload.get('mapping'):
            return {'ok': False, 'error': 'promote 已改為草稿制：請用 drafts=[draft_*.json 路徑]，不再收 mapping'}
        if not drafts or not isinstance(drafts, list):
            return {'ok': False, 'error': 'promote 需要 payload.drafts（draft_*.json 路徑清單）'}
        _res = stage_promote_drafts(drafts)
        _memory_note_promote(payload, _res)
        return _res
    if cmd == 'cleanup':
        drafts = payload.get('drafts') or payload.get('caches') or []
        n = stage_cleanup(drafts)
        return {'ok': True, 'deleted': n}
    if cmd == 'log':
        entries = payload.get('entries')
        if entries:
            n = append_arcus_log_batch([tuple(e) for e in entries])
            return {'ok': n > 0, 'wrote': n}
        cat = payload.get('category')
        content = payload.get('content')
        if cat and content is not None:
            ok = append_arcus_log(cat, content)
            return {'ok': bool(ok), 'wrote': 1 if ok else 0}
        return {'ok': False, 'error': 'log 需要 payload.entries 或 payload.category+content'}
    if cmd in ('paper_ingest', 'paper_cards', 'paper_quote', 'paper_list', 'paper_verify', 'paper_delete'):
        # 延後匯入：臨時 MCP 行程每輪都載入本模組，不該為此付抽取套件的啟動成本
        _pc = sys.modules[__name__]
        if cmd == 'paper_ingest':
            srcs = (payload.get('sources') or payload.get('source')
                    or payload.get('ids') or payload.get('id'))
            if not srcs:
                return {'ok': False, 'error': 'paper_ingest 需要 payload.ids（arXiv id／DOI）或 sources（直接 pdf 網址／虛擬機路徑）'}
            return _pc.ingest(srcs, force=bool(payload.get('force')))
        if cmd == 'paper_cards':
            return _pc.cards(payload.get('ids'), payload.get('fields'))
        if cmd == 'paper_quote':
            pid = payload.get('id')
            if not pid:
                return {'ok': False, 'error': 'paper_quote 需要 payload.id（可給 id 前綴或標題子字串）'}
            return _pc.quote(pid, payload.get('keys'), payload.get('search'))
        if cmd == 'paper_verify':
            pid = payload.get('id')
            if not pid:
                return {'ok': False, 'error': 'paper_verify 需要 payload.id'}
            return _pc.verify(pid)
        if cmd == 'paper_delete':
            pid = payload.get('id')
            if not pid:
                return {'ok': False, 'error': 'paper_delete 需要 payload.id(要刪的卡號)'}
            return _pc.delete_paper(pid)
        return _pc.list_papers()

    if cmd == 'publish':
        name = (payload.get('name') or '').strip()
        content = payload.get('content')
        src = payload.get('src')
        if not name or ('/' in name) or name.startswith('.'):
            return {'ok': False, 'error': 'publish 需要 payload.name（單純檔名，不可含斜線）'}
        root = os.path.join(CHANGES, '_page')
        try:
            os.makedirs(root)
        except OSError:
            pass
        dst = os.path.join(root, name)
        try:
            if content is not None:
                _atomic_write_text(dst, content)
            elif src:
                with open(_arcus_read_path(src), 'rb') as f:
                    data = f.read()
                with open(dst, 'wb') as f:
                    f.write(data)
            else:
                return {'ok': False, 'error': 'publish 需要 payload.content（正文）或 payload.src（來源檔）'}
        except Exception as e:
            return {'ok': False, 'error': str(e)}
        return {'ok': True, 'name': name, 'path': dst,
                'url': 'https://forest-carbon.duckdns.org/arcus/page/' + name,
                'bytes': os.path.getsize(dst)}
    if cmd == 'draw':
        _res = _arcus_draw(payload)
        _memory_note_draw(payload, _res)
        return _res
    if cmd == 'discover':
        return _arcus_discover(payload)
    if cmd == 'research':
        return _arcus_research(payload)
    if cmd == 'recall':
        return recall(payload.get('subjects') or payload.get('subject'), payload.get('query'))
    if cmd == 'list_subjects':
        return list_subjects()
    if cmd == 'remember':
        return remember(payload.get('subjects') or payload.get('subject'),
                        payload.get('kind') or 'term',
                        payload.get('title'),
                        payload.get('body') or payload.get('gloss') or payload.get('summary') or '',
                        payload.get('ref'))
    if cmd == 'principle_add':
        return principle_add(payload.get('when'), payload.get('then'),
                             payload.get('because'), payload.get('scope'))
    if cmd == 'principle_hit':
        return principle_hit(payload.get('id') or payload.get('pid'), payload.get('kind'))
    if cmd == 'principle_list':
        return principle_list(payload.get('status'))
    if cmd == 'restart':
        import subprocess as _sp
        _delay = payload.get('delay')
        try:
            _delay = int(_delay)
        except (TypeError, ValueError):
            _delay = 2
        _delay = max(1, min(_delay, 30))
        # arcus 由 systemd 管（Restart=always）。自我重啟不可直接 systemctl restart：
        # 停止本服務那一刀會連坐殺光同一 cgroup 的行程（含發動重啟者）。改用 systemd-run
        # 排一個瞬時計時器，在服務 cgroup 之外、_delay 秒後才 systemctl restart arcus；
        # 延遲讓本輪工具回應先送回引擎。NOPASSWD sudo 已就位，--collect 讓瞬時單元跑完自刪。
        try:
            _sp.run(
                ['sudo', '-n', 'systemd-run', '--collect',
                 '--on-active=%ds' % _delay, 'systemctl', 'restart', 'arcus'],
                check=True, capture_output=True, text=True, timeout=15)
        except Exception as _e:
            _err = getattr(_e, 'stderr', '') or str(_e)
            return {'ok': False, 'error': '重啟排程失敗：%s' % _err}
        return {'ok': True, 'result': '重啟已排程：%d 秒後由 systemd 於服務 cgroup 外執行 '
                'systemctl restart arcus，本輪回應送出後即生效，約 %d+3 秒後新版程式上線。' % (_delay, _delay)}
    return {'ok': False, 'error': '不支援的 cmd：%s（人工介入）' % cmd}


_ARCUS_TOOL = {
    'name': 'arcus_do',
    'description': (
        'arcus 引擎維護的唯一入口。cmd 選一個動作、payload 帶參數：\n'
        "map     payload {names?: 字串或字串陣列} → 功能地圖/漂移查詢\n"
        "read    payload {paths:[檔名 或 http(s) 網址 或 spill:// handle...], max_chars?, spill_over?, offset?, limit?} → 一次讀完整批：小正文直接回；大正文（>spill_over，預設500000字）自動外溢存檔，只回 preview＋saved(spill:// handle)＋chars，避免結果過大被外層轉存到讀不回的絕對路徑；要全文就把該 handle 放回 paths 再 read（可帶 offset/limit 分段），或用 query 掃該 handle 找關鍵詞\n"
        "query   payload {files: [檔名], keywords: [關鍵字]} → 回命中行號與文字\n"
        "stage_new payload {drafts:[{name, draft}]} → 建 draft_*.json（draft 是信封 {schema_version:1, kind:code/content/append/op, target, content 或 op}；code/content/append 都用 content 欄，op 用 op 欄；追加到檔尾（如 memory.md）用 kind=append，不要用 content，content 會整檔覆寫），回草稿路徑\n"
        "stage_run payload {drafts:[草稿路徑]} → 型別感知測試(code只編譯/content驗解析/op驗前置)，回 ok 與報表；不執行程式碼\n"
        "promote payload {drafts:[草稿路徑]} → 型別測試全過後原子落地(code/content寫檔、op搬移或刪除)，全有全無、失敗整批回滾\n"
        "cleanup payload {drafts:[草稿路徑]} → 刪除草稿\n"
        "log     payload {entries: [[分類,內容]...]} 或 {category, content} → 寫 arcus_log.md\n"
        "draw    payload {spec: JSON形狀規格, name?, verify_only?} → 先自檢再產向量圖，回文字報告與 PNG 的 URL\n"
        "discover payload {queries:[關鍵詞...], budget?, max_results?} → 對外搜尋只回候選網址清單（標題＋網址＋摘要，不抓正文）；審視後把要抓的網址放進 read 的 paths 一次抓回\n"
        "publish payload {name: 檔名, content?: 正文, src?: 來源檔} → 把檔案送上公開網頁，回可直接開啟的網址\n"
        "paper_ingest payload {ids 或 sources:[arXiv id／pdf 網址／虛擬機路徑...]} → 從權威索引庫取逐字著錄（作者/科別/日期）並抽取全文建卡；裸關鍵字不建卡，請先 discover 挑定再交識別碼\n"
        "paper_list / paper_cards payload {ids?, fields?} / paper_quote payload {id, keys?, search?} / paper_verify payload {id} → 列卡片、取欄位、取原文引句、驗證。id 可給 id 前綴或標題子字串；paper_quote 不給 keys 會回該篇可用段落鍵，給 search 依內文片語搜出對應段落；read 不讀 PDF 等二進位，PDF 一律先 paper_ingest\n"
        "paper_delete payload {id} → 刪掉一張卡片(同時移除卡片目錄與 index.json 那筆,兩步一起做,避免留殘卡)\n"
        "remember payload {subjects:科別(字串或陣列), title, body, kind?='term', ref?} → 寫一張記憶卡片進長期抽屜\n"
        "recall   payload {subjects?:科別, query?:關鍵字} → 讀卡片的唯一正道；一律用 recall 讀卡片，禁止改用 read 去讀 memory.md 或任何卡片原始檔（原始檔是內部結構、含未整理內容，直接讀會拿到雜訊，recall 才回乾淨卡片）\n"
        "list_subjects → 列出記憶抽屜現有哪些科別\n"
        "restart payload {delay?:秒數，預設2、範圍1-30} → 讓 arcus 自我重啟（由 systemd 在服務 cgroup 外排程 systemctl restart，不會自殺式連坐）。改動 arcus_core.py／server.py 並 promote 後，磁碟已是新碼但記憶體仍跑舊版，呼叫本指令重啟才會載入新碼；本輪回應送出後約 delay+3 秒新版上線"
    ),
    'inputSchema': {
        'type': 'object',
        'properties': {
            'cmd': {
                'type': 'string',
                'enum': ['map', 'read', 'query', 'discover', 'research', 'stage_new', 'stage_run', 'promote', 'cleanup', 'log', 'draw', 'publish', 'paper_ingest', 'paper_cards', 'paper_quote', 'paper_list', 'paper_verify', 'paper_delete',
                         'remember', 'recall', 'list_subjects', 'restart'],
                'description': '要執行的動作',
            },
            'payload': {'type': 'object', 'description': '依 cmd 而定的參數，見工具說明'},
        },
        'required': ['cmd'],
    },
}


def _arcus_mcp_send(msg):
    sys.stdout.write(json.dumps(msg, ensure_ascii=False) + '\n')
    sys.stdout.flush()


def _arcus_mcp_handle(req):
    m = req.get('method')
    rid = req.get('id')
    if m == 'initialize':
        pv = (req.get('params') or {}).get('protocolVersion', '2025-06-18')
        return {'jsonrpc': '2.0', 'id': rid, 'result': {
            'protocolVersion': pv,
            'capabilities': {'tools': {}},
            'serverInfo': {'name': 'arcus', 'version': '1.0'},
        }}
    if m == 'tools/list':
        return {'jsonrpc': '2.0', 'id': rid, 'result': {'tools': [_ARCUS_TOOL]}}
    if m == 'tools/call':
        p = req.get('params') or {}
        args = p.get('arguments') or {}
        name = p.get('name')
        if name != 'arcus_do':
            return {'jsonrpc': '2.0', 'id': rid, 'result': {
                'content': [{'type': 'text', 'text': json.dumps({'ok': False, 'error': '未知工具 %s' % name}, ensure_ascii=False)}],
                'isError': True}}
        try:
            out = _arcus_do_dispatch(args.get('cmd'), args.get('payload'))
        except Exception as e:
            out = {'ok': False, 'error': '%s: %s' % (type(e).__name__, e)}
        return {'jsonrpc': '2.0', 'id': rid, 'result': {
            'content': [{'type': 'text', 'text': json.dumps(out, ensure_ascii=False)}],
            'isError': not out.get('ok', False)}}
    if m and m.startswith('notifications/'):
        return None
    if rid is not None:
        return {'jsonrpc': '2.0', 'id': rid, 'error': {'code': -32601, 'message': 'method not found: %s' % m}}
    return None


def _arcus_mcp_serve():
    """stdio 主迴圈：一行一個 JSON-RPC 訊息，讀到 EOF 乾淨結束。"""
    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue
        try:
            req = json.loads(line)
        except Exception:
            continue
        resp = _arcus_mcp_handle(req)
        if resp is not None:
            _arcus_mcp_send(resp)


# === ARCUS ADDON · 歸檔改三段式，慢的模型呼叫不佔鎖 (2026-07-19) ===
# 後定義覆蓋前定義。收料（短鎖）→ 模型分批收斂（不佔鎖）→ 併入索引（短鎖）。
# 卡片一旦落進成品區就清空收件匣；模型失敗時原料留在成品區，下一輪重試，卡片不會掉。

MEMORY_SPOOL_DIR = os.path.join(TOOLS_DIR, '_memory_spool')
MEMORY_ARCHIVE_BATCH_SIZE = 50
MEMORY_ARCHIVE_CALL_TIMEOUT = 600
MEMORY_LOCK_STALE_SEC = 1800


def _memory_lock_acquire():
    """取得歸檔佔用標記；持有者行程已死就立刻回收，不必等滿陳舊秒數。"""
    def _grab():
        try:
            fd = os.open(MEMORY_LOCK_PATH, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
            os.write(fd, str(os.getpid()).encode('utf-8'))
            os.close(fd)
            return True
        except FileExistsError:
            return False
        except Exception as e:
            print('[memory] 取鎖失敗：%s' % e)
            return False
    if _grab():
        return True
    try:
        with open(MEMORY_LOCK_PATH, 'r', encoding='utf-8') as f:
            pid = int((f.read() or '0').strip() or 0)
    except Exception:
        pid = 0
    dead = False
    if pid <= 0:
        dead = True
    elif pid != os.getpid():
        try:
            os.kill(pid, 0)
        except ProcessLookupError:
            dead = True
        except Exception:
            pass
    stale = False
    try:
        stale = (time.time() - os.path.getmtime(MEMORY_LOCK_PATH)) > MEMORY_LOCK_STALE_SEC
    except Exception:
        pass
    if not (dead or stale):
        return False
    print('[memory] 回收佔用標記（%s，pid=%s）' % ('持有者已死' if dead else '超過陳舊秒數', pid))
    try:
        os.unlink(MEMORY_LOCK_PATH)
    except Exception:
        pass
    return _grab()


def _memory_spool_ensure():
    try:
        os.makedirs(MEMORY_SPOOL_DIR)
    except OSError:
        pass
    return MEMORY_SPOOL_DIR


def _memory_spool_write(name, obj):
    import json
    _memory_spool_ensure()
    p = os.path.join(MEMORY_SPOOL_DIR, name)
    _atomic_write_text(p, json.dumps(obj, ensure_ascii=False, indent=1))
    return p


def _memory_spool_list():
    _memory_spool_ensure()
    try:
        return sorted(os.path.join(MEMORY_SPOOL_DIR, f)
                      for f in os.listdir(MEMORY_SPOOL_DIR) if f.endswith('.json'))
    except Exception:
        return []


def _memory_spool_load(p):
    import json
    try:
        with open(p, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return None


def _memory_archive_batch(force=False):
    """三段式歸檔：收料落地（短鎖）→ 模型分批收斂（不佔鎖）→ 併入索引（短鎖）。"""
    if not force and not _memory_should_archive():
        return {'ok': True, 'skipped': 'not_due'}

    picked = 0
    if _memory_lock_acquire():
        try:
            projects, cards = [], []
            try:
                names = sorted(os.listdir(CHANGES))
            except Exception:
                names = []
            for nm in names:
                pp = os.path.join(CHANGES, nm)
                if not os.path.isdir(pp):
                    continue
                got = _memory_inbox_read(pp)
                if got:
                    for c in got:
                        if isinstance(c, dict):
                            c.setdefault('project', nm)
                    projects.append(pp)
                    cards.extend(got)
            if cards:
                _memory_spool_write('raw_%d_%d.json' % (int(time.time()), os.getpid()),
                                    {'stage': 'raw', 'cards': cards,
                                     'projects': [os.path.basename(x) for x in projects]})
                for pp in projects:
                    _memory_inbox_clear(pp)
                picked = len(cards)
        finally:
            _memory_lock_release()

    failed = 0
    for sp in _memory_spool_list():
        obj = _memory_spool_load(sp)
        if not obj or obj.get('stage') != 'raw':
            continue
        cards = obj.get('cards') or []
        terms = [c for c in cards if (c.get('kind') or 'term') != 'work']
        works = [c for c in cards if (c.get('kind') or 'term') == 'work']
        idx = _memory_load()
        drawers = _memory_drawer_names(idx)
        listing = _NL.join('- %s（名詞 %d 則）' % (k, len((idx.get(k) or {}).get('terms') or {}))
                           for k in drawers) or '（目前沒有任何抽屜）'
        merged, bad = [], 0
        for i in range(0, len(terms), MEMORY_ARCHIVE_BATCH_SIZE):
            chunk = terms[i:i + MEMORY_ARCHIVE_BATCH_SIZE]
            payload = _NL.join('- %s：%s' % (c.get('title'), (c.get('gloss') or '')[:120])
                               for c in chunk)
            raw = _call_model(_MEMORY_ARCHIVE_PROMPT % (listing, payload),
                              model='haiku', timeout=MEMORY_ARCHIVE_CALL_TIMEOUT)
            got = _memory_parse_terms(raw) if raw else []
            if got:
                merged.extend(got)
            else:
                bad += 1
        if terms and not merged:
            failed += 1
            continue
        obj['stage'] = 'ready'
        obj['merged'] = merged
        obj['works'] = works
        obj['batches_failed'] = bad
        _memory_spool_write(os.path.basename(sp), obj)

    if not _memory_lock_acquire():
        return {'ok': True, 'skipped': 'merge_locked', 'picked': picked}
    try:
        idx = _memory_load()
        n = w = files = 0
        used = []
        for sp in _memory_spool_list():
            obj = _memory_spool_load(sp)
            if not obj or obj.get('stage') != 'ready':
                continue
            for it in (obj.get('merged') or []):
                if not isinstance(it, dict):
                    continue
                title = (it.get('title') or '').strip()
                subj = (it.get('subject') or '').strip()
                if not title or not subj:
                    continue
                if _memory_merge_term(idx, subj, title, (it.get('gloss') or '').strip()):
                    n += 1
            for c in (obj.get('works') or []):
                if _memory_merge_work(idx, c):
                    w += 1
            used.append(sp)
            files += 1
        _memory_stamp_archive(save=False)
        if not _memory_save():
            return {'ok': False, 'error': '索引寫入失敗，成品區保留'}
        for sp in used:
            try:
                os.unlink(sp)
            except Exception:
                pass
        return {'ok': True, 'picked': picked, 'spool_files': files, 'terms': n,
                'works': w, 'model_failed': failed,
                'drawers': len(_memory_drawer_names(idx))}
    finally:
        _memory_lock_release()


# 模組載入即嘗試啟動歸檔排程；臨時 MCP 行程由函式內的 --mcp 自檢擋掉。
try:
    start_memory_scheduler()
except Exception as _e:
    print('[memory] 排程啟動失敗：%s' % _e)


# === ARCUS ADDON MCP END ===


# === ARCUS ADDON · 白名單擴充覆寫 generate_project_map (2026-07-18) ===
# 後定義覆蓋前定義：本函式與上方 L884 同名，模組載入時最後執行、取代舊版。
# 唯一差異：改為黑名單，除二進位雜訊外所有副檔名都列入專案地圖。
# 其餘邏輯與原函式逐字相同，避免行為漂移。

def generate_project_map(project_path, project_name=None):
    """掃描專案目錄，生成帶路徑和說明的結構地圖文字（只展開最近 7 天異動路徑）。"""
    EXCLUDE_EXT = {
        # 只擋二進位雜訊，其餘一律看得見
        '.pyc', '.pyo', '.so', '.o', '.a', '.dll', '.dylib', '.class',
        '.swp', '.swo', '.bak~', '.pid', '.lock', '.db-journal',
    }
    EXCLUDE_DIRS = {'.git', 'node_modules', '__pycache__', 'venv', '.venv', 'env'}
    file_entries = []
    for root, dirs, files in os.walk(project_path):
        dirs[:] = sorted(d for d in dirs if d not in EXCLUDE_DIRS)
        for fname in sorted(files):
            ext = os.path.splitext(fname)[1]
            if ext.lower() in EXCLUDE_EXT:
                continue
            fpath = os.path.join(root, fname)
            try:
                mtime = os.path.getmtime(fpath)
            except OSError:
                continue
            rel = os.path.relpath(fpath, project_path)
            file_entries.append((rel, mtime))
    hash_key = hashlib.md5(
        json.dumps(file_entries, ensure_ascii=False).encode()
    ).hexdigest()
    if project_path in _map_cache and _map_cache[project_path][0] == hash_key:
        return _map_cache[project_path][1]
    recent_threshold = time.time() - 7 * 24 * 3600
    lines = []
    for rel, mtime in file_entries:
        fpath = os.path.join(project_path, rel)
        ext = os.path.splitext(rel)[1]
        is_recent = mtime >= recent_threshold
        if is_recent:
            if ext == '.py':
                desc = _get_or_create_py_desc(fpath)
            elif ext == '.md':
                desc = '（無說明）'
                try:
                    with open(fpath, encoding='utf-8', errors='replace') as f:
                        for line in f:
                            s = line.strip()
                            if s:
                                desc = (s.lstrip('#').strip())[:80]
                                break
                except OSError:
                    desc = '（無法讀取）'
            else:
                desc = '（無說明）'
            lines.append(f'★ {rel:<43} → {desc}')
        else:
            lines.append(f'  {rel}')
    if not lines:
        map_text = '（目錄為空或無符合的檔案）'
    else:
        header = f'根目錄：{project_path}'
        if project_name:
            header = f'專案：{project_name}　{header}'
        map_text = header + '\n\n' + '\n'.join(lines)
    _map_cache[project_path] = (hash_key, map_text)
    return map_text


_MERGED_TOOLS_MARK = "2026-07-19"


# === ARCUS 併入工具 · 上網搜尋與網頁正文抽取 (2026-07-19) ===
# 原 tools/read/deep_research.py，併入後不再跨專案借用。

#!/usr/bin/env python3
"""
deep_research.py — Kitchen 研究型任務的確定性搜尋腳本
取代 deep-research 技能，把搜尋行為的控制權從 AI 自主判斷移到 Python 確定性邏輯。

設計分工：
  Python 負責：預算控制、矩陣缺口計算、搜尋排程、內容清洗
  Claude 負責：從乾淨的 snippet 做最終語意辨識與格式化

演算法：10 步動態矩陣搜尋
  Step 1-2：Query Decomposition + Budget Pool
  Step 3-4：Stage 1 鷹眼掃描 + 缺口計算
  Step 5-6：邊際收益排序 + 精準定向搜尋
  Step 7-8：延遲加載 + 深度下鑽（fetch + prune）
  Step 9：Context Pruning（BeautifulSoup，每頁 1,500 字）
  Step 10：輸出結構化 JSON，由 Claude 做最終格式化

用法：
  python deep_research.py "提示詞" [--budget 25] [--depth-threshold 0.8] [--max-chars 1500]
"""

import re
import sys
import json
import argparse
import random
import time
# 延遲載入：requests/bs4 約佔模組載入 200ms，會拖慢 --mcp 冷啟，
# 使工具來不及在第一輪登記；改為第一次真正抓網頁時才載入。
requests = None
BeautifulSoup = None
_web_libs_loaded = False
def _ensure_web_libs():
    global requests, BeautifulSoup, _web_libs_loaded
    if _web_libs_loaded:
        return
    _web_libs_loaded = True
    try:
        import requests as _rq
        from bs4 import BeautifulSoup as _bs
        requests = _rq
        BeautifulSoup = _bs
    except Exception as _e:  # 缺套件只讓本節功能失效，不阻斷引擎啟動
        print('[arcus] 上網搜尋套件缺失：%s' % _e)

# ─────────────────────────────────────────────────────────────────────────────
# 靜態資料：GDP 前 20 國（2024 年 IMF 排名，含大洲分類）
# ─────────────────────────────────────────────────────────────────────────────
GDP_TOP20 = [
    {"rank": 1,  "country": "美國",         "continent": "北美", "en": "USA"},
    {"rank": 2,  "country": "中國",         "continent": "亞洲", "en": "China"},
    {"rank": 3,  "country": "德國",         "continent": "歐洲", "en": "Germany"},
    {"rank": 4,  "country": "日本",         "continent": "亞洲", "en": "Japan"},
    {"rank": 5,  "country": "印度",         "continent": "亞洲", "en": "India"},
    {"rank": 6,  "country": "英國",         "continent": "歐洲", "en": "UK"},
    {"rank": 7,  "country": "法國",         "continent": "歐洲", "en": "France"},
    {"rank": 8,  "country": "義大利",       "continent": "歐洲", "en": "Italy"},
    {"rank": 9,  "country": "巴西",         "continent": "南美", "en": "Brazil"},
    {"rank": 10, "country": "加拿大",       "continent": "北美", "en": "Canada"},
    {"rank": 11, "country": "俄羅斯",       "continent": "歐洲", "en": "Russia"},
    {"rank": 12, "country": "墨西哥",       "continent": "北美", "en": "Mexico"},
    {"rank": 13, "country": "澳洲",         "continent": "大洋洲", "en": "Australia"},
    {"rank": 14, "country": "韓國",         "continent": "亞洲", "en": "South Korea"},
    {"rank": 15, "country": "荷蘭",         "continent": "歐洲", "en": "Netherlands"},
    {"rank": 16, "country": "西班牙",       "continent": "歐洲", "en": "Spain"},
    {"rank": 17, "country": "土耳其",       "continent": "亞洲", "en": "Turkey"},
    {"rank": 18, "country": "印尼",         "continent": "亞洲", "en": "Indonesia"},
    {"rank": 19, "country": "沙烏地阿拉伯", "continent": "亞洲", "en": "Saudi Arabia"},
    {"rank": 20, "country": "瑞士",         "continent": "歐洲", "en": "Switzerland"},
]
GDP_RANK_MAP = {r["country"]: r for r in GDP_TOP20}
GDP_EN_MAP = {r["country"]: r["en"] for r in GDP_TOP20}

# SearXNG 公開實例（按優先順序嘗試）
SEARXNG_INSTANCES = [
    "https://searx.be",
    "https://paulgo.io",
    "https://search.inetol.net",
    "https://opnxng.com",
    "https://searx.tiekoetter.com",
    "https://searx.work",
]

# ─────────────────────────────────────────────────────────────────────────────
# Step 2：BudgetPool — 硬性預算計數器
# ─────────────────────────────────────────────────────────────────────────────
class BudgetPool:
    """強制搜尋次數上限，在 Python 層機械性執行，Claude 看不到內部狀態。"""

    def __init__(self, b_max: int = 25):
        self.b_max = b_max
        self._used = 0
        self._log = []

    def use(self, n: int = 1, label: str = "") -> bool:
        if self._used + n > self.b_max:
            self._log.append(f"[EXHAUSTED] 嘗試使用 {n} 次，剩餘 {self.remaining()} 次，拒絕。")
            return False
        self._used += n
        self._log.append(f"[OK] +{n}（{label[:50]}），已用 {self._used}/{self.b_max}")
        return True

    def remaining(self) -> int:
        return self.b_max - self._used

    def is_exhausted(self) -> bool:
        return self._used >= self.b_max

    def summary(self) -> dict:
        return {"used": self._used, "max": self.b_max, "remaining": self.remaining(), "log": self._log}


# ─────────────────────────────────────────────────────────────────────────────
# Step 1：Query Decomposition
# ─────────────────────────────────────────────────────────────────────────────
def parse_query(prompt: str) -> dict:
    """從使用者提示詞抽取條件、地區、關鍵字。"""
    conditions = []
    condition_patterns = {
        "paypal": r"paypal|PayPal",
        "foreign_account": r"外國人|外籍|非本國|foreign.?(account|register|signup)",
        "api": r"\bapi\b|API|有\s*API|提供\s*API|免費\s*API",
        "playwright": r"playwright|Playwright|爬蟲|自動化爬|headless",
        "auto_upload": r"自動化上架|自動上架|auto.?upload|automated.?listing",
        "free_key": r"免費\s*(api)?\s*key|free.?api.?key|免費金鑰",
    }
    for key, pattern in condition_patterns.items():
        if re.search(pattern, prompt, re.IGNORECASE):
            conditions.append(key)

    # 解析 GDP 前 N 國
    gdp_n = 20
    gdp_match = re.search(r"GDP.{0,10}前\s*(\d+)|前\s*(\d+).{0,10}(國|GDP)|top\s*(\d+)", prompt, re.IGNORECASE)
    if gdp_match:
        gdp_n = int(next(g for g in gdp_match.groups() if g))

    regions = [r["country"] for r in GDP_TOP20[:gdp_n]]

    # 關鍵字
    keywords = []
    if re.search(r"平台|marketplace|platform", prompt, re.IGNORECASE):
        keywords.append("marketplace")
    if re.search(r"電商|e.?commerce|跨境", prompt, re.IGNORECASE):
        keywords.append("e-commerce")
    if re.search(r"數位產品|digital.?product", prompt, re.IGNORECASE):
        keywords.append("digital product")
    if not keywords:
        keywords = ["platform"]

    return {
        "conditions": conditions,
        "regions": regions,
        "keywords": keywords,
        "gdp_n": gdp_n,
        "is_global": bool(re.search(r"五大洲|全球|worldwide|global", prompt, re.IGNORECASE)),
    }


def build_state_matrix(regions: list) -> dict:
    """
    建立狀態矩陣：{region: {"_search_count": 0, "_results": []}}
    _results 存放搜尋結果的 snippet（乾淨文字），由 Claude 做語意辨識。
    """
    return {
        region: {"_search_count": 0, "_results": []}
        for region in regions
    }


# ─────────────────────────────────────────────────────────────────────────────
# 搜尋工具：SearXNG + DuckDuckGo 備援
# ─────────────────────────────────────────────────────────────────────────────
_searxng_cache: str | None = None  # 快取已確認可用的 SearXNG 實例

def search_searxng(query: str, max_results: int = 6) -> list:
    global _searxng_cache
    _ensure_web_libs()
    instances = ([_searxng_cache] + [i for i in SEARXNG_INSTANCES if i != _searxng_cache]
                 if _searxng_cache else SEARXNG_INSTANCES)
    for base in instances:
        try:
            r = requests.get(
                f"{base}/search",
                params={"q": query, "format": "json", "categories": "general"},
                timeout=6,
                headers={"User-Agent": "Mozilla/5.0 (compatible; kitchen-research/1.0)"},
            )
            if r.status_code == 200:
                data = r.json()
                results = data.get("results", [])[:max_results]
                if results:
                    _searxng_cache = base
                    return [{"title": x.get("title", ""), "url": x.get("url", ""), "snippet": x.get("content", "")} for x in results]
        except Exception:
            continue
    return []


def search_duckduckgo(query: str, max_results: int = 6) -> list:
    for module in ["ddgs", "duckduckgo_search"]:
        try:
            if module == "ddgs":
                from ddgs import DDGS
            else:
                from duckduckgo_search import DDGS
            results = DDGS().text(query, max_results=max_results)
            return [{"title": r.get("title", ""), "url": r.get("href", ""), "snippet": r.get("body", "")} for r in results]
        except (ImportError, Exception):
            continue
    return []


# ── 學術索引搜尋：OpenAlex / Crossref / arXiv / Semantic Scholar（皆免金鑰）──
# 統一回傳 [{"title","url","snippet"}]，與一般搜尋一致；snippet 內含作者·年份·摘要，
# url 優先給 DOI／arXiv abs 連結，供引擎判讀後交 paper_ingest 建卡。
_SCHOLAR_MAILTO = "arcus-research@forest-carbon.duckdns.org"  # OpenAlex/Crossref 禮貌池識別

def _oa_abstract(inv):
    """把 OpenAlex 的 abstract_inverted_index 還原成一段文字；無則回空字串。"""
    if not inv or not isinstance(inv, dict):
        return ""
    try:
        pos = {}
        for word, idxs in inv.items():
            for i in idxs:
                pos[i] = word
        return " ".join(pos[i] for i in sorted(pos))
    except Exception:
        return ""


def _norm_title(t):
    import re
    return re.sub(r"\W+", "", (t or "").lower())[:80]


def search_openalex(query: str, max_results: int = 6) -> list:
    _ensure_web_libs()
    try:
        r = requests.get(
            "https://api.openalex.org/works",
            params={"search": query, "per-page": max_results, "mailto": _SCHOLAR_MAILTO},
            timeout=8, headers={"User-Agent": "arcus-research/1.0 (mailto:%s)" % _SCHOLAR_MAILTO},
        )
        if r.status_code != 200:
            return []
        out = []
        for w in (r.json().get("results", []) or [])[:max_results]:
            title = w.get("display_name") or ""
            if not title:
                continue
            year = w.get("publication_year") or ""
            doi = w.get("doi") or ""
            authors = ", ".join(
                (a.get("author", {}) or {}).get("display_name", "")
                for a in (w.get("authorships", []) or [])[:5] if a.get("author"))
            url = doi or (w.get("id") or "")
            abstract = _oa_abstract(w.get("abstract_inverted_index"))
            snippet = " · ".join(x for x in [authors, str(year), abstract] if x)
            out.append({"title": title, "url": url, "snippet": snippet[:400]})
        return out
    except Exception:
        return []


def search_crossref(query: str, max_results: int = 6) -> list:
    import re
    _ensure_web_libs()
    try:
        r = requests.get(
            "https://api.crossref.org/works",
            params={"query": query, "rows": max_results,
                    "select": "title,author,issued,DOI,abstract", "mailto": _SCHOLAR_MAILTO},
            timeout=8, headers={"User-Agent": "arcus-research/1.0 (mailto:%s)" % _SCHOLAR_MAILTO},
        )
        if r.status_code != 200:
            return []
        out = []
        for it in (r.json().get("message", {}).get("items", []) or [])[:max_results]:
            title = (it.get("title") or [""])[0]
            if not title:
                continue
            doi = it.get("DOI") or ""
            try:
                year = it.get("issued", {}).get("date-parts", [[""]])[0][0] or ""
            except Exception:
                year = ""
            authors = ", ".join(
                ("%s %s" % (a.get("given", ""), a.get("family", ""))).strip()
                for a in (it.get("author", []) or [])[:5])
            abstract = re.sub(r"<[^>]+>", "", it.get("abstract", "") or "")
            snippet = " · ".join(x for x in [authors, str(year), abstract] if x)
            url = ("https://doi.org/%s" % doi) if doi else ""
            out.append({"title": title, "url": url, "snippet": snippet[:400]})
        return out
    except Exception:
        return []


def search_arxiv(query: str, max_results: int = 6) -> list:
    _ensure_web_libs()
    try:
        import urllib.parse as _up
        import xml.etree.ElementTree as _ET
        q = _up.quote(query)
        r = requests.get(
            "http://export.arxiv.org/api/query?search_query=all:%s&start=0&max_results=%d" % (q, max_results),
            timeout=10, headers={"User-Agent": "arcus-research/1.0"},
        )
        if r.status_code != 200:
            return []
        ns = {"a": "http://www.w3.org/2005/Atom"}
        root = _ET.fromstring(r.text)
        out = []
        for e in root.findall("a:entry", ns)[:max_results]:
            title = (e.findtext("a:title", "", ns) or "").strip().replace("\n", " ")
            if not title:
                continue
            url = (e.findtext("a:id", "", ns) or "").strip()
            summary = (e.findtext("a:summary", "", ns) or "").strip().replace("\n", " ")
            year = (e.findtext("a:published", "", ns) or "")[:4]
            authors = ", ".join(
                (a.findtext("a:name", "", ns) or "").strip()
                for a in e.findall("a:author", ns)[:5])
            snippet = " · ".join(x for x in [authors, year, summary] if x)
            out.append({"title": title, "url": url, "snippet": snippet[:400]})
        return out
    except Exception:
        return []


def search_semanticscholar(query: str, max_results: int = 6) -> list:
    _ensure_web_libs()
    try:
        r = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={"query": query, "limit": max_results,
                    "fields": "title,year,authors,abstract,externalIds,url"},
            timeout=8, headers={"User-Agent": "arcus-research/1.0"},
        )
        if r.status_code != 200:
            return []
        out = []
        for p in (r.json().get("data", []) or [])[:max_results]:
            title = p.get("title") or ""
            if not title:
                continue
            year = p.get("year") or ""
            ext = p.get("externalIds") or {}
            doi = ext.get("DOI") or ""
            arx = ext.get("ArXiv") or ""
            url = ("https://doi.org/%s" % doi) if doi else (
                  ("https://arxiv.org/abs/%s" % arx) if arx else (p.get("url") or ""))
            authors = ", ".join((a.get("name", "") for a in (p.get("authors", []) or [])[:5]))
            abstract = p.get("abstract") or ""
            snippet = " · ".join(x for x in [authors, str(year), abstract] if x)
            out.append({"title": title, "url": url, "snippet": snippet[:400]})
        return out
    except Exception:
        return []


def search_scholar(query: str, max_results: int = 6) -> list:
    """依序打四家學術索引（OpenAlex→Crossref→arXiv→Semantic Scholar，皆免金鑰），
    以標題正規化去重，累積到 max_results 為止；任一家失敗略過、不阻斷其餘。"""
    out, seen = [], set()
    for fn in (search_openalex, search_crossref, search_arxiv, search_semanticscholar):
        if len(out) >= max_results:
            break
        try:
            for r in fn(query, max_results):
                k = _norm_title(r.get("title"))
                if not k or k in seen:
                    continue
                seen.add(k)
                out.append(r)
                if len(out) >= max_results:
                    break
        except Exception:
            continue
    return out


def search(query: str, budget_pool: BudgetPool, max_results: int = 6) -> list:
    """消耗一次預算，學術索引優先（OpenAlex/Crossref/arXiv/Semantic Scholar，皆免金鑰），
    回空才退一般網路 DuckDuckGo，再退 SearXNG。每次搜尋後隨機休息 1–3 秒，避免 IP 被封鎖。
    §接上四家學術介面（2026-07-22）：本引擎主要查論文，學術庫命中率與著錄品質遠勝一般搜尋，
    且各家皆免金鑰；非學術查詢在學術庫回空時自動落到一般網路，行為向下相容。"""
    if not budget_pool.use(1, label=query):
        return []
    results = search_scholar(query, max_results)
    if not results:
        results = search_duckduckgo(query, max_results)
    if not results:
        results = search_searxng(query, max_results)
    # 隨機間隔：1.0–3.0 秒，分散請求時間點，降低被封鎖風險
    time.sleep(random.uniform(1.0, 3.0))
    return results


# ─────────────────────────────────────────────────────────────────────────────
# Step 9：Context Pruning
# ─────────────────────────────────────────────────────────────────────────────
def _web_auth_headers(url: str) -> dict:
    """依主機名從 gitignore 憑證檔查授權標頭；查無回空。權杖只存於憑證檔，不進程式碼、不進版控。"""
    try:
        from urllib.parse import urlparse
        import os as _os, json as _json
        host = (urlparse(url).hostname or "").lower()
        if not host:
            return {}
        cp = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), ".web_creds.json")
        if not _os.path.exists(cp):
            return {}
        creds = _json.load(open(cp, encoding="utf-8"))
        for h, hdrs in creds.items():
            hl = str(h).lower()
            if host == hl or host.endswith("." + hl):
                return {"Authorization": hdrs} if isinstance(hdrs, str) else dict(hdrs)
    except Exception:
        pass
    return {}


def fetch_and_prune(url: str, max_chars: int = 1500) -> str:
    """_WEB_JSON_MARK 下載網址並依回應型別處理：HTML 走去雜訊清洗；
    JSON/純文字/XML 等機器格式逐字回傳（超長明講截斷），不做有損轉換。"""
    try:
        _hdrs = {"User-Agent": "Mozilla/5.0"}
        _hdrs.update(_web_auth_headers(url))
        _ensure_web_libs()
        r = requests.get(url, timeout=10, headers=_hdrs)
        ctype = (r.headers.get("Content-Type") or "").lower()
        body = r.text or ""
        sniff = body.lstrip()[:1]
        is_html = ("text/html" in ctype) or ("application/xhtml" in ctype)
        is_structured = (("json" in ctype) or ("text/plain" in ctype)
                         or ("csv" in ctype) or ("xml" in ctype)
                         or (sniff in ("{", "[")))
        if is_structured and not is_html:
            cap = max(max_chars, 8000)
            if len(body) > cap:
                return body[:cap] + ("\n…（已截斷：原始長度 %d 字，僅顯示前 %d 字）" % (len(body), cap))
            return body
        r.encoding = r.apparent_encoding
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside", "form", "noscript"]):
            tag.decompose()
        content = soup.find("main") or soup.find("article") or soup.find("body")
        if not content:
            txt = soup.get_text(separator=" ", strip=True) or body
            txt = re.sub(r"\s{3,}", "  ", txt)
            return txt[:max(max_chars, 4000)]
        text = content.get_text(separator=" ", strip=True)
        text = re.sub(r"\s{3,}", "  ", text)
        return text[:max_chars]
    except Exception:
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# Step 4：Gap Ratio 計算
# ─────────────────────────────────────────────────────────────────────────────
def calc_gap_ratio(matrix: dict) -> dict:
    """已有任何搜尋結果的地區算「已覆蓋」。"""
    total = len(matrix)
    covered = sum(1 for d in matrix.values() if d["_results"])
    gaps = [region for region, data in matrix.items() if not data["_results"]]
    return {
        "covered": covered,
        "total": total,
        "coverage_rate": covered / total if total else 0,
        "gaps": gaps,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Step 5：Marginal Utility Scoring
# ─────────────────────────────────────────────────────────────────────────────
def score_marginal_utility(gaps: list, matrix: dict) -> list:
    """Priority_Score = 1 / (GDP排名/3 × (已搜次數+1))  → 分數越大越優先搜尋"""
    scored = []
    for country in gaps:
        info = GDP_RANK_MAP.get(country, {"rank": 999})
        rank = info["rank"]
        search_count = matrix[country]["_search_count"]
        score = 1 / ((rank / 3) * (search_count + 1))
        scored.append({"country": country, "rank": rank, "search_count": search_count, "score": score})
    return sorted(scored, key=lambda x: x["score"], reverse=True)


# ─────────────────────────────────────────────────────────────────────────────
# Step 3：Stage 1 鷹眼掃描
# ─────────────────────────────────────────────────────────────────────────────
def stage1_sweep(matrix: dict, budget_pool: BudgetPool,
                 keywords: list, conditions: list) -> None:
    """2 次寬泛搜尋，結果廣播到所有未覆蓋的地區（填矩陣初始值）。"""
    cond_en = []
    if "paypal" in conditions: cond_en.append("PayPal")
    if "api" in conditions: cond_en.append("API")
    if "foreign_account" in conditions: cond_en.append("foreign registration")
    cond_str = " ".join(cond_en)

    sweep_queries = [
        f"top global {keywords[0]} {cond_str} platforms list 2024",
        f"best international {keywords[0]} {cond_str} cross-border",
    ]

    for query in sweep_queries:
        if budget_pool.is_exhausted():
            break
        results = search(query, budget_pool)
        if results:
            # 把全球性搜尋結果廣播到所有尚無資料的地區
            for region in matrix:
                if not matrix[region]["_results"]:
                    matrix[region]["_results"].append({
                        "query": query,
                        "scope": "global",
                        "items": results,
                    })
                    break  # 每份結果只填一個空白地區


# ─────────────────────────────────────────────────────────────────────────────
# Step 6：Targeted Retrieval
# ─────────────────────────────────────────────────────────────────────────────
def targeted_search(top_gaps: list, matrix: dict, budget_pool: BudgetPool,
                    keywords: list, conditions: list) -> None:
    """針對 Priority_Score 最高的前 3 個缺口做精準搜尋。"""
    cond_str = " ".join([
        "PayPal" if "paypal" in conditions else "",
        "API" if "api" in conditions else "",
        "foreign seller" if "foreign_account" in conditions else "",
    ]).strip()

    for gap_info in top_gaps[:3]:
        if budget_pool.is_exhausted():
            break
        country = gap_info["country"]
        country_en = GDP_EN_MAP.get(country, country)
        query = f"{country_en} {keywords[0]} {cond_str} platform 2024"

        results = search(query, budget_pool)
        matrix[country]["_search_count"] += 1

        if results:
            matrix[country]["_results"].append({
                "query": query,
                "scope": country,
                "items": results,
            })
        else:
            # 搜尋成功但無結果，標記為已嘗試
            matrix[country]["_results"].append({
                "query": query,
                "scope": country,
                "items": [],
                "note": "搜尋無結果",
            })


# ─────────────────────────────────────────────────────────────────────────────
# Step 7-8：延遲加載 + 深度下鑽（fetch + prune）
# ─────────────────────────────────────────────────────────────────────────────
def deep_drill(matrix: dict, budget_pool: BudgetPool,
               conditions: list, keywords: list,
               depth_max_chars: int = 3000, top_n: int = 5) -> list:
    """
    廣度達標後，對最常出現的 URL 做深度 fetch + prune。
    Step 7（延遲加載）：查 Playwright、免費 API key 等細節。
    """
    # 收集所有出現過的 URL（依出現頻率排序）
    url_count = {}
    for data in matrix.values():
        for search_result in data["_results"]:
            for item in search_result.get("items", []):
                url = item.get("url", "")
                if url:
                    url_count[url] = url_count.get(url, 0) + 1

    top_urls = sorted(url_count.items(), key=lambda x: x[1], reverse=True)[:top_n]

    # 深度 fetch：查 Playwright + 免費 API key
    detail_results = []
    for url, count in top_urls:
        if budget_pool.is_exhausted():
            break
        # 深度下鑽搜尋（針對這個 URL 對應的平台）
        title = next(
            (item.get("title", "") for data in matrix.values()
             for sr in data["_results"] for item in sr.get("items", []) if item.get("url") == url),
            url
        )
        detail_query = f'site:{_extract_domain(url)} API playwright automation free key'
        detail_results_raw = search(detail_query, budget_pool, max_results=3)

        # Fetch 頁面本身
        page_text = fetch_and_prune(url, max_chars=depth_max_chars)

        combined_text = page_text + " " + " ".join(
            r.get("snippet", "") for r in detail_results_raw
        )
        detail_results.append({
            "url": url,
            "title": title,
            "appear_count": count,
            "page_excerpt": page_text[:500] if page_text else "",
            "playwright_signal": bool(re.search(r"playwright|selenium|puppeteer|headless|automation", combined_text, re.I)),
            "free_api_signal": bool(re.search(r"free.{0,20}(api|key|tier)|api.{0,10}free|免費.{0,10}(api|key)", combined_text, re.I)),
        })

    return detail_results


def _extract_domain(url: str) -> str:
    m = re.search(r"https?://([^/]+)", url)
    return m.group(1) if m else url[:40]


# ─────────────────────────────────────────────────────────────────────────────
# 通用 Pipeline 工具函式
# ─────────────────────────────────────────────────────────────────────────────

# 過濾 snippet 中無意義的英文詞，讓實體提取更乾淨
COMMON_EN_STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "as", "is", "are", "was", "were", "be",
    "been", "being", "have", "has", "had", "do", "does", "did", "will",
    "would", "could", "should", "may", "might", "shall", "can", "need",
    "this", "that", "these", "those", "it", "its", "they", "their", "we",
    "our", "you", "your", "he", "she", "his", "her", "who", "which",
    "what", "how", "when", "where", "why", "if", "then", "than", "more",
    "most", "not", "no", "so", "up", "out", "about", "into", "through",
    "also", "just", "only", "all", "any", "both", "few", "many", "such",
    "new", "use", "get", "make", "using", "based", "best", "top", "free",
    "list", "support", "compare", "comparison", "platform", "open",
    "tool", "tools", "product", "products", "service", "services",
    "cloud", "web", "site", "page", "com", "www", "here", "there",
    "the", "an", "some", "you", "your", "we", "our", "my", "their",
    "vs", "and", "but", "or", "for", "nor", "yet", "so",
}
# 過濾假冒實體的常見首字大寫詞
_STOPWORDS_TITLE = {
    "The", "This", "That", "These", "Those", "There", "Here",
    "When", "Where", "What", "How", "Why", "And", "But", "Or",
    "For", "With", "From", "Into", "About", "After", "Before",
    "If", "Then", "While", "As", "At", "In", "On", "To", "Of",
    "It", "Its", "We", "Our", "You", "Your", "He", "She", "They",
    "Just", "Also", "Not", "All", "Both", "Each", "Many", "Such",
    "New", "Use", "Get", "Make", "Using", "Based",
}


def extract_entities_from_results(results: list, max_entities: int = 15) -> list:
    """
    從廣度搜尋結果的 snippet 與 title 中，抽取高頻 Title Case 詞彙，
    作為通用矩陣的「行」（實體名稱：工具、產品、框架等）。

    策略：
      1. 掃描 Title Case 英文詞（首字母大寫）
      2. 過濾停用詞和常見動詞
      3. 依出現頻率排序，保留 >= 2 次的實體
    """
    freq: dict = {}
    for item in results:
        text = (item.get("title", "") + " " + item.get("snippet", ""))
        # 抓 Title Case 英文詞（1–4 個詞的連續組合）
        words = re.findall(r'\b[A-Z][a-zA-Z]{1,}(?:\.[a-zA-Z]+)?(?:\s[A-Z][a-zA-Z]+){0,3}\b', text)
        for w in words:
            if w in _STOPWORDS_TITLE:
                continue
            w_lower = w.lower()
            if w_lower in COMMON_EN_STOPWORDS or len(w) < 3:
                continue
            freq[w] = freq.get(w, 0) + 1

    # 頻率 >= 2 才保留，按頻率排序
    sorted_entities = sorted(
        [(name, cnt) for name, cnt in freq.items() if cnt >= 2],
        key=lambda x: x[1], reverse=True
    )
    return [name for name, _ in sorted_entities[:max_entities]]


def score_entity_frequency(gaps: list, matrix: dict, freq_map: dict) -> list:
    """
    通用 Priority_Score：appearance_freq / (search_count + 1)
    以搜尋結果出現頻率取代地理版的 GDP 排名。
    回傳格式與 score_marginal_utility 相容（country 欄位存實體名稱）。
    """
    scored = []
    for entity in gaps:
        freq = freq_map.get(entity, 1)
        search_count = matrix[entity]["_search_count"]
        score = freq / (search_count + 1)
        scored.append({
            "country": entity,  # 欄位名沿用 geographic 版本（供 targeted_search 相容）
            "freq": freq,
            "search_count": search_count,
            "score": score,
        })
    return sorted(scored, key=lambda x: x["score"], reverse=True)


# ─────────────────────────────────────────────────────────────────────────────
# Step 10：結構化輸出 JSON（通用版，不依賴 GDP_RANK_MAP）
# ─────────────────────────────────────────────────────────────────────────────
def build_output_json_generic(matrix: dict, budget_pool: BudgetPool,
                              conditions: list, keywords: list,
                              detail_results: list, entities: list) -> dict:
    """
    通用版輸出 JSON，以動態抽取的實體取代固定地區清單。
    Coverage = 至少一個條件為 True（或有結果）的實體 / 全部實體。
    JSON 結構與地理版一致，方便 claude --print 做最終格式化。
    """
    # 計算覆蓋率：有任何搜尋結果的實體算覆蓋
    covered_entities = [e for e in entities if matrix.get(e, {}).get("_results")]
    coverage_rate = len(covered_entities) / len(entities) if entities else 0

    # 蒐集所有不重複的搜尋結果條目
    all_items = []
    seen_urls: set = set()
    for entity, data in matrix.items():
        for search_result in data.get("_results", []):
            for item in search_result.get("items", []):
                url = item.get("url", "")
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    all_items.append({
                        "title": item.get("title", ""),
                        "url": url,
                        "snippet": item.get("snippet", ""),
                        "search_scope": search_result.get("scope", entity),
                        "entity": entity,
                    })

    # 深度細節查找表
    detail_map = {d["url"]: d for d in detail_results}
    for item in all_items:
        detail = detail_map.get(item["url"])
        if detail:
            item["playwright_signal"] = detail.get("playwright_signal", "待補充")
            item["free_api_signal"] = detail.get("free_api_signal", "待補充")
            item["page_excerpt"] = detail.get("page_excerpt", "")
        else:
            item["playwright_signal"] = "待補充"
            item["free_api_signal"] = "待補充"

    uncovered = [e for e in entities if e not in covered_entities]

    return {
        "meta": {
            "pipeline": "generic",
            "conditions": conditions,
            "keywords": keywords,
            "entities": entities,
            "instruction_for_claude": (
                "以下是 Python 腳本蒐集的搜尋結果。請根據 title + snippet + page_excerpt 辨識具體的工具／產品名稱，"
                "整理每個實體支援的條件（" + "、".join(conditions) + "），"
                "輸出結構化表格。未在 snippet 中確認的欄位標注「待確認」。"
            ),
        },
        "search_results": all_items,
        "deep_drill": detail_results,
        "coverage": {
            "total_entities": len(entities),
            "covered_entities": len(covered_entities),
            "coverage_rate": round(coverage_rate, 2),
            "uncovered_entities": uncovered,
        },
        "budget": {
            "used": budget_pool.summary()["used"],
            "max": budget_pool.summary()["max"],
            "log": budget_pool.summary()["log"],
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# Step 10：結構化輸出 JSON（地理版）
# ─────────────────────────────────────────────────────────────────────────────
def build_output_json(matrix: dict, budget_pool: BudgetPool,
                      conditions: list, keywords: list,
                      detail_results: list, query_info: dict) -> dict:
    """
    整合所有資料，輸出 Claude 可直接讀取的標準 JSON。
    包含：原始搜尋 snippet（供 Claude 語意辨識）+ 深度 fetch 摘要。
    """
    gap_info = calc_gap_ratio(matrix)

    # 蒐集所有不重複的搜尋結果條目
    all_items = []
    seen_urls = set()
    for region, data in matrix.items():
        continent = GDP_RANK_MAP.get(region, {}).get("continent", "未知")
        for search_result in data["_results"]:
            for item in search_result.get("items", []):
                url = item.get("url", "")
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    all_items.append({
                        "title": item.get("title", ""),
                        "url": url,
                        "snippet": item.get("snippet", ""),
                        "search_scope": search_result.get("scope", region),
                        "region": region,
                        "continent": continent,
                    })

    # 深度細節查找表
    detail_map = {d["url"]: d for d in detail_results}

    # 為每個條目附加深度細節（若有）
    for item in all_items:
        detail = detail_map.get(item["url"])
        if detail:
            item["playwright_signal"] = detail.get("playwright_signal", "待補充")
            item["free_api_signal"] = detail.get("free_api_signal", "待補充")
            item["page_excerpt"] = detail.get("page_excerpt", "")
        else:
            item["playwright_signal"] = "待補充"
            item["free_api_signal"] = "待補充"

    uncovered = [r for r, d in matrix.items() if not d["_results"]]

    return {
        "meta": {
            "conditions": conditions,
            "keywords": keywords,
            "gdp_regions": [r["country"] for r in GDP_TOP20[:query_info["gdp_n"]]],
            "instruction_for_claude": (
                "以下是 Python 腳本蒐集的搜尋結果。請根據 title + snippet + page_excerpt 辨識具體的平台名稱，"
                "整理每個平台支援的條件（" + "、".join(conditions) + "），"
                "輸出結構化表格。未在 snippet 中確認的欄位標注「待確認」。"
            ),
        },
        "search_results": all_items,
        "deep_drill": detail_results,
        "coverage": {
            "total_regions": gap_info["total"],
            "covered_regions": gap_info["covered"],
            "coverage_rate": round(gap_info["coverage_rate"], 2),
            "uncovered_regions": uncovered,
        },
        "budget": {
            "used": budget_pool.summary()["used"],
            "max": budget_pool.summary()["max"],
            "log": budget_pool.summary()["log"],
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# Pipeline 封裝：地理版 vs 通用版
# ─────────────────────────────────────────────────────────────────────────────

def run_geographic_pipeline(args, log) -> dict:
    """
    地理版 pipeline（原 _dr_main() 邏輯，原封不動封裝成函式）。
    以 GDP 前 N 國為矩陣行，GDP 排名為 Priority_Score 依據。
    """
    from collections import Counter  # noqa（local import 防呼叫端未引入）

    # Step 1
    log("Step 1：解析提示詞...")
    query_info = parse_query(args.query)
    conditions = query_info["conditions"] or ["api"]
    regions = query_info["regions"]
    keywords = query_info["keywords"]
    log(f"  條件：{conditions}")
    log(f"  地區：{len(regions)} 個（{', '.join(regions[:4])}...）")
    log(f"  關鍵字：{keywords}")

    # Step 2
    log("Step 2：建立狀態矩陣與預算池...")
    matrix = build_state_matrix(regions)
    budget_pool = BudgetPool(b_max=args.budget)
    log(f"  矩陣維度：{len(regions)} 地區，預算上限：{args.budget} 次")

    # Step 3
    log("Step 3：Stage 1 鷹眼掃描（2 次寬泛搜尋）...")
    stage1_sweep(matrix, budget_pool, keywords, conditions)
    gap_info = calc_gap_ratio(matrix)
    log(f"  覆蓋率：{gap_info['covered']}/{gap_info['total']} ({gap_info['coverage_rate']:.0%})，剩餘預算：{budget_pool.remaining()} 次")

    # Step 4-6 主迴圈
    iteration = 0
    while gap_info["coverage_rate"] < args.depth_threshold and not budget_pool.is_exhausted():
        iteration += 1
        log(f"Step 4-6 迴圈 #{iteration}：覆蓋 {gap_info['coverage_rate']:.0%}，缺口 {len(gap_info['gaps'])} 個...")
        scored_gaps = score_marginal_utility(gap_info["gaps"], matrix)
        log(f"  高優先缺口：{[g['country'] for g in scored_gaps[:3]]}")
        targeted_search(scored_gaps, matrix, budget_pool, keywords, conditions)
        gap_info = calc_gap_ratio(matrix)
        log(f"  → 覆蓋 {gap_info['coverage_rate']:.0%}，剩餘預算 {budget_pool.remaining()} 次")

    log(f"廣度完成：覆蓋 {gap_info['coverage_rate']:.0%}（門檻 {args.depth_threshold:.0%}），已用 {budget_pool.summary()['used']}/{args.budget} 次")

    # Step 7-8
    detail_results = []
    if not args.skip_drill and not budget_pool.is_exhausted():
        log("Step 7-8：深度下鑽（前 5 個高頻 URL）...")
        detail_results = deep_drill(matrix, budget_pool, conditions, keywords,
                                    depth_max_chars=args.max_chars * 2)
        log(f"  深度下鑽完成：{len(detail_results)} 個 URL")

    # Step 10
    log("Step 10：建立輸出 JSON...")
    output = build_output_json(matrix, budget_pool, conditions, keywords, detail_results, query_info)
    log(f"  蒐集結果條目：{len(output['search_results'])} 筆")
    log(f"  最終預算使用：{output['budget']['used']}/{output['budget']['max']} 次")

    return output


def run_generic_pipeline(args, log) -> dict:
    """
    通用版 pipeline（雙軌設計選 B）。
    Stage 1 廣度搜尋抽取實體，Stage 2 逐一精準搜尋填充矩陣，Stage 3 deep_drill。
    """
    from collections import Counter

    budget_pool = BudgetPool(b_max=args.budget)

    # 從 query 推導 conditions 與 keywords（重用現有 parse_query）
    query_info = parse_query(args.query)
    conditions = query_info["conditions"] or ["free_tier"]
    keywords = query_info["keywords"]
    log(f"[Generic] 條件：{conditions}，關鍵字：{keywords}")
    log(f"[Generic] 預算上限：{args.budget} 次")

    # ── Stage 1：廣度搜尋，從結果抽取實體 ──────────────────────────────────────
    log("Stage 1：廣度搜尋，抽取實體...")
    broad_results = search(args.query, budget_pool, max_results=10)
    if not broad_results:
        log("  廣度搜尋無結果，嘗試備用查詢...")
        fallback_q = f"best {keywords[0]} tools comparison 2024"
        broad_results = search(fallback_q, budget_pool, max_results=10)

    entities = extract_entities_from_results(broad_results, max_entities=15)
    log(f"  抽取到 {len(entities)} 個實體：{entities[:5]}...")

    if not entities:
        return {"error": "無法從搜尋結果抽取實體，請嘗試更具體的查詢。", "budget": budget_pool.summary()}

    # 建立 freq_map：每個結果的 snippet + title 中出現實體的次數
    freq_map: dict = Counter()
    for item in broad_results:
        text = (item.get("title", "") + " " + item.get("snippet", "")).lower()
        for entity in entities:
            if entity.lower() in text:
                freq_map[entity] += 1

    # ── Stage 2：逐一精準搜尋，填充矩陣 ────────────────────────────────────────
    log("Stage 2：逐一精準搜尋...")
    matrix: dict = {entity: {"_search_count": 0, "_results": []} for entity in entities}

    # 先把廣度結果廣播到尚未有資料的實體（按 freq 高低依序填充）
    sorted_entities_by_freq = sorted(entities, key=lambda e: freq_map.get(e, 0), reverse=True)
    for entity in sorted_entities_by_freq:
        if not matrix[entity]["_results"] and broad_results:
            matrix[entity]["_results"].append({
                "query": args.query,
                "scope": "broad",
                "items": broad_results,
            })
            break  # 只把廣度結果填給頻率最高的那個

    # 計算缺口並精準搜尋
    gaps_all = [e for e in entities if not matrix[e]["_results"]]
    scored_gaps = score_entity_frequency(gaps_all, matrix, freq_map)

    for gap_info in scored_gaps:
        if budget_pool.is_exhausted():
            break
        entity = gap_info["country"]  # 欄位名沿用 geographic 版本
        entity_query = f"{entity} {args.query}"
        results = search(entity_query, budget_pool, max_results=6)
        matrix[entity]["_search_count"] += 1

        if results:
            matrix[entity]["_results"].append({
                "query": entity_query,
                "scope": entity,
                "items": results,
            })
            log(f"  [{entity}] 找到 {len(results)} 筆結果")
        else:
            matrix[entity]["_results"].append({
                "query": entity_query,
                "scope": entity,
                "items": [],
                "note": "搜尋無結果",
            })
            log(f"  [{entity}] 無結果")

    covered_count = sum(1 for e in entities if matrix[e]["_results"] and
                        any(len(sr.get("items", [])) > 0 for sr in matrix[e]["_results"]))
    log(f"  精準搜尋完成：覆蓋 {covered_count}/{len(entities)} 個實體，剩餘預算 {budget_pool.remaining()} 次")

    # ── Stage 3：deep_drill（budget 有剩時執行）───────────────────────────────
    detail_results: list = []
    if not args.skip_drill and not budget_pool.is_exhausted():
        log("Stage 3：深度下鑽...")
        detail_results = deep_drill(matrix, budget_pool, conditions, keywords,
                                    depth_max_chars=args.max_chars * 2)
        log(f"  深度下鑽完成：{len(detail_results)} 個 URL")

    log("建立通用版輸出 JSON...")
    return build_output_json_generic(matrix, budget_pool, conditions, keywords, detail_results, entities)


# ─────────────────────────────────────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────────────────────────────────────
def _dr_main():
    parser = argparse.ArgumentParser(description="Kitchen 確定性研究腳本（10 步動態矩陣搜尋）")
    parser.add_argument("--query", required=True, help="使用者提示詞（必填）")
    parser.add_argument("--type", choices=["geographic", "generic"], default="generic",
                        help="Pipeline 類型：generic（通用實體矩陣，安全預設，適用任何主題）"
                             "或 geographic（GDP 前 20 國固定矩陣，僅在明確要做各國市場逐一掃描時 opt-in；"
                             "誤用於非市場查詢會硬拆成 20 國、耗時數百秒）")
    parser.add_argument("--budget", type=int, default=25, help="最大搜尋次數（預設 25）")
    parser.add_argument("--depth-threshold", type=float, default=0.8, help="廣度覆蓋門檻（預設 0.8）")
    parser.add_argument("--max-chars", type=int, default=1500, help="每頁截取字數（預設 1500）")
    parser.add_argument("--skip-drill", action="store_true", help="跳過深度下鑽（快速模式）")
    parser.add_argument("--quiet", action="store_true", help="只輸出 JSON，不顯示進度")
    args = parser.parse_args()

    def log(msg: str):
        if not args.quiet:
            print(f"[DR] {msg}", file=sys.stderr)

    if args.type == "geographic":
        result = run_geographic_pipeline(args, log)
    else:
        result = run_generic_pipeline(args, log)

    print(json.dumps(result, ensure_ascii=False, indent=2))


# === ARCUS 併入工具 · 向量圖繪製與文件嵌入 (2026-07-19) ===
# 原 tools/write/draw_docx_figure.py，命令列進入點改由本檔的 --spec 分支呼叫。

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""draw_docx_figure.py — 由精簡 JSON 形狀規格產生 Word 原生向量圖（省 token）

省 token 原理
------------
模型（Claude）只需寫一份很短的 JSON 形狀規格（方框、線、多邊形、橢圓、文字的座標
與顏色）。冗長的向量描述（SVG）與點陣後援（PNG）全部由本工具在硬碟上產生、直接
嵌進 .docx，永遠不進對話脈絡。這樣「描述一張圖」只花少量 token，圖本身的體積不會
被後續每個工具回合乘上去。

輸出格式
--------
用 matplotlib 對同一張圖同時輸出：
  1. 向量 SVG（svg.fonttype=path，中文字轉成向量路徑，永不缺字型）
  2. 點陣 PNG（舊版 Word / 檢視器的後援）
再以 Word 2016+ 原生的 svgBlip 擴充（ext uri {96DAC541-7B7A-43D3-8B79-37D633B846F1}）
把 SVG 當「真向量圖」嵌入、PNG 當後援。支援 SVG 的 Word 顯示無限縮放不失真的向量；
不支援的舊版自動退回顯示 PNG，絕不會破圖。

零安裝：只用 python-docx / lxml / matplotlib / PIL，皆已安裝，不需 inkscape、
libreoffice、svglib、cairosvg。

JSON 規格
---------
{
  "size": [寬px, 高px],          # 畫布尺寸，座標以左上角為原點、y 向下（螢幕直覺）
  "bg": "#ffffff" 或 "none",     # 背景色，可省略（預設白）
  "elements": [
    {"type":"rect",    "x":,"y":,"w":,"h":, "fill":"#fff","stroke":"#333","lw":1.5,"rx":8},
    {"type":"line",    "x1":,"y1":,"x2":,"y2":, "stroke":"#333","lw":1.2,"dash":[6,4],"arrow":false},
    {"type":"arc_arrow","x1":,"y1":,"x2":,"y2":,"rad":0.3,"stroke":"#333","lw":1.2},  # 真圓弧箭頭，rad正負決定鼓向
    {"type":"connect", "from":"節點id","to":"節點id","gap":18,"spread":26,"rad":-0.18,"stroke":"#555","lw":2.6},
        # 高階連接器：指名兩節點(ellipse/rect須有"id")，自動把端點放在同一橢圓環上算箭頭。
        # 全圖的 connect 一起圍成一個環（如四步驟迴圈）；gap=離節點淨空、spread=端點離節點角度、rad 微鼓(負=外鼓)。
    {"type":"polygon", "points":[[x,y],...], "fill":"none","stroke":"#333","lw":1.5,"close":true},
    {"type":"ellipse", "cx":,"cy":,"rx":,"ry":, "fill":"#fff","stroke":"#333","lw":1.5},
    {"type":"text",    "x":,"y":,"s":"文字","size":14,"color":"#111",
                       "ha":"left|center|right","va":"top|center|bottom","bold":false}
  ]
}
fill/stroke 可為 "none"。顏色接受 #rrggbb 或 matplotlib 具名色。

用法
----
  # 由 JSON 檔產生新 .docx（含該圖）
  python3 draw_docx_figure.py --spec fig.json --out /tmp/figure.docx

  # 直接用行內 JSON
  python3 draw_docx_figure.py --spec-json '{"size":[400,300],"elements":[...]}' --out out.docx

  # 從標準輸入讀 JSON
  cat fig.json | python3 draw_docx_figure.py --out out.docx

  # 把圖附加到既有 .docx 末尾（而非新建）
  python3 draw_docx_figure.py --spec fig.json --append-to /path/thesis.docx

  # 內建冰山範例（重現畢業專題 圖1-1 那類結構示意圖）作自我測試
  python3 draw_docx_figure.py --demo iceberg --out /tmp/iceberg.docx

  # 畫圖後自檢（不需 --out，不產圖）：查缺字＋版面出界/重疊，只回純文字報告
  python3 draw_docx_figure.py --spec fig.json --verify
  #   通過 → exit 0；不通過 → exit 3（可讓腳本或機械層據以攔阻）
  #   目的：取代「把 PNG 讀回對話脈絡目視」，PNG 永不進脈絡、省 token

  常用選項：--caption "圖1-1 ..." 加圖說；--width-cm 15 控制插入寬度；
           --keep-assets 保留產生的 .svg/.png（預設會印出其路徑）。
"""
import argparse
import io
import json
import math
import os
import sys

# ── drawing-spec-v1：機械層預設單一真相源（DEFAULTS）────────────────────────
# 慣例優於設定：所有「該全站統一」的規格集中在此。JSON 未寫該欄＝吃預設；寫了＝覆寫。
# 改一處即全專案一致，避免魔術數字散落各行造成「跨輪決策/記憶負擔」（每輪全新呼叫必飄）。
DEFAULTS = {
    # 畫布（16:9＝本專案 PPT 比例，規格可直接拿去做簡報；框緣＝節點外框）
    "canvas": [1600, 900],   # 邏輯座標；內部節點比例不限制
    "png_dpi": 150,          # PNG 後援解析度（一般解析度，使用者指定）；SVG 向量本就無限清晰
    "width_cm": 15.0,        # 嵌入 Word 寬度
    # 文字
    "text_color": "#000000",                                    # 純黑
    "size_base": 22,                                            # 內文基準 pt（agent 可改大小）
    "size_ratio": {"title": 1.6, "subtitle": 1.3, "body": 1.0, "note": 0.75},  # 四級固定比例
    "line_height": 1.3,      # 多行節點行距倍數
    # 節點
    "border_color": "#333333",
    "lw_node": 1.8,          # 節點框（rect/ellipse/polygon）線寬
    "lw_line": 2.6,          # 連接器箭頭線寬
    "corner_radius": 12,     # 矩形預設圓角
    "text_pad": 14,          # 文字距框邊內距下限（供自檢參考）
    "palette": ["#4C72B0", "#55A868", "#C44E52", "#8172B3", "#CCB974", "#64B5CD"],
    # 箭頭 / 連接
    "gap": 18,               # 埠沿外法線外推距離
    "arrow_head": 20,
    "connect_mode": "face",  # 甲＝face 面向對方（預設）；乙＝spec/元素 ring:true 切環形
    "ring_rad": 0.18,        # ring 模式外鼓曲率「幅度」；0.18＝使用者親眼定案「原本很好」的幅度。只用絕對值：符號由 _resolve_connects 資料驅動量測(算 _arc3_control 真實弧峰離形心遠近)自動朝外，agent 不碰、換版面免疫
}

# 字型檔（drawing-spec-v1）：中文 TW-Kai 標楷體 ＋ 英數 Liberation Serif（Times 度量相容替身，丙案）
# matplotlib 逐字元後備：拉丁/數字命中 Liberation Serif、中文命中 TW-Kai。
FONT_CJK_PATH = "/usr/share/fonts/truetype/cns11643/TW-Kai-98_1.ttf"                 # 教育部全字庫標楷體
FONT_LATIN_PATH = "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf"
FONT_LATIN_BOLD_PATH = "/usr/share/fonts/truetype/liberation/LiberationSerif-Bold.ttf"
FONT_FAMILY_CJK = "TW-Kai"
FONT_FAMILY_LATIN = "Liberation Serif"
# verify 第一層 cmap 缺字檢查沿用下列名稱（union：字在任一字型即不算缺）；TW-Kai 無粗體變體，bold 落回 regular
CJK_FONT_PATH = FONT_CJK_PATH
CJK_FONT_BOLD = FONT_CJK_PATH
SVG_EXT_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
ASVG_NS = "http://schemas.microsoft.com/office/drawing/2016/SVG/main"

# 文字繪製層級：matplotlib 預設 patch=1、line=2、annotate 箭頭/text=3（同層時後畫者在上）。
# 設 5 讓所有文字一律浮到最上層，箭頭/形狀永遠壓不到字（治「箭頭蓋過文字」的可讀性）。
# 只改 zorder、不動任何座標/顏色/內容，故黃金測試特徵指紋不受影響（指紋不記 zorder）。
TEXT_ZORDER = 5


def _err(msg):
    print(f"[錯誤] {msg}", file=sys.stderr)
    sys.exit(2)


# --------------------------------------------------------------------------
# connect 高階連接器：指名兩元素 id，自動算「合成環」的箭頭端點
# --------------------------------------------------------------------------
def _node_bbox(el):
    """節點軸對齊外接矩形 (x0, y0, x1, y1)；支援 ellipse / rect / polygon。"""
    t = el.get("type")
    if t == "ellipse":
        cx, cy = float(el["cx"]), float(el["cy"])
        rx, ry = float(el["rx"]), float(el["ry"])
        return (cx - rx, cy - ry, cx + rx, cy + ry)
    if t == "rect":
        x, y, w, h = float(el["x"]), float(el["y"]), float(el["w"]), float(el["h"])
        return (min(x, x + w), min(y, y + h), max(x, x + w), max(y, y + h))
    if t == "polygon":
        pts = el.get("points") or []
        xs = [float(p[0]) for p in pts]
        ys = [float(p[1]) for p in pts]
        if not xs:
            raise ValueError("connect 的 polygon 節點無 points")
        return (min(xs), min(ys), max(xs), max(ys))
    raise ValueError(f"connect 只能連 ellipse / rect / polygon，不支援 type={t!r}")


def _bbox_center(b):
    return ((b[0] + b[2]) / 2.0, (b[1] + b[3]) / 2.0)


def _port(bbox, side, gap):
    """外接框四方位埠（邊中點）沿外法線外推 gap，回 (px, py)。"""
    x0, y0, x1, y1 = bbox
    cx = (x0 + x1) / 2.0
    cy = (y0 + y1) / 2.0
    if side == "N":
        return (cx, y0 - gap)
    if side == "S":
        return (cx, y1 + gap)
    if side == "W":
        return (x0 - gap, cy)
    return (x1 + gap, cy)          # "E"


def _face_side(from_c, to_c):
    """from 節點朝 to 節點的那一側埠：Δx 主導→E/W；Δy 主導→N/S。"""
    dx = to_c[0] - from_c[0]
    dy = to_c[1] - from_c[1]
    if abs(dx) >= abs(dy):
        return "E" if dx >= 0 else "W"
    return "S" if dy >= 0 else "N"


def _ring_side(node_c, target_c, O):
    """ring 模式挑埠：node 沿「繞形心 O 的切線、朝 target 那一側」的正方位邊中點。

    幾何：半徑向量 r=node-O，其兩條垂直切線 (-ry,rx)/(ry,-rx) 取指向 target 的那條，
    再 snap 到最接近的 N/S/E/W。效果＝相鄰節點各出「一橫邊、一縱邊」接成環
    （例：北節點用東側邊、東節點用北側邊），端點恆在節點邊上、不出框。
    """
    rx, ry = node_c[0] - O[0], node_c[1] - O[1]
    t1 = (-ry, rx)
    t2 = (ry, -rx)
    tvx, tvy = target_c[0] - node_c[0], target_c[1] - node_c[1]
    tx, ty = t1 if (t1[0] * tvx + t1[1] * tvy) >= (t2[0] * tvx + t2[1] * tvy) else t2
    if abs(tx) >= abs(ty):
        return "E" if tx >= 0 else "W"
    return "S" if ty >= 0 else "N"


def _arc3_control(x1, y1, x2, y2, rad):
    """matplotlib 在本工具畫布上「實際 render」的 arc3 二次貝茲控制點（回資料座標）。

    唯一真相函式——凡是要預測弧的實際彎向（符號決策、自檢邊界框、凸度量測）一律呼叫本
    函式，禁止再散落裸公式 mid+rad*(dy,-dx)。

    為何不是裸公式：FancyArrowPatch 是「先把端點轉顯示座標、再套 arc3」，而本工具畫布
    set_ylim(H,0) 讓 y 軸上下倒置（且 set_aspect equal 等比、無 x/y 形變），等於對整條弧
    做一次 y 鏡射。故實際控制點＝資料座標裸公式 mid+rad*(dy,-dx) 的「號相反」side
    ＝ (mid − rad·dy, mid + rad·dx)。已用 matplotlib 內部 ConnectionStyle.Arc3 探針位元
    驗證：REAL(rad=+r) 座標 == 裸公式(rad=−r) 座標，四條 ring 箭頭無一例外（見陷阱#29/#32）。
    """
    mx, my = (x1 + x2) / 2.0, (y1 + y2) / 2.0
    dx, dy = x2 - x1, y2 - y1
    return (mx - rad * dy, my + rad * dx)


def _resolve_connects(spec):
    """把 spec 內所有 connect 解算成箭頭端點，回 {id(el): (Sx, Sy, Ex, Ey, rad)}。

    作法（drawing-spec-v1・外接框四方位埠，取代舊「共用大橢圓環」）：
      每個被連節點只認自己的軸對齊外接矩形，四埠＝四邊中點，端點沿外法線外推 gap。
      端點永遠落在節點自己的邊上、不再飄到與節點無關的大橢圓 → 任何佈局都不出框。
    兩種挑埠規則並存（甲乙皆保留、不抹殺圖樣變化）：
      甲 face（預設）：from/to 各自 snap 到「朝對方」那側埠，rad 預設 0（直線、最通用）。
      乙 ring（元素或 spec 帶 ring:true）：同樣用面向埠，但自動加外鼓曲率成旋轉對稱環；
        曲率正負由「朝形心外側」自動決定（agent 永不碰 rad，根治舊 rad 正負號陷阱）。
    回傳簽章與舊版相同 (Sx,Sy,Ex,Ey,rad)，故繪製分支與自檢 _shape_extent 皆不需改。
    相容欄位 spread 保留但忽略。
    """
    els = spec.get("elements", [])
    conns = [el for el in els if isinstance(el, dict) and el.get("type") == "connect"]
    if not conns:
        return {}
    idmap = {el["id"]: el for el in els if isinstance(el, dict) and "id" in el}
    bboxes = {}
    for el in conns:
        for key in ("from", "to"):
            rid = el.get(key)
            if rid not in idmap:
                raise ValueError(f"connect 的 {key}={rid!r} 找不到對應 id（請在目標元素加 \"id\"）")
            if rid not in bboxes:
                bboxes[rid] = _node_bbox(idmap[rid])
    centers = {r: _bbox_center(b) for r, b in bboxes.items()}
    Ox = sum(c[0] for c in centers.values()) / len(centers)
    Oy = sum(c[1] for c in centers.values()) / len(centers)
    ring_all = bool(spec.get("ring", False))

    out = {}
    for el in conns:
        A, B = el["from"], el["to"]
        gap = float(el.get("gap", DEFAULTS["gap"]))
        cA, cB = centers[A], centers[B]
        ring = bool(el.get("ring", ring_all))
        if ring:
            # 乙 ring：沿繞形心切線挑埠，相鄰節點各出一橫邊一縱邊接成環（治菱形對插）
            sideA = _ring_side(cA, cB, (Ox, Oy))
            sideB = _ring_side(cB, cA, (Ox, Oy))
        else:
            # 甲 face：各自 snap 到面向對方那側埠
            sideA = _face_side(cA, cB)
            sideB = _face_side(cB, cA)
        Sx, Sy = _port(bboxes[A], sideA, gap)
        Ex, Ey = _port(bboxes[B], sideB, gap)
        if ring:
            mag = abs(float(el.get("rad", DEFAULTS["ring_rad"])))
            midx, midy = (Sx + Ex) / 2.0, (Sy + Ey) / 2.0
            d_mid = math.hypot(midx - Ox, midy - Oy)
            # 資料驅動挑號（不靠公式正負記憶、不靠讀圖）：用 _arc3_control 算「matplotlib 實際
            # 會 render 的弧峰」(t=0.5)，量它離形心 O 是否比弦中點遠；遠＝外鼓，取該號、否則取反。
            # 因 _arc3_control 已內含 y 倒置補償，此判斷對任何節點佈局都自動朝外，換版面免疫。
            cxp, cyp = _arc3_control(Sx, Sy, Ex, Ey, mag)          # 試 +mag 的真實控制點
            apex_x = 0.25 * Sx + 0.5 * cxp + 0.25 * Ex             # 二次貝茲 t=0.5 弧峰
            apex_y = 0.25 * Sy + 0.5 * cyp + 0.25 * Ey
            rad = mag if math.hypot(apex_x - Ox, apex_y - Oy) >= d_mid else -mag
        else:
            rad = float(el.get("rad", 0.0))   # 面向埠直線最乾淨
        out[id(el)] = (Sx, Sy, Ex, Ey, rad)
    return out


# --------------------------------------------------------------------------
# 字型（drawing-spec-v1）：註冊三檔並設逐字元後備清單
# --------------------------------------------------------------------------
_FONTS_READY = False


def _setup_fonts():
    """把中英文三個字型檔註冊進 matplotlib font_manager；缺檔即回報、不靜默退回舊字型。"""
    global _FONTS_READY
    if _FONTS_READY:
        return
    from matplotlib import font_manager as fm
    missing = []
    for p in (FONT_LATIN_PATH, FONT_LATIN_BOLD_PATH, FONT_CJK_PATH):
        if os.path.exists(p):
            try:
                fm.fontManager.addfont(p)
            except Exception:
                pass
        else:
            missing.append(p)
    if missing:
        _err("字型檔缺失（drawing-spec-v1 規定不靜默退回）：\n  " + "\n  ".join(missing)
             + "\n請確認：apt install fonts-cns11643-kai fonts-liberation")
    _FONTS_READY = True


def _resolve_size(el):
    """文字字級：顯式 size 優先；否則依 level 套四級固定比例（agent 決定基準 size_base）。"""
    if el.get("size") is not None:
        return float(el["size"])
    lvl = el.get("level", "body")
    ratio = float(DEFAULTS["size_ratio"].get(lvl, 1.0))
    return float(DEFAULTS["size_base"]) * ratio


# --------------------------------------------------------------------------
# 1. 用 matplotlib 把 JSON 規格畫成 SVG + PNG
# --------------------------------------------------------------------------
def _build_figure(spec):
    """建圖核心：把 JSON 規格畫成 matplotlib figure，回傳 (fig, ax, text_artists, W, H)。
    text_artists = [(元素索引, matplotlib artist, 元素 dict)]。
    render_spec 與 verify_spec 共用同一套繪圖邏輯，避免兩份程式漂移。"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib import patches
        from matplotlib.font_manager import FontProperties
    except Exception as e:
        _err(f"matplotlib 無法載入：{e}")

    _setup_fonts()
    plt.rcParams["svg.fonttype"] = "path"  # 文字轉向量路徑，SVG 永不缺字型
    plt.rcParams["font.family"] = [FONT_FAMILY_LATIN, FONT_FAMILY_CJK]  # 逐字後備：拉丁 Serif、中文楷體
    plt.rcParams["axes.unicode_minus"] = False

    # family 清單 → matplotlib 逐字元後備混排（拉丁/數字 Liberation Serif、中文 TW-Kai）
    reg_font = FontProperties(family=[FONT_FAMILY_LATIN, FONT_FAMILY_CJK])
    bold_font = FontProperties(family=[FONT_FAMILY_LATIN, FONT_FAMILY_CJK], weight="bold")

    size = spec.get("size", DEFAULTS["canvas"])
    if not (isinstance(size, (list, tuple)) and len(size) == 2):
        _err("size 必須是 [寬, 高]")
    W, H = float(size[0]), float(size[1])
    if W <= 0 or H <= 0:
        _err("size 寬高必須為正數")

    dpi = 100.0
    fig = plt.figure(figsize=(W / dpi, H / dpi), dpi=dpi)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, W)
    ax.set_ylim(H, 0)          # y 向下：左上角為原點
    ax.set_aspect("equal")
    ax.axis("off")

    bg = spec.get("bg", "#ffffff")
    if bg and bg != "none":
        fig.patch.set_facecolor(bg)
        ax.add_patch(patches.Rectangle((0, 0), W, H, facecolor=bg, edgecolor="none", zorder=-100))
    else:
        fig.patch.set_alpha(0.0)

    def _fill(v):
        return "none" if (v is None or v == "none") else v

    def _dash(el, ls_default="-"):
        d = el.get("dash")
        if not d:
            return ls_default
        return (0, tuple(d))

    def _z(el, default):
        # 圖層層級：JSON 有寫 zorder 就用它、沒寫沿用該原語的 matplotlib 自然預設。
        # matplotlib 排序＝先比 zorder 大小（大者在上）、同 zorder 才看加入順序。
        # 治「資料點 ellipse(1) 被格線 line(2) 壓在下面」：資料點 JSON 加 zorder 即可抬層。
        z = el.get("zorder")
        return default if z is None else float(z)

    try:
        connect_geom = _resolve_connects(spec)
    except Exception as ce:
        _err(f"connect 連接器解析失敗：{ce}")

    text_artists = []
    for i, el in enumerate(spec.get("elements", [])):
        t = el.get("type")
        try:
            if t == "rect":
                x, y, w, h = float(el["x"]), float(el["y"]), float(el["w"]), float(el["h"])
                rx = float(el.get("rx", DEFAULTS["corner_radius"]) or 0)
                rx = max(0.0, min(rx, abs(w) / 2.0, abs(h) / 2.0))   # 防 FancyBboxPatch 負尺寸
                common = dict(facecolor=_fill(el.get("fill", "none")),
                              edgecolor=_fill(el.get("stroke", DEFAULTS["border_color"])),
                              linewidth=float(el.get("lw", DEFAULTS["lw_node"])),
                              linestyle=_dash(el), zorder=_z(el, 1))
                if rx > 0:
                    p = patches.FancyBboxPatch((x + rx, y + rx), w - 2 * rx, h - 2 * rx,
                                               boxstyle=f"round,pad={rx},rounding_size={rx}",
                                               **common)
                else:
                    p = patches.Rectangle((x, y), w, h, **common)
                ax.add_patch(p)
            elif t == "line":
                x1, y1, x2, y2 = float(el["x1"]), float(el["y1"]), float(el["x2"]), float(el["y2"])
                color = _fill(el.get("stroke", "#333333"))
                lw = float(el.get("lw", 1.2))
                ls = _dash(el)
                if el.get("arrow"):
                    ax.annotate("", xy=(x2, y2), xytext=(x1, y1), zorder=_z(el, 3),
                                arrowprops=dict(arrowstyle="->", color=color, lw=lw, linestyle=ls))
                else:
                    ax.plot([x1, x2], [y1, y2], color=color, linewidth=lw, linestyle=ls,
                            solid_capstyle="round", zorder=_z(el, 2))
            elif t == "polygon":
                pts = el.get("points") or []
                if len(pts) < 2:
                    continue
                p = patches.Polygon(pts, closed=bool(el.get("close", True)),
                                    facecolor=_fill(el.get("fill", "none")),
                                    edgecolor=_fill(el.get("stroke", DEFAULTS["border_color"])),
                                    linewidth=float(el.get("lw", DEFAULTS["lw_node"])),
                                    linestyle=_dash(el), joinstyle="round", zorder=_z(el, 1))
                ax.add_patch(p)
            elif t == "ellipse":
                cx, cy = float(el["cx"]), float(el["cy"])
                erx, ery = float(el["rx"]), float(el["ry"])
                p = patches.Ellipse((cx, cy), 2 * erx, 2 * ery,
                                    facecolor=_fill(el.get("fill", "none")),
                                    edgecolor=_fill(el.get("stroke", DEFAULTS["border_color"])),
                                    linewidth=float(el.get("lw", DEFAULTS["lw_node"])),
                                    linestyle=_dash(el), zorder=_z(el, 1))
                ax.add_patch(p)
            elif t == "arc_arrow":
                x1, y1, x2, y2 = float(el["x1"]), float(el["y1"]), float(el["x2"]), float(el["y2"])
                rad = float(el.get("rad", 0.3))
                color = _fill(el.get("stroke", "#333333"))
                lw = float(el.get("lw", 1.2))
                ls = _dash(el)
                p = patches.FancyArrowPatch((x1, y1), (x2, y2),
                                            connectionstyle=f"arc3,rad={rad}", arrowstyle="->",
                                            mutation_scale=15, color=color, lw=lw, linestyle=ls,
                                            shrinkA=0, shrinkB=0, zorder=_z(el, 1))
                ax.add_patch(p)
            elif t == "connect":
                Sx, Sy, Ex, Ey, rad = connect_geom[id(el)]
                color = _fill(el.get("stroke", DEFAULTS["border_color"]))
                lw = float(el.get("lw", DEFAULTS["lw_line"]))
                p = patches.FancyArrowPatch((Sx, Sy), (Ex, Ey),
                                            connectionstyle=f"arc3,rad={rad}", arrowstyle="-|>",
                                            mutation_scale=float(el.get("head", DEFAULTS["arrow_head"])),
                                            color=color, lw=lw, shrinkA=0, shrinkB=0, zorder=_z(el, 1))
                ax.add_patch(p)
            elif t == "text":
                x, y = float(el["x"]), float(el["y"])
                ha = el.get("ha", "left")
                va = {"top": "top", "center": "center", "bottom": "bottom"}.get(el.get("va", "center"), "center")
                fp = bold_font if el.get("bold") else reg_font
                _t = ax.text(x, y, str(el.get("s", "")), fontproperties=fp,
                             fontsize=_resolve_size(el),
                             color=el.get("color", DEFAULTS["text_color"]),
                             ha=ha, va=va, wrap=True, zorder=_z(el, TEXT_ZORDER),
                             linespacing=float(DEFAULTS["line_height"]))
                text_artists.append((i, _t, el))
            else:
                print(f"[警告] 第 {i} 個元素 type='{t}' 不認得，略過", file=sys.stderr)
        except KeyError as ke:
            _err(f"第 {i} 個元素（type={t}）缺欄位 {ke}")
        except Exception as ee:
            _err(f"第 {i} 個元素（type={t}）繪製失敗：{ee}")

    return fig, ax, text_artists, W, H


def render_spec(spec, svg_path, png_path):
    """畫圖並存檔：呼叫 _build_figure 後同時輸出向量 SVG + 點陣 PNG 後援。"""
    import matplotlib.pyplot as plt
    fig, ax, _text_artists, W, H = _build_figure(spec)
    bg = spec.get("bg", "#ffffff")
    fig.savefig(svg_path, format="svg", facecolor=fig.get_facecolor(),
                edgecolor="none", transparent=(bg == "none"))
    fig.savefig(png_path, format="png", dpi=int(DEFAULTS["png_dpi"]), facecolor=fig.get_facecolor(),
                edgecolor="none", transparent=(bg == "none"))
    plt.close(fig)
    return W, H


# --------------------------------------------------------------------------
# 1b. --verify：畫圖後自檢（缺字＋版面幾何），只回純文字報告、PNG 不進脈絡
# --------------------------------------------------------------------------
_CMAP_CACHE = {}


def _font_codepoints(path):
    """載入字型的字元表（cmap），回傳可用字碼 set；載入失敗回傳 None（該層跳過、不當機）。"""
    if path in _CMAP_CACHE:
        return _CMAP_CACHE[path]
    try:
        from fontTools.ttLib import TTFont, TTCollection
        if path.lower().endswith(".ttc"):
            font = TTCollection(path).fonts[0]
        else:
            font = TTFont(path)
        cps = set(font.getBestCmap().keys())
    except Exception:
        cps = None
    _CMAP_CACHE[path] = cps
    return cps


def verify_spec(spec):
    """畫圖自檢：回傳 (ok: bool, report: str)。
    第一層 字型覆蓋：逐字查 cmap，抓豆腐字（缺字）。
    第二層 版面幾何：文字邊界框出界／重疊。
    只回純文字報告，不產生任何 PNG／SVG、不進對話脈絡。"""
    import warnings
    from collections import Counter
    lines = []

    elements = spec.get("elements", []) or []
    type_counts = Counter(el.get("type", "?") for el in elements)
    size = spec.get("size", [0, 0])
    try:
        W, H = float(size[0]), float(size[1])
    except Exception:
        W, H = 0.0, 0.0
    type_str = " / ".join(f"{k} {v}" for k, v in sorted(type_counts.items()))
    lines.append("[畫圖自檢 draw_docx_figure --verify]")
    lines.append(f"畫布：{int(W)} x {int(H)} px；元素 {len(elements)}（{type_str}）")

    # ── 第一層：字型覆蓋（cmap 缺字）──
    # 混排：字在「中文 TW-Kai」或「英數 Liberation Serif」任一字型即不算缺（union）
    _kai = _font_codepoints(CJK_FONT_PATH)
    _lat = _font_codepoints(FONT_LATIN_PATH)
    reg_cps = None if (_kai is None and _lat is None) else ((_kai or set()) | (_lat or set()))
    bold_cps = reg_cps
    missing = []
    if reg_cps is None and bold_cps is None:
        lines.append("第一層 字型覆蓋：略過（fontTools 不可用或字型讀取失敗）")
        layer1_ok = True
    else:
        for i, el in enumerate(elements):
            if el.get("type") != "text":
                continue
            s = str(el.get("s", ""))
            cps = bold_cps if el.get("bold") else reg_cps
            if cps is None:
                cps = reg_cps if reg_cps is not None else bold_cps
            miss = []
            for ch in s:
                o = ord(ch)
                if o in (0x20, 0x09, 0x0A, 0x0D):  # 空白/tab/換行不算
                    continue
                if cps is not None and o not in cps:
                    miss.append((ch, f"U+{o:04X}"))
            if miss:
                missing.append((i, s, miss))
        layer1_ok = not missing
        if layer1_ok:
            lines.append("第一層 字型覆蓋：缺字 0")
        else:
            lines.append(f"第一層 字型覆蓋：缺字 {len(missing)}")
            for i, s, miss in missing:
                mstr = "、".join(f"{code}({ch})" for ch, code in miss)
                stxt = s if len(s) <= 24 else s[:24] + "…"
                lines.append(f'  · 元素#{i} text "{stxt}"：缺 {mstr}')

    # ── 第二層：版面幾何（出界＋重疊＋貼邊）──
    # 文字用 matplotlib 渲染框（字寬取決於字型度量、非算得出）；
    # 非文字元素（矩形/線/箭頭/多邊形/橢圓）改用「由 spec 座標解析算出的邊界框」——
    # 精確、不依賴渲染，且避開 annotate 空字串註解的 get_window_extent 只回文字錨點、
    # 量不到箭頭實際跨距的坑。出界與貼邊兩項因此同時涵蓋文字與非文字。
    MARGIN_WARN_PX = 20.0   # 距畫布邊 < 此值即「貼邊不通過」（§B 升級為擋關 exit3；門檻沿用 20px 不改 30px：校準自 C-1 後版最小合法邊距 24.6px，現存 8 圖 24.6>20 皆不受影響，改 30 會誤殺既有圖）
    # 第三層 文字淨空門檻（皆提醒、不影響通過；校準自全 8 張真實圖）：
    SHAPE_GAP_WARN_PX = 12.0     # 文字距實心形狀邊 < 此值即「貼近提醒」。設提醒非擋關——圖表型圖的座標軸刻度標籤天生貼軸線，硬擋會誤傷 4 張合法圖
    ARROW_ENDPOINT_SKIP_PX = 90.0  # 文字框中心距箭頭端點 < 此值＝該箭頭連的節點標籤（刻意），不算被橫越；校準自 B-1/C-1 共 11 支箭頭零誤報
    out_of_canvas = []
    overlaps = []
    near_edge = []          # 貼邊提醒（不翻成不通過）
    shape_press = []        # 第三層：文字跨界壓到實心形狀（提醒）
    shape_near = []         # 第三層：文字貼近實心形狀 < 門檻（提醒）
    arrow_cross = []        # 第三層：箭頭橫越無關文字（提醒）
    layer2_note = None
    _KIND_ZH = {"line": "線/箭頭", "rect": "矩形", "polygon": "多邊形", "ellipse": "橢圓", "arc_arrow": "圓弧箭頭", "connect": "連接器箭頭", "text": "text"}

    try:
        _connect_geom = _resolve_connects(spec)
    except Exception:
        _connect_geom = {}   # 真正的錯誤會在下方 _build_figure 呼叫時由 _err 回報

    def _shape_extent(el):
        """非文字元素的解析式邊界框，回 (xmin, ymin, xmax, ymax) 或 None。"""
        t = el.get("type")
        try:
            if t == "rect":
                x, y, w, h = float(el["x"]), float(el["y"]), float(el["w"]), float(el["h"])
                return (min(x, x + w), min(y, y + h), max(x, x + w), max(y, y + h))
            if t == "line":
                x1, y1, x2, y2 = float(el["x1"]), float(el["y1"]), float(el["x2"]), float(el["y2"])
                return (min(x1, x2), min(y1, y2), max(x1, x2), max(y1, y2))
            if t == "polygon":
                pts = el.get("points") or []
                xs = [float(p[0]) for p in pts]
                ys = [float(p[1]) for p in pts]
                if not xs:
                    return None
                return (min(xs), min(ys), max(xs), max(ys))
            if t == "ellipse":
                cx, cy = float(el["cx"]), float(el["cy"])
                erx, ery = float(el["rx"]), float(el["ry"])
                return (cx - erx, cy - ery, cx + erx, cy + ery)
            if t == "arc_arrow":
                x1, y1, x2, y2 = float(el["x1"]), float(el["y1"]), float(el["x2"]), float(el["y2"])
                rad = float(el.get("rad", 0.3))
                cx, cy = _arc3_control(x1, y1, x2, y2, rad)  # 真實 render 控制點(含 y 倒置補償)
                xs = []; ys = []
                for k in range(41):
                    tt = k / 40.0; mt = 1 - tt
                    xs.append(mt * mt * x1 + 2 * mt * tt * cx + tt * tt * x2)
                    ys.append(mt * mt * y1 + 2 * mt * tt * cy + tt * tt * y2)
                return (min(xs), min(ys), max(xs), max(ys))
            if t == "connect":
                g = _connect_geom.get(id(el))
                if g is None:
                    return None
                x1, y1, x2, y2, rad = g
                cx, cy = _arc3_control(x1, y1, x2, y2, rad)  # 真實 render 控制點(含 y 倒置補償)
                xs = []; ys = []
                for k in range(41):
                    tt = k / 40.0; mt = 1 - tt
                    xs.append(mt * mt * x1 + 2 * mt * tt * cx + tt * tt * x2)
                    ys.append(mt * mt * y1 + 2 * mt * tt * cy + tt * tt * y2)
                return (min(xs), min(ys), max(xs), max(ys))
        except Exception:
            return None
        return None

    # ── 第三層幾何小工具 ──
    def _aabb_gap(a, b):
        """兩軸對齊矩形間隙：<0 表示兩軸皆重疊；>0 表示最短外距（含對角）。"""
        import math as _m
        ax0, ay0, ax1, ay1 = a; bx0, by0, bx1, by1 = b
        gx = max(bx0 - ax1, ax0 - bx1)   # >0 表 x 軸上分離
        gy = max(by0 - ay1, ay0 - by1)   # >0 表 y 軸上分離
        if gx > 0 and gy > 0:
            return _m.hypot(gx, gy)
        if gx > 0:
            return gx
        if gy > 0:
            return gy
        return -1.0

    def _contained(inner, outer, tol=1.0):
        """inner 矩形是否幾乎完全落在 outer 內（＝標籤在形狀內、刻意）。"""
        ix0, iy0, ix1, iy1 = inner; ox0, oy0, ox1, oy1 = outer
        return (ix0 >= ox0 - tol and iy0 >= oy0 - tol and
                ix1 <= ox1 + tol and iy1 <= oy1 + tol)

    def _box_center_near(box, px, py, r):
        import math as _m
        cx = (box[0] + box[2]) / 2.0; cy = (box[1] + box[3]) / 2.0
        return _m.hypot(cx - px, cy - py) <= r

    def _seg_hits_box(x1, y1, x2, y2, box):
        """線段是否穿過矩形（端點落在矩形內、或線段與任一邊相交皆算）。"""
        rx0, ry0, rx1, ry1 = box
        def _inside(px, py):
            return rx0 <= px <= rx1 and ry0 <= py <= ry1
        if _inside(x1, y1) or _inside(x2, y2):
            return True
        def _ccw(A, B, C):
            return (C[1] - A[1]) * (B[0] - A[0]) - (B[1] - A[1]) * (C[0] - A[0])
        def _inter(A, B, C, Dp):
            d1 = _ccw(C, Dp, A); d2 = _ccw(C, Dp, B)
            d3 = _ccw(A, B, C); d4 = _ccw(A, B, Dp)
            return ((d1 > 0) != (d2 > 0)) and ((d3 > 0) != (d4 > 0))
        corners = [(rx0, ry0), (rx1, ry0), (rx1, ry1), (rx0, ry1)]
        for k in range(4):
            if _inter((x1, y1), (x2, y2), corners[k], corners[(k + 1) % 4]):
                return True
        return False

    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            fig, ax, text_artists, W2, H2 = _build_figure(spec)
            fig.canvas.draw()
            renderer = fig.canvas.get_renderer()
            inv = ax.transData.inverted()
            boxes = []
            for idx, artist, el in text_artists:
                bb = artist.get_window_extent(renderer)
                (dx0, dy0) = inv.transform((bb.x0, bb.y0))
                (dx1, dy1) = inv.transform((bb.x1, bb.y1))
                xmin, xmax = sorted((dx0, dx1))
                ymin, ymax = sorted((dy0, dy1))
                boxes.append((idx, str(el.get("s", "")), xmin, ymin, xmax, ymax))
            plt.close(fig)

        tol = 1.0
        Wc = W2 or W
        Hc = H2 or H

        # 合併：文字（渲染框）＋非文字（解析框），供出界與貼邊檢查
        geo = []   # (idx, kind, label, xmin, ymin, xmax, ymax)
        for idx, s, xmin, ymin, xmax, ymax in boxes:
            geo.append((idx, "text", s, xmin, ymin, xmax, ymax))
        for i, el in enumerate(elements):
            if el.get("type") == "text":
                continue
            ext = _shape_extent(el)
            if ext is None:
                continue
            geo.append((i, el.get("type", "?"), "", ext[0], ext[1], ext[2], ext[3]))

        for idx, kind, s, xmin, ymin, xmax, ymax in geo:
            out = False
            if xmin < -tol:
                out_of_canvas.append((idx, kind, s, f"左緣 {xmin:.0f} < 0")); out = True
            if ymin < -tol:
                out_of_canvas.append((idx, kind, s, f"上緣 {ymin:.0f} < 0")); out = True
            if xmax > Wc + tol:
                out_of_canvas.append((idx, kind, s, f"右緣 {xmax:.0f} 超出畫布寬 {int(Wc)}")); out = True
            if ymax > Hc + tol:
                out_of_canvas.append((idx, kind, s, f"下緣 {ymax:.0f} 超出畫布高 {int(Hc)}")); out = True
            if out:
                continue
            # 貼邊提醒（僅對未出界者；跳過鋪滿整版的邊框元素避免噪音）
            spans_full = (xmax - xmin) >= 0.95 * Wc and (ymax - ymin) >= 0.95 * Hc
            if spans_full:
                continue
            cand = [(xmin, "左緣"), (ymin, "上緣"), (Wc - xmax, "右緣"), (Hc - ymax, "下緣")]
            m, edge = min(cand, key=lambda z: z[0])
            if m < MARGIN_WARN_PX:
                near_edge.append((idx, kind, s, m, edge))

        for a in range(len(boxes)):
            for b in range(a + 1, len(boxes)):
                ia, sa, ax0, ay0, ax1, ay1 = boxes[a]
                ib, sb, bx0, by0, bx1, by1 = boxes[b]
                ix = max(0.0, min(ax1, bx1) - max(ax0, bx0))
                iy = max(0.0, min(ay1, by1) - max(ay0, by0))
                inter = ix * iy
                if inter <= 0:
                    continue
                area_a = max(1e-6, (ax1 - ax0) * (ay1 - ay0))
                area_b = max(1e-6, (bx1 - bx0) * (by1 - by0))
                frac = inter / min(area_a, area_b)
                if frac >= 0.08:  # 重疊超過較小框 8% 才報，避免貼邊誤報
                    overlaps.append((ia, ib, sa, sb, frac))

        # ── 第三層：文字淨空（文字 vs 實心形狀＝框對框；文字 vs 箭頭＝線段判準）──
        # 皆為提醒、不影響通過。把「一次回看才知道」的壓字/貼近資訊搬進文字報告。
        shape_exts = []   # (idx, type, (x0,y0,x1,y1)) 實心形狀 rect/polygon/ellipse
        for i, el in enumerate(elements):
            if el.get("type") in ("rect", "polygon", "ellipse"):
                ext = _shape_extent(el)
                if ext is not None:
                    shape_exts.append((i, el.get("type"), ext))
        arrow_segs = []   # (idx, x1,y1,x2,y2) 帶箭頭的線
        for i, el in enumerate(elements):
            if el.get("type") == "line" and el.get("arrow"):
                try:
                    arrow_segs.append((i, float(el["x1"]), float(el["y1"]),
                                       float(el["x2"]), float(el["y2"])))
                except Exception:
                    pass
            elif el.get("type") in ("arc_arrow", "connect"):
                try:
                    if el.get("type") == "connect":
                        g = _connect_geom.get(id(el))
                        if g is None:
                            continue
                        x1, y1, x2, y2, rad = g
                    else:
                        x1, y1, x2, y2 = float(el["x1"]), float(el["y1"]), float(el["x2"]), float(el["y2"])
                        rad = float(el.get("rad", 0.3))
                    cx, cy = _arc3_control(x1, y1, x2, y2, rad)  # 真實 render 控制點(含 y 倒置補償)
                    pts = []
                    for k in range(9):
                        tt = k / 8.0; mt = 1 - tt
                        pts.append((mt * mt * x1 + 2 * mt * tt * cx + tt * tt * x2,
                                    mt * mt * y1 + 2 * mt * tt * cy + tt * tt * y2))
                    for k in range(8):
                        xa, ya = pts[k]; xb, yb = pts[k + 1]
                        arrow_segs.append((i, xa, ya, xb, yb))
                except Exception:
                    pass
        for (ti, ts, txn, tyn, txx, tyx) in boxes:
            tbox = (txn, tyn, txx, tyx)
            for (si, st, sext) in shape_exts:
                g = _aabb_gap(tbox, sext)
                if g < 0:                       # 兩軸皆重疊
                    if _contained(tbox, sext):
                        continue                # 標籤完全在形狀內＝刻意，放行
                    shape_press.append((ti, ts, si, st))
                elif g < SHAPE_GAP_WARN_PX:
                    shape_near.append((ti, ts, si, st, g))
            for (ai, x1, y1, x2, y2) in arrow_segs:
                if (_box_center_near(tbox, x1, y1, ARROW_ENDPOINT_SKIP_PX) or
                        _box_center_near(tbox, x2, y2, ARROW_ENDPOINT_SKIP_PX)):
                    continue                    # 文字在箭頭端點附近＝該箭頭節點標籤，刻意
                if _seg_hits_box(x1, y1, x2, y2, tbox):
                    arrow_cross.append((ai, ti, ts))
    except Exception as e:
        layer2_note = f"第二層 版面幾何：略過（渲染量測失敗：{e}）"

    if layer2_note:
        lines.append(layer2_note)
        layer2_ok = True
    else:
        layer2_ok = (not out_of_canvas) and (not overlaps) and (not near_edge)   # §B：貼邊升級為擋關（三項任一非空即不通過）
        lines.append(f"第二層 版面幾何：出界 {len(out_of_canvas)}；重疊 {len(overlaps)}；貼邊 {len(near_edge)}")
        for idx, kind, s, why in out_of_canvas:
            klab = _KIND_ZH.get(kind, kind)
            desc = f'"{s[:20]}…"' if len(s) > 20 else (f'"{s}"' if s else "")
            lines.append(f'  · 元素#{idx} {klab} {desc} {why}')
        for ia, ib, sa, sb, frac in overlaps:
            ta = sa if len(sa) <= 14 else sa[:14] + "…"
            tb = sb if len(sb) <= 14 else sb[:14] + "…"
            lines.append(f'  · 元素#{ia}"{ta}" 與 元素#{ib}"{tb}" 重疊 {frac*100:.0f}%（以較小框計）')
        for idx, kind, s, m, edge in near_edge:
            klab = _KIND_ZH.get(kind, kind)
            desc = f'"{s[:16]}…"' if len(s) > 16 else (f'"{s}"' if s else "")
            lines.append(f'  · [貼邊不通過] 元素#{idx} {klab} {desc} 距{edge}僅 {m:.0f}px（門檻 {int(MARGIN_WARN_PX)}px，請加大留白或擴畫布）')

    # ── 第三層 文字淨空報告（皆提醒、不影響通過與否）──
    if not layer2_note:
        _CAP = 8
        lines.append(f"第三層 文字淨空：文字壓形狀 {len(shape_press)}；文字貼近形狀 {len(shape_near)}；箭頭橫越文字 {len(arrow_cross)}（皆提醒、不影響通過）")
        for ti, ts, si, st in shape_press[:_CAP]:
            klab = _KIND_ZH.get(st, st)
            t = ts if len(ts) <= 14 else ts[:14] + "…"
            lines.append(f'  · [提醒] 文字#{ti}"{t}" 壓到{klab}#{si}（跨越邊界、未完全在形狀內）')
        for ti, ts, si, st, g in shape_near[:_CAP]:
            klab = _KIND_ZH.get(st, st)
            t = ts if len(ts) <= 14 else ts[:14] + "…"
            lines.append(f'  · [提醒] 文字#{ti}"{t}" 貼近{klab}#{si} 僅 {g:.0f}px（門檻 {int(SHAPE_GAP_WARN_PX)}px）')
        for ai, ti, ts in arrow_cross[:_CAP]:
            t = ts if len(ts) <= 14 else ts[:14] + "…"
            lines.append(f'  · [提醒] 箭頭#{ai} 橫越文字#{ti}"{t}"（非端點節點標籤、疑穿過無關文字）')
        for nm, coll in (("文字壓形狀", shape_press), ("文字貼近形狀", shape_near), ("箭頭橫越文字", arrow_cross)):
            if len(coll) > _CAP:
                lines.append(f'  · …{nm}另有 {len(coll) - _CAP} 處未列（避免洗版）')

    # ── ring 凸度讀數：把「弧內凹/外鼓」變成數字，取代讀 PNG 目視（§66：agent 一律不讀圖）──
    _curved = [(el, _connect_geom[id(el)]) for el in elements
               if isinstance(el, dict) and el.get("type") == "connect"
               and id(el) in _connect_geom and abs(_connect_geom[id(el)][4]) > 1e-6]
    if _curved:
        _pts = []
        for _el, (sx, sy, ex, ey, _r) in _curved:
            _pts.append((sx, sy)); _pts.append((ex, ey))
        _Ox = sum(p[0] for p in _pts) / len(_pts)
        _Oy = sum(p[1] for p in _pts) / len(_pts)
        _conv = []
        for _el, (sx, sy, ex, ey, r) in _curved:
            _cx, _cy = _arc3_control(sx, sy, ex, ey, r)          # 真實 render 控制點
            _ax = 0.25 * sx + 0.5 * _cx + 0.25 * ex              # 弧峰 t=0.5
            _ay = 0.25 * sy + 0.5 * _cy + 0.25 * ey
            _mmx, _mmy = (sx + ex) / 2.0, (sy + ey) / 2.0
            _dmid = math.hypot(_mmx - _Ox, _mmy - _Oy) or 1e-9
            _conv.append(math.hypot(_ax - _Ox, _ay - _Oy) / _dmid)
        _cmin, _cmax = min(_conv), max(_conv)
        lines.append(f"[ring 凸度] {len(_curved)} 條曲線箭頭；凸度＝弧峰半徑/弦中點半徑（>1 外鼓凸、<1 內凹）"
                     f" 最小 {_cmin:.3f}／最大 {_cmax:.3f}")
        if _cmin < 1.0:
            lines.append(f"  · [提醒] 有箭頭凸度<1＝內凹；若本意是外鼓成環，符號決策有誤，請查 _arc3_control/_resolve_connects")

    has_advice = bool(shape_press) or bool(shape_near) or bool(arrow_cross)  # §B：near_edge 已升級擋關、通過分支不會有貼邊，故移除
    ok = layer1_ok and layer2_ok
    if ok:
        if has_advice:
            lines.append("判定：通過（免讀 PNG 目視）；有文字淨空提醒，請確認是否為刻意設計，勿為此讀 PNG")
        else:
            lines.append("判定：通過（免讀 PNG 目視）")
    else:
        lines.append("判定：不通過（出界／重疊／貼邊，請修 JSON 形狀規格後重跑 --verify；勿把 PNG 讀回脈絡目視）")
    return ok, "\n".join(lines)


# --------------------------------------------------------------------------
# 2. 把 SVG（真向量）+ PNG（後援）嵌進 .docx
# --------------------------------------------------------------------------
def embed_into_docx(svg_path, png_path, out_docx, append_to=None,
                    caption=None, width_cm=15.0, canvas_wh=None):
    try:
        import docx
        from docx.shared import Cm
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        import lxml.etree as etree
    except Exception as e:
        _err(f"python-docx / lxml 無法載入：{e}")

    if append_to:
        if not os.path.exists(append_to):
            _err(f"--append-to 指定的檔案不存在：{append_to}")
        document = docx.Document(append_to)
    else:
        document = docx.Document()

    # 2a. 先用 python-docx 正常插入 PNG（自動建立 blip + 關聯 + media 部件）
    para = document.add_paragraph()
    para.alignment = 1  # 置中
    run = para.add_run()
    inline_shape = run.add_picture(png_path, width=Cm(float(width_cm)))

    # 2b. 把 SVG 加成新的 image 部件並關聯
    with open(svg_path, "rb") as f:
        svg_bytes = f.read()
    pkg = document.part.package
    # 找一個不衝突的 media 檔名
    n = 1
    existing = {p.partname for p in pkg.iter_parts()}
    while PackURI(f"/word/media/vecfig{n}.svg") in existing:
        n += 1
    svg_partname = PackURI(f"/word/media/vecfig{n}.svg")
    svg_part = Part(svg_partname, "image/svg+xml", svg_bytes, pkg)
    rId_svg = document.part.relate_to(svg_part, RT.IMAGE)

    # 2c. 在剛建立的 a:blip 底下掛 svgBlip 擴充
    from docx.oxml.ns import qn
    drawing = run._r.find(qn("w:drawing"))
    if drawing is None:
        _err("找不到剛插入圖片的 w:drawing 節點")
    blip = None
    for el in drawing.iter("{%s}blip" % A_NS):
        blip = el
        break
    if blip is None:
        _err("找不到 a:blip 節點，無法掛 SVG 向量")
    extLst = etree.SubElement(blip, "{%s}extLst" % A_NS)
    ext = etree.SubElement(extLst, "{%s}ext" % A_NS)
    ext.set("uri", SVG_EXT_URI)
    svgBlip = etree.SubElement(ext, "{%s}svgBlip" % ASVG_NS, nsmap={"asvg": ASVG_NS})
    svgBlip.set("{%s}embed" % R_NS, rId_svg)

    # 2d. 圖說
    if caption:
        cap = document.add_paragraph(caption)
        cap.alignment = 1

    target = append_to if append_to else out_docx
    document.save(target)
    return target, rId_svg, str(svg_partname)


# --------------------------------------------------------------------------
# 3. 內建範例（冰山概念圖，重現畢業專題 圖1-1 那類結構示意圖）
# --------------------------------------------------------------------------
def demo_iceberg():
    W, H = 1080, 920
    rows = [
        ("氣候變遷挑戰", "溫室效應、海平面上升、生物多樣性威脅", "#e05a4d"),
        ("政策方向", "SDGs、NDCs、碳紀錄管理制度、管理會計", "#3b6fd4"),
        ("技術挑戰", "維度：碳紀錄成本、資料完整度、交易條件複雜度", "#8a8a8a"),
        ("研究動機", "制度性障礙、中小企業參與度、資訊技術解決方案", "#e0b400"),
        ("研究目的", "光達 LiDAR、區塊鏈、智能合約驅動林地碳紀錄管理平台", "#2fa565"),
        ("研究成果對政策推動的影響", "提升中小企業內部政策執行效率", "#e07b2f"),
    ]
    els = []
    # 中間冰山多邊形
    els.append({"type": "polygon",
                "points": [[540, 110], [610, 210], [560, 300], [640, 360], [560, 460],
                           [660, 560], [560, 660], [620, 760], [540, 850],
                           [460, 760], [520, 660], [420, 560], [520, 460], [440, 360],
                           [520, 300], [470, 210]],
                "fill": "none", "stroke": "#555555", "lw": 1.6, "close": True})
    # 虛線水線
    els.append({"type": "line", "x1": 60, "y1": 300, "x2": 1020, "y2": 300,
                "stroke": "#333333", "lw": 1.4, "dash": [8, 6]})
    top = 210
    rowh = 96
    gap = 20
    for i, (title, desc, color) in enumerate(rows):
        cy = top + i * (rowh + gap)
        # 左框（標題）
        els.append({"type": "rect", "x": 70, "y": cy, "w": 330, "h": rowh,
                    "fill": "#ffffff", "stroke": "#333333", "lw": 1.4, "rx": 10})
        els.append({"type": "text", "x": 110, "y": cy + rowh / 2, "s": title,
                    "size": 19, "color": color, "ha": "left", "va": "center", "bold": True})
        # 右框（說明）
        els.append({"type": "rect", "x": 700, "y": cy, "w": 310, "h": rowh,
                    "fill": "#ffffff", "stroke": "#333333", "lw": 1.4, "rx": 10})
        els.append({"type": "text", "x": 990, "y": cy + rowh / 2, "s": desc,
                    "size": 13, "color": "#444444", "ha": "right", "va": "center"})
        # 連接線
        els.append({"type": "line", "x1": 400, "y1": cy + rowh / 2, "x2": 470, "y2": cy + rowh / 2,
                    "stroke": "#999999", "lw": 1.0})
        els.append({"type": "line", "x1": 610, "y1": cy + rowh / 2, "x2": 700, "y2": cy + rowh / 2,
                    "stroke": "#999999", "lw": 1.0})
    return {"size": [W, H], "bg": "#ffffff", "elements": els}


# --------------------------------------------------------------------------
def load_spec(args):
    if args.demo:
        if args.demo == "iceberg":
            return demo_iceberg()
        _err(f"未知的 --demo 範例：{args.demo}（目前只有 iceberg）")
    raw = None
    if args.spec_json:
        raw = args.spec_json
    elif args.spec:
        if not os.path.exists(args.spec):
            _err(f"--spec 檔案不存在：{args.spec}")
        with open(args.spec, "r", encoding="utf-8") as f:
            raw = f.read()
    elif not sys.stdin.isatty():
        raw = sys.stdin.read()
    if not raw or not raw.strip():
        _err("沒有提供形狀規格：請用 --spec / --spec-json / 標準輸入 / --demo 其中一種")
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        _err(f"JSON 解析失敗：{e}")


def _draw_main():
    ap = argparse.ArgumentParser(
        description="由精簡 JSON 形狀規格產生 Word 原生向量圖（SVG 真向量 + PNG 後援，省 token）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="範例：python3 draw_docx_figure.py --demo iceberg --out /tmp/iceberg.docx")
    ap.add_argument("--spec", help="JSON 形狀規格檔路徑")
    ap.add_argument("--spec-json", help="直接以字串傳入 JSON 形狀規格")
    ap.add_argument("--demo", help="內建範例名稱（目前：iceberg）")
    ap.add_argument("--out", help="輸出的新 .docx 路徑")
    ap.add_argument("--append-to", help="把圖附加到既有 .docx（而非新建）")
    ap.add_argument("--caption", help="圖說文字（置中放在圖下方）")
    ap.add_argument("--width-cm", type=float, default=15.0, help="插入寬度（公分，預設 15）")
    ap.add_argument("--keep-assets", action="store_true", help="保留中繼 .svg/.png（預設仍會印出路徑）")
    ap.add_argument("--verify", action="store_true",
                    help="自檢模式：只檢查 JSON 規格的缺字與版面（文字出界/重疊），回純文字報告、不產生圖、不需 --out；通過 exit 0、不通過 exit 3")
    args = ap.parse_args()

    if args.verify:
        spec = load_spec(args)
        ok, report = verify_spec(spec)
        print(report)
        sys.exit(0 if ok else 3)

    if not args.out and not args.append_to:
        _err("必須指定 --out（新建）或 --append-to（附加到既有檔）其中一個")

    spec = load_spec(args)

    base = os.path.splitext(args.out or args.append_to)[0]
    svg_path = base + ".vecfig.svg"
    png_path = base + ".vecfig.png"

    W, H = render_spec(spec, svg_path, png_path)
    target, rId, svg_part = embed_into_docx(
        svg_path, png_path, args.out, append_to=args.append_to,
        caption=args.caption, width_cm=args.width_cm, canvas_wh=(W, H))

    svg_sz = os.path.getsize(svg_path)
    png_sz = os.path.getsize(png_path)
    print(f"[完成] 已產生 Word 原生向量圖")
    print(f"  畫布尺寸 : {int(W)} x {int(H)} px")
    print(f"  向量 SVG : {svg_path}  ({svg_sz:,} bytes)")
    print(f"  後援 PNG : {png_path}  ({png_sz:,} bytes)")
    print(f"  輸出 docx: {target}")
    print(f"  SVG 關聯 : {rId} -> {svg_part}")
    if not args.keep_assets:
        print(f"  （中繼 .svg/.png 已保留於上列路徑，可自行刪除；加 --keep-assets 只是明示保留）")


# === ARCUS 併入工具 · 論文卡片抽取（含 pdf） (2026-07-19) ===
# 原 runtime/paper_cards.py，快取區改在 arcus 專案資料夾底下。

# -*- coding: utf-8 -*-
"""論文卡片機制。

一篇論文走一次：抽文字 → 切段編鍵 → 開一次性模型行程萃取六欄 → 存卡。
之後重看只讀卡片（約 1500 字元），要核對出處才點名取原文段落。

三個檔分工（同一個資料夾，名稱取檔案內容雜湊前 16 碼）：
  meta.json  檔頭：標題、作者、年份、頁數、來源、抽取時間
  card.json  六欄卡片，常讀
  text.json  切段全文，備查才讀
另有 index.json 當總表，跨論文比較靠它。
"""
import os
import json
import time
import hashlib
import urllib.request

PAPERS_DIR = os.path.join(TOOLS_DIR, '_papers')
INDEX_PATH = os.path.join(PAPERS_DIR, 'index.json')

SEP = '\n\n'                 # 段落接合字串；還原就是照鍵的順序用它串回去
SEG_TARGET = 1000            # 碎塊黏到這個字元數就收尾成一段
BATCH_LIMIT = 150000         # 超過這個字元數改分批讀（兩百頁以上的專書才會遇到）
CARD_TIMEOUT = 600           # 讀六萬字元加萃取六欄，120 秒一定不夠
CARD_MODEL = 'opus'          # arcus 一律用 Opus

FIELDS = ['研究問題', '研究方法', '研究流程', '研究成果', '爭議', '未來發展']
META_FIELDS = ['標題', '作者', '年份']


# ── 基礎工具 ──────────────────────────────────────────────────────────────────
def _sha16(data):
    return hashlib.sha256(data).hexdigest()[:16]


def _paper_dir(pid):
    return os.path.join(PAPERS_DIR, pid)


def _pc_read_json(path, default=None):
    try:
        with open(path, encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return default


def _pc_write_json(path, obj):
    """先寫暫存再換名，中途斷掉不會留下半份壞檔。"""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = path + '.tmp'
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(obj, f, ensure_ascii=False, indent=1)
    os.replace(tmp, path)


def _fetch(source):
    """來源可以是網址或虛擬機上的檔案路徑，一律回傳位元組內容。"""
    if source.startswith('http://') or source.startswith('https://'):
        req = urllib.request.Request(source, headers={'User-Agent': 'arcus-paper-cards/1.0'})
        with urllib.request.urlopen(req, timeout=120) as r:
            return r.read()
    p = os.path.expanduser(source)
    if not os.path.isabs(p):
        # 相對路徑以專案資料夾為準，與 read／query 同一套解析
        p = _arcus_read_path(source)
    with open(p, 'rb') as f:
        return f.read()


# ── 抽取與切段 ────────────────────────────────────────────────────────────────
def _blocks(pdf_bytes):
    """抽出所有非空白區塊，回傳 [(頁碼, 文字)]；同時回傳總頁數。"""
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype='pdf')
    out = []
    for pno in range(len(doc)):
        for b in doc[pno].get_text('blocks'):
            t = (b[4] or '').strip()
            if t:
                out.append((pno + 1, t))
    pages = len(doc)
    pdf_meta = dict(doc.metadata or {})
    doc.close()
    return out, pages, pdf_meta


def _merge(blocks):
    """碎塊黏到約 SEG_TARGET 字元一段。鍵取該段起始頁加頁內序號，過濾後重新編號。"""
    segs = []
    cur = []
    start = None
    per_page = {}

    def flush():
        if not cur:
            return
        per_page[start] = per_page.get(start, 0) + 1
        segs.append(('p%03d#%02d' % (start, per_page[start]), SEP.join(cur)))

    for p, t in blocks:
        if start is None:
            start = p
        cur.append(t)
        if len(SEP.join(cur)) >= SEG_TARGET:
            flush()
            cur = []
            start = None
    flush()
    return segs


def _full_text(segments):
    """整份文字直接定義為各段用 SEP 串起來，所以還原不可能出錯。"""
    return SEP.join(v for _, v in segments)


def _labelled(segments):
    """餵給模型的樣子：每段前面掛著鍵，模型才有辦法報出處。"""
    return '\n\n'.join('[%s]\n%s' % (k, v) for k, v in segments)


# ── 萃取卡片 ──────────────────────────────────────────────────────────────────
_CARD_PROMPT = '''以下是一篇論文的全文，每一段前面用中括號標了段落鍵，例如 [p012#03]。

請閱讀全文，只輸出一個 JSON 物件，不要輸出任何其他文字，不要用程式碼圍籬包起來。

格式：
{
  "標題": "論文標題",
  "作者": "作者，多位用頓號分隔",
  "年份": "西元年份，查不到就填未標示",
  "科別": "學科領域，例如公共行政、教育科技、資訊工程，多個用頓號分隔",
  "研究問題": {"text": "……", "refs": ["p002#01"]},
  "研究方法": {"text": "……", "refs": ["p012#03"]},
  "研究流程": {"text": "……", "refs": ["p014#01"]},
  "研究成果": {"text": "……", "refs": ["p031#02"]},
  "爭議": {"text": "……", "refs": ["p045#01"]},
  "未來發展": {"text": "……", "refs": ["p047#03"]}
}

規則：
- refs 只能填實際出現在上面全文裡的段落鍵，禁止自行編造或推測。
- 每一欄的 text 用完整中文敘述，長度 200 到 400 字，不要條列。
- 全文沒有述及的欄位，text 填「原文未述及」，refs 填空陣列。
- 爭議指的是研究本身的限制、與既有研究衝突之處、作者自承的疑慮。

全文開始：

%s
'''


def _parse_card(raw):
    """模型偶爾會在前後多講幾句，取第一個左大括號到最後一個右大括號。"""
    if not raw:
        return None
    i = raw.find('{')
    j = raw.rfind('}')
    if i < 0 or j <= i:
        return None
    try:
        return json.loads(raw[i:j + 1])
    except Exception:
        return None


def _verify_refs(card, valid_keys):
    """出處鍵由程式回查，查不到就標記出處存疑。模型只負責內容是什麼。"""
    suspect = 0
    for f in FIELDS:
        cell = card.get(f)
        if not isinstance(cell, dict):
            card[f] = {'text': str(cell or '原文未述及'), 'refs': [], '出處存疑': True}
            suspect += 1
            continue
        refs = [r for r in (cell.get('refs') or []) if isinstance(r, str)]
        good = [r for r in refs if r in valid_keys]
        bad = [r for r in refs if r not in valid_keys]
        cell['refs'] = good
        if bad:
            cell['bad_refs'] = bad
            cell['出處存疑'] = True
            suspect += 1
    return suspect


def _call(prompt):
    """借用 arcus_core 的一次性模型行程；延後匯入避免兩個模組互相依賴。"""
    from arcus_core import _call_model
    return _call_model(prompt, model=CARD_MODEL, timeout=CARD_TIMEOUT)


# ── 總表 ──────────────────────────────────────────────────────────────────────
def _index_load():
    return _pc_read_json(INDEX_PATH, {}) or {}


def _index_put(pid, entry):
    idx = _index_load()
    idx[pid] = entry
    _pc_write_json(INDEX_PATH, idx)


def _resolve_pid(given):
    """把給的字串解析成真正的論文 id（16 碼雜湊）。依序：精確命中索引 → 磁碟有該目錄
    → id 前綴唯一命中 → 標題含該子字串唯一命中。回傳 (pid, candidates)；
    pid 為 None 時 candidates 供呼叫端回報消歧義。"""
    if not given or not isinstance(given, str):
        return None, []
    g = given.strip()
    idx = _index_load()
    if g in idx:
        return g, []
    if g and '/' not in g and '\\' not in g and os.path.isdir(_paper_dir(g)):
        return g, []
    gl = g.lower()
    pref = [k for k in idx if k.lower().startswith(gl) or gl.startswith(k.lower())]
    if len(pref) == 1:
        return pref[0], []
    if len(pref) > 1:
        return None, [{'id': k, '標題': idx[k].get('標題', '')} for k in pref]
    tt = [k for k in idx if gl in str(idx[k].get('標題', '')).lower()]
    if len(tt) == 1:
        return tt[0], []
    if len(tt) > 1:
        return None, [{'id': k, '標題': idx[k].get('標題', '')} for k in tt]
    return None, []


# ── 對外四個動作 ──────────────────────────────────────────────────────────────
def ingest(sources, force=False):
    """批次收論文。sources 是網址或虛擬機路徑的清單。已收過就直接命中快取。"""
    if isinstance(sources, str):
        sources = [sources]
    results = []
    for src in (sources or []):
        t0 = time.time()
        try:
            results.append(_ingest_one(src, force, t0))
        except Exception as e:
            results.append({'source': src, 'ok': False, 'error': '%s：%s' % (type(e).__name__, e)})
    return {'ok': any(r.get('ok') for r in results), 'papers': results}


_PAPER_CONNECTOR_MARK = "2026-07-20"


def _arxiv_id_from(sv):
    """從 arXiv id 或 arXiv 網址取出純 id；不是就回 None。"""
    import re
    x = (sv or '').strip()
    m = re.match(r'^(?:arxiv:)?(\d{4}\.\d{4,5})(v\d+)?$', x, re.I)
    if m:
        return m.group(1) + (m.group(2) or '')
    m = re.search(r'arxiv\.org/(?:abs|pdf)/([^\s?#]+?)(?:\.pdf)?$', x, re.I)
    if m:
        return m.group(1)
    return None


def _arxiv_lookup(arxiv_id):
    """回 arXiv 權威著錄，全逐字，不經模型。無全文回 None。"""
    import urllib.parse, xml.etree.ElementTree as ET
    ns = {'a': 'http://www.w3.org/2005/Atom', 'arxiv': 'http://arxiv.org/schemas/atom'}
    url = ('http://export.arxiv.org/api/query?id_list=%s&max_results=1'
           % urllib.parse.quote(arxiv_id))
    req = urllib.request.Request(url, headers={'User-Agent': 'arcus-paper-cards/1.0'})
    with urllib.request.urlopen(req, timeout=60) as r:
        raw = r.read()
    e = ET.fromstring(raw).find('a:entry', ns)
    if e is None:
        return None
    # 老舊著錄的 author 區塊會混入雜項（如 Keywords、MSC-class），只收真正的人名
    junk = ('keywords', 'msc-class', 'msc class', 'report-no', 'acm-class')
    authors = []
    for a in e.findall('a:author', ns):
        nm = (a.findtext('a:name', '', ns) or '').strip()
        if nm and nm.lower() not in junk:
            authors.append(nm)
    cats = [c.get('term') for c in e.findall('a:category', ns) if c.get('term')]
    pdf = ''
    for ln in e.findall('a:link', ns):
        if ln.get('title') == 'pdf' or ln.get('type') == 'application/pdf':
            pdf = ln.get('href')
            break
    if not pdf:
        return None
    return {'title': (e.findtext('a:title', '', ns) or '').strip(),
            'authors': authors, 'categories': cats,
            'date': (e.findtext('a:published', '', ns) or '').strip(),
            'source': 'arxiv',
            'landing_url': (e.findtext('a:id', '', ns) or '').strip(),
            'pdf_url': pdf}


def _journal_lookup(identifier):
    """回權威著錄（全逐字）。目前只認 arXiv；查無或無全文回 None。
       慢慢增加來源時，在這裡多接一個轉接器即可，主流程不動。"""
    aid = _arxiv_id_from(identifier)
    if aid:
        try:
            return _arxiv_lookup(aid)
        except Exception:
            return None
    return None


def _resolve_source(src):
    """判斷輸入屬於哪一種，回 (rec, fetch_src, err)。
       識別碼 → 走索引庫著錄；直接 pdf 來源 → 逕行下載、meta 留白不猜；
       裸關鍵字 → 不建卡，導向 discover（避免關鍵字搜尋張冠李戴）。"""
    s = (src or '').strip()
    rec = _journal_lookup(s)
    if rec is not None:
        return rec, rec['pdf_url'], None
    low = s.lower()
    if (low.startswith('http://') or low.startswith('https://') or low.endswith('.pdf')
            or os.path.isabs(s)):
        return None, s, None
    return None, None, ('無法辨識來源「%s」。建卡入口只吃 arXiv id 或直接 pdf 來源（DOI 等其他來源待後續轉接器）；'
                        '想用關鍵字找論文請改用 discover，挑定後再把識別碼交來建卡。' % s[:60])


def _direct_title(pdf_meta, fetch_src, src):
    """直接上傳的 PDF 無權威著錄時取標題：PDF 內建標題與來源檔名字幹，取較像標題者。
    排版原稿副檔名、頁碼範圍、數字多於字母者視為垃圾(負分)。"""
    import re
    def _score(t):
        t = (t or '').strip()
        if not t:
            return -1
        low = t.lower()
        if re.search(r'\.(indd|doc|docx|qxd|pmd|fm)\b', low):
            return -1
        if re.search(r'\d+\s*\.\.\s*\d+', t):
            return -1
        letters = sum(c.isalpha() for c in t)
        digits = sum(c.isdigit() for c in t)
        if letters == 0 or digits > letters:
            return -1
        return len([w for w in re.split(r'[\s_\-]+', t)
                    if sum(c.isalpha() for c in w) >= 3])
    pm = (pdf_meta.get('title') or '').strip() if isinstance(pdf_meta, dict) else ''
    ss = (fetch_src or src or '').split('?')[0]
    stem = os.path.splitext(os.path.basename(ss.rstrip('/')))[0].replace('_', ' ').strip()
    best = max([pm, stem], key=_score)
    return best if _score(best) >= 0 else (stem or pm or '')


def _salvage_biblio(card):
    """直接上傳無權威索引時，救回模型從標題頁抽出的書目欄（標題/作者/科別/年份）。
    作者、科別統一正規化成清單；年份查無則留空。"""
    import re
    def _as_list(v):
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
        return [p.strip() for p in re.split(r'[、；;／/]+', str(v or '')) if p.strip()]
    title = card.get('標題')
    title = title.strip() if isinstance(title, str) else ''
    year = card.get('年份')
    year = year.strip() if isinstance(year, str) else ''
    if year in ('未標示', '未知', '無', 'N/A', 'na', 'NA'):
        year = ''
    return title, _as_list(card.get('作者')), _as_list(card.get('科別')), year


def _ingest_one(src, force, t0):
    rec, fetch_src, err = _resolve_source(src)
    if err:
        return {'source': src, 'ok': False, 'error': err}
    raw = _fetch(fetch_src)
    pid = _sha16(raw)
    d = _paper_dir(pid)
    card_path = os.path.join(d, 'card.json')

    if os.path.exists(card_path) and not force:
        meta = _pc_read_json(os.path.join(d, 'meta.json'), {}) or {}
        return {'source': src, 'ok': True, 'id': pid, 'cached': True,
                'title': meta.get('標題', ''), 'card': _pc_read_json(card_path)}

    blocks, pages, pdf_meta = _blocks(raw)
    if not blocks:
        return {'source': src, 'ok': False, 'id': pid,
                'error': '抽不到任何文字，可能是掃描型檔案，需要文字辨識'}

    segments = _merge(blocks)
    full = _full_text(segments)
    if len(full) > BATCH_LIMIT:
        return {'source': src, 'ok': False, 'id': pid,
                'error': '全文 %d 字元超過單次上限 %d，尚未支援分批' % (len(full), BATCH_LIMIT)}

    _pc_write_json(os.path.join(d, 'text.json'), {
        'sep': SEP,
        'sha256': hashlib.sha256(full.encode('utf-8')).hexdigest(),
        'chars': len(full),
        'segments': dict(segments),
    })

    card = _parse_card(_call(_CARD_PROMPT % _labelled(segments)))
    if not isinstance(card, dict):
        return {'source': src, 'ok': False, 'id': pid,
                'error': '模型沒有回傳可解析的卡片，切段全文已保留'}

    suspect = _verify_refs(card, set(k for k, _ in segments))
    _r = rec or {}
    # 索引庫沒值時，回填模型從標題頁抽出的書目欄（標題頁比檔名更準）
    _sv_title, _sv_authors, _sv_cats, _sv_year = _salvage_biblio(card)
    meta = {
        '標題': _r.get('title') or _sv_title or _direct_title(pdf_meta, fetch_src, src),
        '作者': _r.get('authors') or _sv_authors,
        '科別': _r.get('categories') or _sv_cats,
        '日期': _r.get('date') or _sv_year,
        '來源': _r.get('source') or 'direct',
        '來源網址': _r.get('landing_url') or '',
        'pdf網址': fetch_src,
        '頁數': pages,
        '字元數': len(full),
        '段數': len(segments),
        '抽取時間': time.strftime('%Y-%m-%d %H:%M:%S'),
        '已編修': False,
    }
    # meta 一律以權威著錄為準，把模型吐的 meta 欄全部丟掉
    for k in ('標題', '作者', '年份', '科別', '日期'):
        card.pop(k, None)

    _pc_write_json(os.path.join(d, 'meta.json'), meta)
    _pc_write_json(card_path, card)
    _index_put(pid, {'標題': meta['標題'], '日期': meta['日期'],
                     '科別': meta['科別'], '來源': meta['來源'], '段數': len(segments)})

    return {'source': src, 'ok': True, 'id': pid, 'cached': False,
            'title': meta['標題'], 'pages': pages, 'chars': len(full),
            'segments': len(segments), 'suspect_fields': suspect,
            'seconds': round(time.time() - t0, 1), 'card': card}


def cards(ids=None, fields=None):
    """批次取卡片。只給 fields 就跨論文比較同一欄，這是省最多的用法。"""
    idx = _index_load()
    ids = ids or list(idx.keys())
    fields = fields or FIELDS
    out = []
    for pid in ids:
        card = _pc_read_json(os.path.join(_paper_dir(pid), 'card.json'))
        if not card:
            out.append({'id': pid, 'ok': False, 'error': '沒有這篇的卡片'})
            continue
        meta = _pc_read_json(os.path.join(_paper_dir(pid), 'meta.json'), {}) or {}
        out.append({'id': pid, 'ok': True, '標題': meta.get('標題', ''),
                    '欄位': {f: card.get(f) for f in fields if f in card}})
    return {'ok': True, 'papers': out}


def quote(paper_id, keys=None, search=None):
    """取原文段落核對卡片有無依據。三種用法：keys=段落鍵清單→取那幾段；
    search=片語→回傳含該片語的段落鍵與片段（不必先知道鍵）；兩者都不給→回傳這篇有哪些段落鍵。
    paper_id 可給 id 前綴或標題子字串。"""
    pid, cands = _resolve_pid(paper_id)
    if not pid:
        return {'ok': False, 'error': '找不到這篇（給的 id／標題無法唯一對應）', 'candidates': cands}
    data = _pc_read_json(os.path.join(_paper_dir(pid), 'text.json'))
    if not data:
        return {'ok': False, 'id': pid, 'error': '沒有這篇的切段全文'}
    segs = data.get('segments') or {}
    all_keys = list(segs.keys())
    if isinstance(keys, str):
        keys = [keys]
    if search:
        needle = str(search).lower()
        hits = {}
        for k, v in segs.items():
            t = str(v)
            pos = t.lower().find(needle)
            if pos >= 0:
                s = max(0, pos - 80)
                hits[k] = ('…' if s > 0 else '') + t[s:pos + len(str(search)) + 160]
        return {'ok': bool(hits), 'id': pid, 'search': search,
                'matches': hits, 'match_count': len(hits), 'total_segments': len(all_keys)}
    if not keys:
        return {'ok': True, 'id': pid, 'available_keys': all_keys,
                'total_segments': len(all_keys),
                'hint': '再用 keys=[...] 取指定段，或 search="片語" 依內文搜尋'}
    got = {k: segs[k] for k in keys if k in segs}
    missing = [k for k in keys if k not in segs]
    out = {'ok': bool(got), 'id': pid, 'segments': got, 'missing': missing}
    if missing:
        out['available_keys'] = all_keys
    return out


def list_papers():
    """總表，只回標題與年份，不碰任何卡片內容。"""
    idx = _index_load()
    return {'ok': True, 'count': len(idx),
            'papers': [dict(id=k, **v) for k, v in idx.items()]}

def delete_paper(paper_id):
    """刪掉一張卡片:先從 index.json 拿掉那筆,再刪卡片目錄。兩步一起做,避免留殘卡。
    安全上限:只准刪 PAPERS_DIR 的直接子目錄,任何超出範圍的路徑一律拒絕。"""
    if (not paper_id or not isinstance(paper_id, str)
            or '/' in paper_id or '\\' in paper_id or paper_id.startswith('.')):
        return {'ok': False, 'error': 'paper_delete 需要合法的 payload.id(單純卡號,不含斜線)'}
    idx = _index_load()
    d = _paper_dir(paper_id)
    dir_exists = os.path.isdir(d)
    in_index = paper_id in idx
    if not dir_exists and not in_index:
        return {'ok': False, 'error': '找不到這張卡片:%s(index 與磁碟都沒有)' % paper_id}
    real = os.path.realpath(d)
    if dir_exists:
        base = os.path.realpath(PAPERS_DIR)
        if os.path.dirname(real) != base or real == base:
            return {'ok': False, 'error': '安全檢查未過,拒絕刪除範圍外路徑'}
    removed_index = False
    if in_index:
        del idx[paper_id]
        _pc_write_json(INDEX_PATH, idx)
        removed_index = True
    removed_dir = False
    if dir_exists:
        import shutil as _shutil
        _shutil.rmtree(real)
        removed_dir = True
    return {'ok': True, 'id': paper_id, 'removed_index': removed_index,
            'removed_dir': removed_dir, 'remaining': len(idx)}



def verify(paper_id):
    """把所有段落照鍵的順序接回去，比對雜湊值，證明沒掉東西。
    注意：這裡的雜湊是『重組全文的 sha256』，和論文 id（原始檔位元組的 sha 前16碼）是兩碼事，不可互當。"""
    pid, cands = _resolve_pid(paper_id)
    if not pid:
        return {'ok': False, 'error': '找不到這篇（給的 id／標題無法唯一對應）', 'candidates': cands}
    data = _pc_read_json(os.path.join(_paper_dir(pid), 'text.json'))
    if not data:
        return {'ok': False, 'id': pid, 'error': '沒有這篇的切段全文'}
    segs = data.get('segments') or {}
    restored = (data.get('sep') or SEP).join(segs.values())
    got = hashlib.sha256(restored.encode('utf-8')).hexdigest()
    return {'ok': got == data.get('sha256'), 'id': pid, 'chars': len(restored),
            '全文雜湊_預期': data.get('sha256', '')[:12], '全文雜湊_實得': got[:12]}


# 命令列進入點置於檔尾：併入的工具函式都定義完了才輪到它執行。
if __name__ == '__main__':
    _argv = sys.argv[1:]
    if '--mcp' in _argv:
        _arcus_mcp_serve()
    elif '--index' in _argv:
        _print_index()
    elif '--spec' in _argv:
        _draw_main()
    else:
        print('usage: python arcus_core.py [--mcp | --index | --spec ...]')


# ══ 覆寫 cards()：每欄自動附上 refs 對應原文段落供溯源核對（2026-07-22） ══
def cards(ids=None, fields=None):
    """批次取卡片。每欄自動附上 refs 對應的原文段落（原文 key）供溯源核對。
    傳入的每個 id 可為 id 前綴或標題子字串，會先解析成真正的卡號。"""
    idx = _index_load()
    if ids:
        _rr = []
        for _g in ids:
            _p, _ = _resolve_pid(_g)
            _rr.append(_p or _g)
        ids = _rr
    ids = ids or list(idx.keys())
    fields = fields or FIELDS
    out = []
    for pid in ids:
        card = _pc_read_json(os.path.join(_paper_dir(pid), 'card.json'))
        if not card:
            out.append({'id': pid, 'ok': False, 'error': '沒有這篇的卡片'})
            continue
        meta = _pc_read_json(os.path.join(_paper_dir(pid), 'meta.json'), {}) or {}
        text_data = _pc_read_json(os.path.join(_paper_dir(pid), 'text.json')) or {}
        segs = text_data.get('segments') or {}
        欄位 = {}
        for f in fields:
            if f not in card:
                continue
            fval = card[f]
            if isinstance(fval, dict):
                refs = fval.get('refs') or []
                原文 = {k: segs[k] for k in refs if k in segs}
                欄位[f] = dict(fval, 原文=原文)
            else:
                欄位[f] = fval
        out.append({'id': pid, 'ok': True, '標題': meta.get('標題', ''), '欄位': 欄位})
    return {'ok': True, 'papers': out}
